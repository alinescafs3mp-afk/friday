"""Готовые файлы: Word, Excel, PDF, картинка.

Требование владельца (2026-08-01): «сделай мне отчёт с выводом по тем-то
документам» и подобные просьбы должны заканчиваться настоящим файлом, а не
текстом в чате.

Документ описывается СТРУКТУРОЙ, а не разметкой: заголовки, абзацы, списки,
таблицы. Одна и та же структура кладётся в любой из форматов — иначе для каждого
пришлось бы учить модель отдельному языку разметки, и три формата разошлись бы по
возможностям в первый же день.

Здесь нет ни базы, ни модели: на вход — что писать, на выход — байты. Откуда
взялось содержимое и правда ли оно, решают выше.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from typing import Any

__all__ = [
    "Block",
    "ReportSpec",
    "SUPPORTED_KINDS",
    "render",
    "sheet_title_from_report_title",
]

#: Что умеем отдать. Расширение — часть договора: по нему Telegram выбирает,
#: как показать файл, а человек — чем открыть.
SUPPORTED_KINDS: dict[str, tuple[str, str]] = {
    "docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"),
    "xlsx": ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx"),
    "pdf": ("application/pdf", "pdf"),
    "png": ("image/png", "png"),
}

#: Кириллица в PDF: встроенные шрифты reportlab её не знают, и текст молча
#: превращается в чёрные квадраты. DejaVu лежит в системе и покрывает кириллицу.
_PDF_FONT_CANDIDATES = (
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
)
#: Потолок высоты картинки. Примерно 130 строк — столько человек ещё
#: разглядывает; дальше формат перестаёт быть картинкой и становится обузой.
_PNG_MAX_HEIGHT = 4000
_PDF_BOLD_CANDIDATES = (
    "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
    "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf",
)


@dataclass(frozen=True)
class Block:
    """Кусок документа: заголовок, абзац, список или таблица.

    `kind` — одно из `heading`, `text`, `bullets`, `table`. Для таблицы `rows` —
    строки, первая считается шапкой.
    """

    kind: str
    text: str = ""
    items: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)


@dataclass(frozen=True)
class ReportSpec:
    title: str
    blocks: list[Block]
    subtitle: str = ""


def _as_blocks(raw: Any) -> list[Block]:
    """Структура от модели — в блоки, молча отбрасывая мусор.

    Модель ошибётся в форме рано или поздно; падать из-за лишнего ключа посреди
    отчёта хуже, чем пропустить непонятный кусок.
    """
    blocks: list[Block] = []
    if not isinstance(raw, list):
        return blocks
    for item in raw:
        if not isinstance(item, dict):
            continue
        kind = str(item.get("kind") or item.get("type") or "text").strip().casefold()
        if kind in {"heading", "header", "title"}:
            blocks.append(Block("heading", text=str(item.get("text") or "")))
        elif kind in {"bullets", "list", "items"}:
            source = item.get("items") or item.get("rows") or []
            items = [str(entry) for entry in source if str(entry).strip()] if isinstance(source, list) else []
            if items:
                blocks.append(Block("bullets", items=items))
        elif kind == "table":
            rows_raw = item.get("rows")
            rows = _table_rows(rows_raw)
            if rows:
                blocks.append(Block("table", rows=rows))
        else:
            text = str(item.get("text") or "")
            if text.strip():
                blocks.append(Block("text", text=text))
    return blocks


def _table_rows(raw: Any) -> list[list[str]]:
    """Строки таблицы, как бы модель их ни записала.

    Прежде брались только списки, а строки-словари ({"фамилия": …, "сумма": …})
    отбрасывались ВСЕ до одной: блок таблицы тихо исчезал, соседние оставались,
    и человек получал отчёт без главного — без самой таблицы, без единого следа
    в ответе. Словари — естественная форма для модели, и терять их нельзя.

    У словарей первая строка становится шапкой из ключей; порядок ключей берётся
    от первой строки, чтобы колонки не разъезжались.
    """
    if not isinstance(raw, list) or not raw:
        return []
    rows: list[list[str]] = []
    header: list[str] = []
    for item in raw:
        if isinstance(item, list):
            rows.append([str(cell) for cell in item])
        elif isinstance(item, dict):
            if not header:
                header = [str(key) for key in item]
                rows.append(header)
            rows.append([str(item.get(key, "")) for key in header])
        elif item is not None and str(item).strip():
            rows.append([str(item)])
    return rows


def spec_from_payload(title: str, subtitle: str, blocks: Any) -> ReportSpec:
    return ReportSpec(title=str(title or "Отчёт"), subtitle=str(subtitle or ""), blocks=_as_blocks(blocks))


def render(kind: str, spec: ReportSpec) -> bytes:
    kind = str(kind or "").strip().casefold()
    if kind not in SUPPORTED_KINDS:
        raise ValueError(f"Неизвестный формат: {kind!r}. Доступны: {', '.join(sorted(SUPPORTED_KINDS))}")
    if kind == "docx":
        return _render_docx(spec)
    if kind == "xlsx":
        return _render_xlsx(spec)
    if kind == "pdf":
        return _render_pdf(spec)
    return _render_png(spec)


def _render_docx(spec: ReportSpec) -> bytes:
    from docx import Document  # type: ignore[import-untyped]
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT  # type: ignore[import-untyped]
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
    from docx.oxml import OxmlElement  # type: ignore[import-untyped]
    from docx.oxml.ns import qn  # type: ignore[import-untyped]
    from docx.shared import Pt  # type: ignore[import-untyped]

    def shade(cell: Any, fill: str) -> None:
        properties = cell._tc.get_or_add_tcPr()
        element = properties.find(qn("w:shd"))
        if element is None:
            element = OxmlElement("w:shd")
            properties.append(element)
        element.set(qn("w:fill"), fill)

    def repeat_as_header(row: Any) -> None:
        properties = row._tr.get_or_add_trPr()
        element = OxmlElement("w:tblHeader")
        element.set(qn("w:val"), "true")
        properties.append(element)

    document = Document()
    document.add_heading(spec.title, level=0)
    if spec.subtitle:
        document.add_paragraph(spec.subtitle)
    for block in spec.blocks:
        if block.kind == "heading":
            document.add_heading(block.text, level=1)
        elif block.kind == "bullets":
            for item in block.items:
                document.add_paragraph(item, style="List Bullet")
        elif block.kind == "table" and any(block.rows):
            columns = max(len(row) for row in block.rows)
            table = document.add_table(rows=0, cols=columns)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for index, row in enumerate(block.rows):
                cells = table.add_row().cells
                for column in range(columns):
                    cells[column].text = row[column] if column < len(row) else ""
                    cells[column].vertical_alignment = (
                        WD_CELL_VERTICAL_ALIGNMENT.CENTER if index == 0 else (WD_CELL_VERTICAL_ALIGNMENT.TOP)
                    )
                    paragraph = cells[column].paragraphs[0]
                    paragraph.paragraph_format.space_after = Pt(0)
                    if index == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        shade(cells[column], "D9EAF7")
                        for run in paragraph.runs:
                            run.bold = True
                    elif index % 2 == 0:
                        shade(cells[column], "F4F8FB")
                if index == 0:
                    repeat_as_header(table.rows[-1])
        else:
            document.add_paragraph(block.text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_xlsx(spec: ReportSpec) -> bytes:
    from openpyxl import Workbook  # type: ignore[import-untyped]
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side  # type: ignore[import-untyped]
    from openpyxl.utils import get_column_letter  # type: ignore[import-untyped]

    book = Workbook()
    sheet = book.active
    # Имя листа пишет модель через заголовок отчёта, а Excel запрещает в нём
    # `\ / * ? : [ ]` — openpyxl на таком заголовке ПАДАЕТ, и человек вместо
    # файла получает сообщение об ошибке. «Отчёт: июль/август» — совершенно
    # обычная просьба.
    sheet.title = sheet_title_from_report_title(spec.title)
    _append_xlsx_literal_row(sheet, [spec.title])
    sheet["A1"].font = Font(bold=True, color="FFFFFF", size=14)
    sheet["A1"].fill = PatternFill("solid", fgColor="1F4E78")
    if spec.subtitle:
        _append_xlsx_literal_row(sheet, [spec.subtitle])
    sheet.append([])
    section_rows: set[int] = set()
    table_ranges: list[tuple[int, int, int]] = []
    for block in spec.blocks:
        if block.kind == "heading":
            _append_xlsx_literal_row(sheet, [block.text])
            sheet.cell(row=sheet.max_row, column=1).font = Font(bold=True, size=12)
            section_rows.add(sheet.max_row)
        elif block.kind == "bullets":
            for item in block.items:
                _append_xlsx_literal_row(sheet, [f"• {item}"])
        elif block.kind == "table" and any(block.rows):
            columns = max(len(row) for row in block.rows)
            start_row = 0
            for row in block.rows:
                padded = list(row[:columns]) + [""] * (columns - len(row))
                _append_xlsx_literal_row(sheet, padded)
                if not start_row:
                    start_row = sheet.max_row
            table_ranges.append((start_row, sheet.max_row, columns))
        else:
            _append_xlsx_literal_row(sheet, [block.text])
        sheet.append([])

    thin = Side(style="thin", color="B8C4CE")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    body_alignment = Alignment(wrap_text=True, vertical="top")
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in sheet.iter_rows():
        for cell in row:
            if cell.value is None:
                continue
            cell.alignment = body_alignment
            cell.border = border
    for row_index in section_rows:
        cell = sheet.cell(row=row_index, column=1)
        cell.fill = PatternFill("solid", fgColor="D9EAF7")
        cell.alignment = Alignment(vertical="center", wrap_text=True)
    for start_row, end_row, columns in table_ranges:
        for row_index in range(start_row, end_row + 1):
            for column in range(1, columns + 1):
                cell = sheet.cell(row=row_index, column=column)
                cell.border = border
                cell.alignment = body_alignment
                if row_index == start_row:
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill("solid", fgColor="1F4E78")
                    cell.alignment = header_alignment
                elif (row_index - start_row) % 2 == 0:
                    cell.fill = PatternFill("solid", fgColor="F4F8FB")

    # Ширина считается для каждой занятой колонки, а не только для A. Верхняя
    # граница оставляет таблицу обозримой: длинный текст переносится внутри
    # ячейки, а не растягивает лист на несколько экранов.
    for column in range(1, sheet.max_column + 1):
        longest = 0
        for cell in sheet[get_column_letter(column)]:
            if cell.value is None:
                continue
            longest = max(longest, max((len(line) for line in str(cell.value).splitlines()), default=0))
        sheet.column_dimensions[get_column_letter(column)].width = min(max(longest + 2, 12), 48)

    if table_ranges:
        # У листа один автофильтр, поэтому выбираем самую содержательную таблицу.
        start_row, end_row, columns = max(
            table_ranges,
            key=lambda bounds: ((bounds[1] - bounds[0] + 1) * bounds[2], bounds[2]),
        )
        sheet.freeze_panes = f"A{start_row + 1}"
        sheet.auto_filter.ref = f"A{start_row}:{get_column_letter(columns)}{end_row}"

    buffer = io.BytesIO()
    book.save(buffer)
    return buffer.getvalue()


#: Знаки, запрещённые Excel в имени листа.
_SHEET_FORBIDDEN = str.maketrans({char: " " for char in "\\/*?:[]"})


def _append_xlsx_literal_row(sheet: Any, values: list[str]) -> None:
    """Write report payload as literal cells, never as executable formulas.

    ``openpyxl`` interprets a string beginning with ``=`` as a formula.  Report
    content is model/data output, not workbook code; a formula such as
    ``WEBSERVICE`` could otherwise perform an outbound request when the person
    opens the attachment.  Setting the cell type after assignment preserves the
    visible text (without an apostrophe prefix) while serialising it as an
    inline string.  All strings are forced literal, so ``+``, ``-`` and ``@``
    are safe as well if the workbook is later converted to CSV.
    """

    sheet.append(list(values))
    for cell in sheet[sheet.max_row]:
        if isinstance(cell.value, str):
            cell.data_type = "s"


def sheet_title_from_report_title(title: str) -> str:
    """Заголовок отчёта — в допустимое имя листа (не длиннее 31 знака)."""
    cleaned = " ".join(str(title or "").translate(_SHEET_FORBIDDEN).split())
    return cleaned[:31].strip() or "Отчёт"


def _pdf_font() -> tuple[str, str]:
    """Имена зарегистрированных шрифтов: обычный и полужирный."""
    from pathlib import Path

    from reportlab.pdfbase import pdfmetrics  # type: ignore[import-untyped]
    from reportlab.pdfbase.ttfonts import TTFont  # type: ignore[import-untyped]

    regular = next((path for path in _PDF_FONT_CANDIDATES if Path(path).exists()), "")
    bold = next((path for path in _PDF_BOLD_CANDIDATES if Path(path).exists()), regular)
    if not regular:
        # Ни одного подходящего шрифта — лучше честная латиница Helvetica, чем
        # чёрные квадраты вместо русского текста.
        return "Helvetica", "Helvetica-Bold"
    if "FridayText" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("FridayText", regular))
        pdfmetrics.registerFont(TTFont("FridayText-Bold", bold or regular))
    return "FridayText", "FridayText-Bold"


def _pdf_table_widths(rows: list[list[str]], available_width: float) -> list[float]:
    """Fit every table column inside the printable page width.

    ReportLab otherwise derives width from unbroken content and a wide table can
    silently leave the page.  Relative content lengths still decide which
    columns get more room, while very long values wrap instead of monopolising
    the sheet.
    """

    columns = max((len(row) for row in rows), default=1)
    desired: list[float] = []
    for column in range(columns):
        longest = max(
            (
                max((len(line) for line in row[column].splitlines()), default=0)
                for row in rows
                if column < len(row)
            ),
            default=0,
        )
        desired.append((min(max(longest, 6), 42) * 4.6) + 12)
    total = sum(desired) or 1.0
    if total <= available_width:
        return desired
    scale = available_width / total
    return [width * scale for width in desired]


def _render_pdf(spec: ReportSpec) -> bytes:
    from reportlab.lib import colors  # type: ignore[import-untyped]
    from reportlab.lib.enums import TA_LEFT  # type: ignore[import-untyped]
    from reportlab.lib.pagesizes import A4, landscape  # type: ignore[import-untyped]
    from reportlab.lib.styles import ParagraphStyle  # type: ignore[import-untyped]
    from reportlab.lib.units import mm  # type: ignore[import-untyped]
    from reportlab.platypus import (  # type: ignore[import-untyped]
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    font, bold_font = _pdf_font()
    body = ParagraphStyle("body", fontName=font, fontSize=10.5, leading=15, alignment=TA_LEFT)
    heading = ParagraphStyle("heading", fontName=bold_font, fontSize=13, leading=18, spaceBefore=8)
    title_style = ParagraphStyle("title", fontName=bold_font, fontSize=18, leading=24, spaceAfter=6)
    table_body = ParagraphStyle("table-body", parent=body, fontSize=8.5, leading=11)
    table_header = ParagraphStyle(
        "table-header",
        parent=table_body,
        fontName=bold_font,
        textColor=colors.white,
    )

    widest_table = max(
        (max((len(row) for row in block.rows), default=0) for block in spec.blocks if block.kind == "table"),
        default=0,
    )
    page_size = landscape(A4) if widest_table >= 6 else A4

    buffer = io.BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=page_size,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=spec.title,
    )
    story: list[Any] = [Paragraph(_escape(spec.title), title_style)]
    if spec.subtitle:
        story.append(Paragraph(_escape(spec.subtitle), body))
    story.append(Spacer(1, 6))
    for block in spec.blocks:
        if block.kind == "heading":
            story.append(Paragraph(_escape(block.text), heading))
        elif block.kind == "bullets":
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(_escape(item), body)) for item in block.items],
                    bulletType="bullet",
                    start="•",
                )
            )
        elif block.kind == "table" and any(block.rows):
            columns = max(len(row) for row in block.rows)
            normalised = [list(row[:columns]) + [""] * (columns - len(row)) for row in block.rows]
            data = [
                [Paragraph(_escape(cell), table_header if row_index == 0 else table_body) for cell in row]
                for row_index, row in enumerate(normalised)
            ]
            table = Table(
                data,
                colWidths=_pdf_table_widths(normalised, document.width),
                repeatRows=1,
                hAlign="LEFT",
                splitByRow=1,
            )
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F4F8FB")]),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B8C4CE")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )
            story.append(table)
        else:
            story.append(Paragraph(_escape(block.text), body))
        story.append(Spacer(1, 6))
    document.build(story)
    return buffer.getvalue()


def _escape(text: str) -> str:
    """`<`, `&` — разметка для reportlab, а в тексте это обычные знаки."""
    return (
        str(text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
    )


@dataclass(frozen=True)
class _Line:
    """Строка картинки: текст либо ячейки таблицы с координатами колонок.

    Колонки нужны потому, что шрифт пропорциональный: склейка пробелами
    превращала таблицу в лесенку, по которой не прочитать, где чьё значение.
    """

    text: str | list[str]
    font: Any
    step: int
    columns: list[float] | None = None


def _render_png(spec: ReportSpec) -> bytes:
    """Картинка — тот же документ, нарисованный на холсте.

    Без внешних библиотек рисования: Pillow есть, а matplotlib нет, и тянуть его
    ради текстовой карточки не за что.
    """
    from PIL import Image, ImageDraw, ImageFont  # type: ignore[import-untyped]

    def _font(size: int, bold: bool = False) -> Any:
        candidates = _PDF_BOLD_CANDIDATES if bold else _PDF_FONT_CANDIDATES
        for path in candidates:
            try:
                return ImageFont.truetype(path, size)
            except OSError:
                continue
        return ImageFont.load_default()

    width, margin = 1200, 48
    lines: list[_Line] = [_Line(spec.title, _font(34, bold=True), 46)]
    if spec.subtitle:
        lines.append(_Line(spec.subtitle, _font(20), 30))
    lines.append(_Line("", _font(10), 12))
    body_font, head_font = _font(20), _font(24, bold=True)
    for block in spec.blocks:
        # Перенос нужен КАЖДОЙ ветке, а не только абзацам: длинный заголовок,
        # длинный пункт списка и широкая строка таблицы одинаково уезжали за
        # правый край картинки и обрывались на середине слова.
        if block.kind == "heading":
            lines.extend(_Line(chunk, head_font, 34) for chunk in _wrap(block.text, 64))
        elif block.kind == "bullets":
            for item in block.items:
                wrapped = _wrap(item, 74)
                lines.append(_Line(f"•  {wrapped[0]}", body_font, 28))
                lines.extend(_Line(f"   {chunk}", body_font, 28) for chunk in wrapped[1:])
        elif block.kind == "table":
            # Колонки — по координатам, а не через склейку пробелами: шрифт
            # пропорциональный, и «Тип Штук / Рапорты 42» превращалось в лесенку,
            # по которой не прочитать, где чьё значение. Ширина колонки — самая
            # широкая ячейка в ней, с отступом.
            rows = [[str(cell) for cell in row] for row in block.rows if row]
            if not rows:
                continue
            column_count = max(len(row) for row in rows)
            # Шапка меряется ЖИРНЫМ шрифтом, которым и рисуется: измерение
            # обычным давало колонку уже реального заголовка, и «Тип документа»
            # налезал на «Штук».
            widths = [
                max(
                    (
                        (head_font if row_index == 0 else body_font).getlength(row[index])
                        if index < len(row)
                        else 0.0
                    )
                    for row_index, row in enumerate(rows)
                )
                for index in range(column_count)
            ]
            offsets: list[float] = []
            running = 0.0
            for width_value in widths:
                offsets.append(running)
                running += width_value + 28
            for index, row in enumerate(rows):
                cells = [(row[column] if column < len(row) else "") for column in range(column_count)]
                lines.append(_Line(cells, head_font if index == 0 else body_font, 30, offsets))
        else:
            lines.extend(_Line(chunk, body_font, 28) for chunk in _wrap(block.text, 78))
        lines.append(_Line("", body_font, 10))

    # Картинка не может расти бесконечно. Замерено: 500 строк таблицы дают
    # 15164 пикселя и 1.1 МБ, 2000 строк — 60164 пикселя и 4.7 МБ. Telegram
    # такую не покажет, а человек просил «картинку со сводкой», а не файл,
    # который не открывается. Обрезаем — и ГОВОРИМ об этом прямо на картинке:
    # молчаливый обрез это ровно тот случай, который система уже чинила в
    # голосе и в разборе документов.
    # Место под саму оговорку резервируется заранее: иначе она выталкивает
    # картинку за собственный потолок.
    budget = _PNG_MAX_HEIGHT - margin * 2 - 40
    kept: list[_Line] = []
    used = 0
    for line in lines:
        if used + line.step > budget:
            break
        kept.append(line)
        used += line.step
    if len(kept) < len(lines):
        dropped = len(lines) - len(kept)
        kept.append(_Line("", body_font, 8))
        kept.append(
            _Line(
                f"…показано не всё: не поместилось строк — {dropped}. "
                "Попросите тот же отчёт в Word или Excel.",
                _font(18, bold=True),
                30,
            )
        )
        lines = kept
    else:
        lines = kept
    height = margin * 2 + sum(line.step for line in lines)
    image = Image.new("RGB", (width, max(height, 200)), "white")
    draw = ImageDraw.Draw(image)
    offset = margin
    for line in lines:
        if line.columns is not None and isinstance(line.text, list):
            for cell, column_offset in zip(line.text, line.columns, strict=False):
                if cell:
                    draw.text((margin + column_offset, offset), cell, font=line.font, fill=(24, 24, 24))
        elif isinstance(line.text, str) and line.text:
            draw.text((margin, offset), line.text, font=line.font, fill=(24, 24, 24))
        offset += line.step
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def _wrap(text: str, width: int) -> list[str]:
    words = str(text or "").split()
    if not words:
        return [""]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        if len(current) + 1 + len(word) <= width:
            current = f"{current} {word}"
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines
