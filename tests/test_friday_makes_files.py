"""«Сделай отчёт» заканчивается файлом, а не текстом в чате.

Требование владельца 2026-08-01: Пятница должна уметь составлять Word, Excel, PDF
и картинки — «сделай мне отчёт с выводом по тем-то документам» и подобное.

Проверяется не «функция вернула байты», а то, что файл ОТКРЫВАЕТСЯ штатной
библиотекой и русский текст в нём читается: PDF со встроенными шрифтами
reportlab молча превращает кириллицу в чёрные квадраты, и такой отчёт выглядит
готовым ровно до момента, когда его откроют.
"""

from __future__ import annotations

import io

import pytest

from friday.reports import SUPPORTED_KINDS, render, spec_from_payload

BLOCKS = [
    {"kind": "heading", "text": "Основные выводы"},
    {"kind": "text", "text": "По документам архива за июль 2026 года видно следующее."},
    {"kind": "bullets", "items": ["Штатное расписание обновлено 30 июля", "Приказ № 214 от 3 мая"]},
    {
        "kind": "table",
        "rows": [["Фамилия", "Месяц", "Начислено"], ["Бутко", "октябрь 2025", "87 450"]],
    },
]


def _spec():
    return spec_from_payload("Отчёт по документам", "Подготовлено Пятницей", BLOCKS)


@pytest.mark.parametrize("kind", sorted(SUPPORTED_KINDS))
def test_every_format_produces_a_file(kind):
    payload = render(kind, _spec())
    assert len(payload) > 1000, f"{kind}: файл подозрительно мал"


def test_the_word_file_opens_and_keeps_the_russian_text():
    import docx

    document = docx.Document(io.BytesIO(render("docx", _spec())))
    text = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    assert "Отчёт по документам" in text
    assert "Основные выводы" in text
    assert any("Штатное расписание" in line for line in text), "список не попал в документ"
    assert [cell.text for cell in document.tables[0].rows[0].cells] == [
        "Фамилия",
        "Месяц",
        "Начислено",
    ]


def test_the_excel_file_opens_and_keeps_the_table():
    import openpyxl

    book = openpyxl.load_workbook(io.BytesIO(render("xlsx", _spec())))
    sheet = book.active
    values = [str(sheet.cell(row=row, column=1).value) for row in range(1, sheet.max_row + 1)]
    assert "Отчёт по документам" in values
    assert "Фамилия" in values
    assert any("Бутко" in value for value in values)


def test_a_large_excel_table_is_complete_filterable_and_visually_structured():
    """A 300-row source remains real cells, including a late requested row."""

    import openpyxl

    rows = [["Позиция", "Фамилия", "Примечание"]]
    rows.extend(
        [str(position), f"Человек {position}", f"Проверяемая строка {position}: " + ("подробности " * 8)]
        for position in range(1, 301)
    )
    spec = spec_from_payload("Реестр", "", [{"kind": "table", "rows": rows}])

    book = openpyxl.load_workbook(io.BytesIO(render("xlsx", spec)))
    sheet = book.active
    header_row = next(cell.row for cell in sheet["A"] if cell.value == "Позиция")
    last_row = header_row + 300

    assert sheet.cell(header_row + 288, 1).value == "288"
    assert sheet.cell(header_row + 288, 2).value == "Человек 288"
    assert sheet.freeze_panes == f"A{header_row + 1}"
    assert sheet.auto_filter.ref == f"A{header_row}:C{last_row}"

    header = sheet.cell(header_row, 1)
    assert header.font.bold is True
    assert header.fill.fill_type == "solid"
    assert header.fill.fgColor.rgb.endswith("1F4E78")
    assert header.alignment.horizontal == "center"
    assert header.alignment.wrap_text is True

    for row in sheet.iter_rows(min_row=header_row, max_row=last_row, min_col=1, max_col=3):
        for cell in row:
            assert {
                cell.border.left.style,
                cell.border.right.style,
                cell.border.top.style,
                cell.border.bottom.style,
            } == {"thin"}
            assert cell.alignment.wrap_text is True
            assert cell.alignment.vertical == "top" or cell.row == header_row
    for column in ("A", "B", "C"):
        assert 12 <= sheet.column_dimensions[column].width <= 48


def test_excel_report_payload_is_always_literal_data():
    """Model/data strings must never turn into executable workbook formulas."""

    import openpyxl

    dangerous = [
        '=WEBSERVICE("https://example.invalid/private")',
        '=HYPERLINK("file:///private/path", "open")',
        "+SUM(1,2)",
        "-1+1",
        "@SUM(1,2)",
    ]
    spec = spec_from_payload(
        dangerous[0],
        dangerous[1],
        [
            {"kind": "heading", "text": dangerous[2]},
            {"kind": "bullets", "items": [dangerous[3]]},
            {"kind": "table", "rows": [[dangerous[4], dangerous[0]]]},
        ],
    )

    book = openpyxl.load_workbook(io.BytesIO(render("xlsx", spec)), data_only=False)
    cells = [cell for row in book.active.iter_rows() for cell in row if cell.value is not None]
    assert all(cell.data_type != "f" for cell in cells)
    assert set(dangerous).issubset({str(cell.value).removeprefix("• ") for cell in cells})


def test_the_word_table_has_a_repeating_shaded_header():
    import docx
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    document = docx.Document(io.BytesIO(render("docx", _spec())))
    table = document.tables[0]
    header = table.rows[0]

    assert header._tr.xpath("./w:trPr/w:tblHeader")
    for cell in header.cells:
        shading = cell._tc.xpath("./w:tcPr/w:shd")
        assert shading and shading[0].get(qn("w:fill")) == "D9EAF7"
        assert cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
        assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert all(run.bold for run in cell.paragraphs[0].runs)


def test_the_pdf_keeps_cyrillic_readable():
    """Мутация: убрать регистрацию DejaVu (вернуть Helvetica) — тест краснеет.

    Встроенные шрифты reportlab кириллицу не знают: текст верстается, файл
    открывается, а вместо букв — квадраты. Проверяется извлечённый текст, а не
    факт создания файла.
    """
    import pypdf

    reader = pypdf.PdfReader(io.BytesIO(render("pdf", _spec())))
    text = " ".join(reader.pages[0].extract_text().split())
    assert "Отчёт по документам" in text
    assert "Основные выводы" in text
    assert "Штатное расписание обновлено 30 июля" in text


def test_a_wide_multipage_pdf_table_fits_landscape_and_repeats_its_header():
    import pypdf

    header = [f"Колонка {column}" for column in range(1, 9)]
    rows = [header]
    rows.extend([f"Строка {position}, поле {column}" for column in range(1, 9)] for position in range(1, 121))
    spec = spec_from_payload("Большой реестр", "", [{"kind": "table", "rows": rows}])

    reader = pypdf.PdfReader(io.BytesIO(render("pdf", spec)))
    assert len(reader.pages) > 1
    assert float(reader.pages[0].mediabox.width) > float(reader.pages[0].mediabox.height)
    for page in reader.pages:
        text = " ".join((page.extract_text() or "").split())
        assert "Колонка 1" in text
        assert "Колонка 8" in text


def test_the_picture_is_a_real_image_with_room_for_the_text():
    from PIL import Image

    image = Image.open(io.BytesIO(render("png", _spec())))
    assert image.format == "PNG"
    assert image.width >= 800
    # Высота считается по содержимому: карточка, обрезанная на середине списка,
    # выглядит как сбой вёрстки.
    assert image.height > 300


def test_an_unknown_format_is_refused_by_name():
    with pytest.raises(ValueError, match="Неизвестный формат"):
        render("djvu", _spec())


def test_a_broken_block_does_not_break_the_report():
    """Модель ошибётся в форме рано или поздно.

    Падать из-за лишнего ключа посреди отчёта хуже, чем пропустить непонятный
    кусок: человек просил документ, а не сообщение об ошибке разбора.
    """
    spec = spec_from_payload(
        "Отчёт",
        "",
        [
            {"kind": "text", "text": "Первый абзац."},
            "строка вместо объекта",
            {"kind": "невиданное", "text": "Станет обычным абзацем."},
            {"kind": "table", "rows": "не список"},
            {"kind": "bullets", "items": []},
            {"kind": "text", "text": "Последний абзац."},
        ],
    )
    assert [block.kind for block in spec.blocks] == ["text", "text", "text"]
    assert len(render("docx", spec)) > 1000


def test_the_filename_cannot_escape_its_folder():
    """Заголовок пишет модель, а он становится именем файла."""
    from friday.execution_kernel import _safe_filename

    assert _safe_filename("../../etc/passwd", "pdf") == "etc passwd.pdf"
    assert _safe_filename("Отчёт: июль/август", "docx") == "Отчёт июль август.docx"
    assert _safe_filename("", "xlsx") == "Отчёт.xlsx"
    assert "/" not in _safe_filename("a/b/c", "png")


@pytest.mark.anyio
async def test_the_tool_returns_the_file_as_an_attachment(settings, storage):
    """Мутация: убрать `_attachment` из ответа — тест краснеет.

    Файл, который никуда не уходит, — это не выполненная просьба.
    """
    from friday.execution_kernel import ExecutionKernel
    from friday.permissions import ActorContext, AuthorizationService

    storage.ensure_user("alice", preset_key="admin")
    kernel = ExecutionKernel(AuthorizationService(storage), settings)
    kernel.bind_services(storage, None, None, None)
    actor = ActorContext(user_id="alice", preset_key="admin", source="test")

    result = await kernel.execute(
        "make_file",
        {"kind": "docx", "title": "Отчёт по документам", "blocks": BLOCKS},
        actor=actor,
    )

    assert result.success, result.error
    assert result.data["created"] is True
    assert result.data["filename"] == "Отчёт по документам.docx"
    assert result.attachment, "файл собран, но человеку не отправляется"
    assert result.attachment["kind"] == "document"
    assert result.attachment["mime_type"].endswith("wordprocessingml.document")
    assert result.attachment["content_base64"]


@pytest.mark.anyio
async def test_an_empty_report_is_refused_rather_than_sent_blank(settings, storage):
    from friday.execution_kernel import ExecutionKernel
    from friday.permissions import ActorContext, AuthorizationService

    storage.ensure_user("alice", preset_key="admin")
    kernel = ExecutionKernel(AuthorizationService(storage), settings)
    kernel.bind_services(storage, None, None, None)
    actor = ActorContext(user_id="alice", preset_key="admin", source="test")

    result = await kernel.execute("make_file", {"kind": "pdf", "title": "Пустой", "blocks": []}, actor=actor)
    assert result.data["created"] is False
    assert result.attachment is None


def test_a_promise_never_becomes_the_document():
    """Мутация: убрать фильтр `_IS_A_PROMISE` — тест краснеет.

    Замерено на живом экземпляре: файлы приходили с именами «Сейчас соберу
    сводку….docx», «Нашёл в личной базе.pdf», «Не удалось безопасно завершить
    вызов инструмента….png». Имя файла — первое, что видит человек, и по нему
    он решает, открывать ли.
    """
    from friday.agent_runtime import _blocks_from_text, _title_from_text

    answer = (
        "Сейчас соберу сводку и оформлю её в PDF.\n\n"
        "Сводка по базе знаний\n"
        "- Документов: 1533\n"
        "- Сущностей: 4608\n"
    )
    assert _title_from_text(answer) == "Сводка по базе знаний"
    texts = " ".join(str(block) for block in _blocks_from_text(answer))
    assert "Сейчас соберу" not in texts


def test_markdown_does_not_leak_into_the_document():
    """В Word и PDF `**Итого**` показывается вместе со звёздочками."""
    from friday.agent_runtime import _blocks_from_text

    blocks = _blocks_from_text("```\n**Итого:**\n- *Документов*: 1533\n---\n### Вывод\n")
    rendered = " ".join(str(block) for block in blocks)
    assert "**" not in rendered and "###" not in rendered and "```" not in rendered
    assert "Документов: 1533" in rendered


def test_the_title_falls_back_to_what_the_person_asked():
    """Когда ответа не получилось, заголовок берётся из просьбы."""
    from friday.agent_runtime import _title_from_request

    assert _title_from_request("сделай отчёт в word: сводка по базе знаний") == "Сводка по базе знаний"
    assert _title_from_request("оформи в excel таблицу с тегами базы")


def test_the_requested_format_is_honoured():
    from friday.agent_runtime import _file_kind_from_request

    assert _file_kind_from_request("сделай pdf со сводкой") == "pdf"
    assert _file_kind_from_request("оформи в excel") == "xlsx"
    assert _file_kind_from_request("сделай картинку со сводкой") == "png"
    assert _file_kind_from_request("сделай отчёт") == "docx"


def test_grounds_come_from_the_context_not_from_thin_air(settings, storage):
    """Мутация: убрать `_grounds_from_context` — тест краснеет.

    Замерено: без оснований второй заход сочинял «15 420 записей», «500 ГБ»,
    «10 миллионов уникальных записей» — при 1533 документах в архиве. Файл с
    выдуманными числами хуже отсутствия файла: его уносят и показывают другим.
    """
    from friday.agent_runtime import AgentContext, _grounds_from_context

    context = AgentContext(
        conversation_id="c1",
        user_id="alice",
        kb_size=1533,
        entity_count=4608,
        relation_count=0,
        knowledge_hits=[{"title": "Приказ 214", "snippet": "О назначении Хасанова"}],
    )
    grounds = _grounds_from_context(context)
    # Числа архива идут в основания только когда находок нет: в отчёте про
    # человека они читаются как характеристика человека. Проверка этого — в
    # tests/test_demo_screens_tell_the_truth.py.
    assert "Приказ 214" in grounds
    assert "О назначении Хасанова" in grounds

    empty = AgentContext(conversation_id="c1", user_id="alice", kb_size=1533, entity_count=4608)
    fallback = _grounds_from_context(empty)
    assert "1533" in fallback and "4608" in fallback


def test_a_sent_document_does_not_trigger_building_another_one():
    """Мутация: убрать проверку `synthetic_document_notice` — тест краснеет.

    Когда человек присылает файл без подписи, backend сочиняет за него текст
    «Загружен документ: отчёт.docx». Слово «отчёт» в ЧУЖОМ имени файла попадало
    под детектор просьбы о документе, и система собирала ещё один файл — на
    пустом месте, в ответ на присланный.

    Флаг `synthetic_document_notice` для этого и заведён: он уже отключает
    графовое расширение по той же причине — имя чужого файла не является
    просьбой человека.
    """
    import inspect

    from friday.agent_runtime import _ASKS_FOR_A_FILE, AgentRuntime

    # Сам детектор на такой строке срабатывает — в этом и была ловушка.
    assert _ASKS_FOR_A_FILE.search("Загружен документ: отчёт по проверке.docx")

    source = inspect.getsource(AgentRuntime.chat)
    guard = source[: source.index("_file_for_a_request_that_wanted_one(")]
    assert "not synthetic_document_notice" in guard, (
        "сгенерированное уведомление о файле судится как просьба человека"
    )


def test_a_list_item_never_becomes_the_document_title():
    """Мутация: убрать пропуск пунктов списка — тест краснеет.

    Замерено на живом архиве: отчёт по июльским документам получил имя
    «1 Рапорт на премии (файл Рапорт на премии май 2024 _ _ копия (2) docx).docx»
    — первая строка содержимого оказалась первым пунктом перечня. Внутри был
    настоящий отчёт, а имя выглядело как сбой.
    """
    from friday.agent_runtime import _title_from_text

    content = (
        "1. Рапорт на премии (файл: Рапорт на премии май 2024.docx)\n"
        "Суть: ходатайство о выплате премии личному составу.\n"
    )
    assert _title_from_text(content) == "Суть: ходатайство о выплате премии личному составу."

    # Маркеры тоже не заголовок.
    assert _title_from_text("- первый пункт\nОтчёт по документам за июль") == ("Отчёт по документам за июль")


def test_a_promise_in_the_present_tense_is_still_a_promise():
    """«Собираю отчёт…» — реплика в чате, а не название документа.

    В списке было «соберу», но не «собираю», и файл получил имя «Собираю отчёт по
    документам которые появились в архиве в июле 2026 года.docx».
    """
    from friday.agent_runtime import _title_from_text

    answer = "Собираю отчёт по документам за июль.\nОтчёт по июльским документам"
    assert _title_from_text(answer) == "Отчёт по июльским документам"
    for verb in ("Составляю", "Оформляю", "Готовлю", "Создаю", "Подготавливаю", "Делаю"):
        assert _title_from_text(f"{verb} документ сейчас.\nСводка за июль") == "Сводка за июль"


def test_the_title_comes_from_the_text_the_blocks_came_from():
    """Мутация: убрать `answer = clean` — тест краснеет.

    Блоки собирались вторым заходом, а заголовок брался из первой реплики: два
    разных текста, и в имя файла попадал не тот.
    """
    import inspect

    from friday.agent_runtime import AgentRuntime

    source = inspect.getsource(AgentRuntime._file_for_a_request_that_wanted_one)  # noqa: SLF001
    marker = source.index("blocks = _blocks_from_text(clean)")
    assert "answer = clean" in source[marker : marker + 500], (
        "заголовок берётся не из того текста, из которого собраны блоки"
    )


def test_a_clarifying_question_is_not_packed_into_a_document():
    """Мутация: убрать проверку `_answer_is_a_question` — тест краснеет.

    Замерено: на «сделай сводку в excel по рапортам» модель справедливо
    переспросила, каких именно, — и рантайм собрал файл с именем «Давай уточню
    что именно нужно собрать в Excel чтобы не выдумывать.xlsx».
    """
    from friday.agent_runtime import _answer_is_a_question

    assert _answer_is_a_question("Давай уточню, что именно собрать в Excel?") is True
    # Вопросительного знака может и не быть: замерено сквозным прогоном —
    # «Сначала уточню, что именно подразумевается под актом 77.» упаковалось в
    # документ, причём с заголовком из ЧУЖОГО документа архива.
    assert _answer_is_a_question("Сначала уточню, что именно подразумевается под актом 77.") is True
    assert _answer_is_a_question("Не совсем понятно, о каком акте речь.") is True
    assert _answer_is_a_question("Каких именно рапортов — на премии или на увольнение.") is True
    assert _answer_is_a_question("Каких рапортов — на премии или на увольнение?") is True
    # Ответ по существу, заканчивающийся вопросом, документом быть может.
    long_answer = (
        "Отчёт по документам за июль.\n"
        + "В архиве 42 рапорта, 17 приказов и 8 актов. " * 12
        + "\nНужно ли добавить сводку по подписантам?"
    )
    assert _answer_is_a_question(long_answer) is False
    assert _answer_is_a_question("Сводка за июль: 42 рапорта, 17 приказов.") is False
    assert _answer_is_a_question("") is False


def test_the_question_guard_is_wired_before_the_file_is_built():
    import inspect

    from friday.agent_runtime import AgentRuntime

    source = inspect.getsource(AgentRuntime._file_for_a_request_that_wanted_one)  # noqa: SLF001
    assert "_answer_is_a_question(answer)" in source
    assert source.index("_answer_is_a_question(answer)") < source.index("_blocks_from_text(answer)")


def test_a_table_in_a_picture_keeps_its_columns():
    """Мутация: вернуть склейку `"   ".join(row)` — тест краснеет.

    Шрифт пропорциональный, поэтому склеенная пробелами таблица превращалась в
    лесенку: «Тип Штук / Рапорты 42 / Приказы 17» — прочитать, где чьё значение,
    нельзя. Колонки рисуются по координатам, а ширина шапки меряется тем же
    жирным шрифтом, которым она рисуется (иначе заголовок налезает на соседа).
    """
    import inspect

    from friday import reports

    source = inspect.getsource(reports._render_png)  # noqa: SLF001
    assert '"   ".join(row)' not in source, "таблица снова склеивается пробелами"
    assert "head_font if row_index == 0 else body_font" in source, (
        "ширина шапки меряется не тем шрифтом, которым она рисуется"
    )

    # И картинка действительно строится: широкая таблица не роняет отрисовку.
    payload = reports.render(
        "png",
        reports.ReportSpec(
            title="Отчёт",
            blocks=[
                reports.Block(
                    "table",
                    rows=[
                        ["Тип документа", "Штук", "Ответственный"],
                        ["Рапорты", "42", "Проскурин В.А."],
                        ["Приказы", "17", "Шматов Р."],
                    ],
                )
            ],
        ),
    )
    assert payload[:8] == b"\x89PNG\r\n\x1a\n"
    import io

    from PIL import Image

    image = Image.open(io.BytesIO(payload))
    assert image.width == 1200
    assert image.height > 150


def test_a_ragged_table_does_not_break_any_format():
    """Строки разной длины — обычное дело у модели; падать на них нельзя."""
    from friday import reports

    spec = reports.ReportSpec(
        title="Неровная таблица",
        blocks=[
            reports.Block("table", rows=[["A", "B", "C"], ["1"], ["2", "3"], []]),
        ],
    )
    for kind in sorted(reports.SUPPORTED_KINDS):
        payload = reports.render(kind, spec)
        assert payload, f"{kind}: пустой файл"


@pytest.mark.parametrize("rows", [[], [[]]])
def test_an_empty_table_is_skipped_not_crashed(rows):
    from friday import reports

    spec = reports.ReportSpec(title="Пусто", blocks=[reports.Block("table", rows=rows)])
    for kind in sorted(reports.SUPPORTED_KINDS):
        assert reports.render(kind, spec)


def test_a_picture_that_cannot_hold_everything_says_so():
    """Мутация: убрать потолок высоты — тест краснеет.

    Замерено: 500 строк таблицы дают картинку 1200×15164 и 1.1 МБ, 2000 строк —
    1200×60164 и 4.7 МБ. Telegram такую не покажет, а человек просил картинку,
    а не файл, который не открывается. Обрезать молча нельзя — это ровно тот
    случай, который система уже чинила в голосе и в разборе документов.
    """
    import inspect
    import io

    from PIL import Image

    from friday import reports

    spec = reports.ReportSpec(
        title="Длинная сводка",
        blocks=[
            reports.Block("table", rows=[["№", "строка"]] + [[str(i), f"значение-{i}"] for i in range(300)])
        ],
    )
    payload = reports.render("png", spec)
    image = Image.open(io.BytesIO(payload))
    assert image.height <= reports._PNG_MAX_HEIGHT, "картинка растёт без предела"  # noqa: SLF001
    assert len(payload) < 1_000_000, "картинка тяжелее мегабайта"

    source = inspect.getsource(reports._render_png)  # noqa: SLF001
    assert "показано не всё" in source
    assert "Word или Excel" in source


def test_a_short_report_is_not_capped():
    """Контроль: обычный отчёт не обрастает предупреждением."""
    import io

    from PIL import Image

    from friday import reports

    spec = reports.ReportSpec(
        title="Короткая сводка",
        blocks=[reports.Block("table", rows=[["№", "строка"], ["1", "значение"]])],
    )
    image = Image.open(io.BytesIO(reports.render("png", spec)))
    assert image.height < 500
