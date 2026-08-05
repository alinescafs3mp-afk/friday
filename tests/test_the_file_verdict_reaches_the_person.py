"""Что система знает о присланном файле, человек должен узнать тоже.

Три дефекта одного клубка, найденные третьим аудитом:

- файл, из которого не вышло ни знака, продвигался в знание, и человеку
  говорили «✅ Файл стал знанием — можно спрашивать» (спрашивать не о чем:
  всё содержимое объекта — строка `[File: akt.docx; type=…; size=37211]`);
- предупреждение «текст извлечь не удалось» было физически недостижимо: оно
  проверяло верхнеуровневый ключ `extraction_success`, а приёмный путь кладёт
  исход разбора во вложенный словарь `extraction`;
- голос, который не распознался, молчал о том, что он не распознан, — ход
  превращался в «Загружен документ: telegram-voice-7.ogg».

Словарь ответа здесь НЕ собирается руками: ровно так предыдущие тесты и
оставались зелёными при нерабочем коде. Он берётся у настоящего `ingest_file`.
"""

from __future__ import annotations

import asyncio

import pytest

from friday.ingestion import IngestionPipeline
from friday.knowledge_graph import KnowledgeGraph
from friday.telegram_bridge._callbacks import _file_fate_line


@pytest.fixture
def pipeline(settings, storage):
    return IngestionPipeline(settings, storage, KnowledgeGraph(storage))


def _ingest(pipeline, data: bytes, filename: str, mime_type: str) -> dict:
    return asyncio.run(
        pipeline.ingest_file(
            "alice",
            None,
            data,
            filename=filename,
            mime_type=mime_type,
            source_ref=f"test:{filename}",
        )
    )


def test_a_file_with_no_text_is_never_called_knowledge(pipeline):
    """Мутация: убрать `assessment.action == "review"` из гейта — тест краснеет.

    Пустой .txt разбирается БЕЗ ошибки: `extraction_succeeded=True`, ассетов для
    vision нет, политика по умолчанию продвигать разрешает. Оценка при этом уже
    говорит «review» и объясняет почему — но гейт её не читал.
    """
    result = _ingest(pipeline, b"   \n  \n ", "пустой.txt", "text/plain")
    assert result["promoted"] is False, "файл без текста стал знанием"
    assert result["queued_for_review"] is True
    assert "✅ Файл стал знанием" not in _file_fate_line(result)
    assert "/inbox" in _file_fate_line(result)


def test_a_file_with_text_still_goes_through(pipeline):
    """Контроль: правка не должна отправлять в Inbox всё подряд.

    Без этой проверки «ничего не продвигать никогда» тоже прошло бы первый тест.
    """
    text = "Приказ №214 от 1 августа 2026 года. Ответственный — Проскурин В.А.".encode()
    result = _ingest(pipeline, text, "приказ.txt", "text/plain")
    assert result["extraction"]["success"] is True
    assert result["extraction"]["text_success"] is True
    # Продвижение зависит от политики установки; здесь важно, что вердикт оценки
    # не «review» — то есть новая ветка гейта на нормальном файле не срабатывает.
    raw = pipeline.storage.get_raw_object(result["raw_object_id"], "alice")
    import json

    metadata = json.loads(str(raw.get("metadata_json") or "{}"))
    assert metadata.get("promotion_assessment", {}).get("action") != "review"


def test_the_bridge_can_actually_say_the_text_was_not_extracted(pipeline):
    """Мутация: вернуть проверку верхнеуровневого `extraction_success` — краснеет.

    Проверяется на ФОРМЕ, которую производит продакшен, а не на собранной руками:
    прежний тест собирал словарь с ключом `extraction_success`, которого приёмный
    путь не возвращает никогда, и оставался зелёным при мёртвой ветке.
    """
    result = _ingest(pipeline, b"\x89PNG\r\n\x1a\n" + b"\x00" * 64, "qr.png", "image/png")
    assert "extraction_success" not in result, (
        "форма ответа изменилась — тест снова проверяет не то, что приходит"
    )
    assert result["extraction"]["success"] is False
    line = _file_fate_line(result)
    assert "Текст извлечь не удалось" in line, "предупреждение по-прежнему недостижимо"
    assert "/inbox" in line


def test_an_unrecognised_voice_says_so():
    """Мутация: убрать флаг в server.py или ветку в мосте — тест краснеет.

    Человек наговорил вопрос, Whisper не разобрал ни слова. До правки ход
    превращался в «Загружен документ: telegram-voice-7.ogg» — про нераспознанное
    ни слова.
    """
    line = _file_fate_line({"queued_for_review": True, "voice_unrecognised": True})
    assert "Голос не распознался" in line
    assert "Повторите текстом" in line


def test_the_server_marks_an_unrecognised_voice():
    """Флаг ставится там, где известен и вид вложения, и наличие транскрипта."""
    import inspect

    from friday import server

    source = inspect.getsource(server)
    assert 'file_ingestion["voice_unrecognised"] = True' in source
    marker = source.index('file_ingestion["voice_unrecognised"] = True')
    guard = source[max(0, marker - 400) : marker]
    assert "is_voice and not transcript" in guard, "флаг ставится не по факту «голос без транскрипта»"


def test_a_parse_stopped_by_the_deadline_says_only_the_beginning_was_read():
    """`parse_deadline_reached` писался в ответ и не читался ни одним потребителем.

    Успех и полнота — разные вещи: разбор, оборванный по сроку, приходит с
    `success=True` и частичным текстом, и человек узнавал «файл принят» без
    единого слова о том, что принято лишь начало.
    """
    partial = {
        "queued_for_review": True,
        "extraction": {
            "success": True,
            "text_success": True,
            "parse_deadline_reached": True,
            "parse_pages_read": 12,
        },
    }
    line = _file_fate_line(partial)
    assert "принято только начало" in line
    assert "12" in line

    promoted = dict(partial, promoted=True, queued_for_review=False)
    promoted_line = _file_fate_line(promoted)
    assert "стал знанием" in promoted_line
    assert "принято только начало" in promoted_line, "продвинутый файл тоже мог быть прочитан наполовину"


def test_a_whole_file_says_nothing_extra():
    """Контроль: обычный разбор не обрастает оговорками."""
    line = _file_fate_line({"promoted": True, "extraction": {"success": True, "text_success": True}})
    assert line == "✅ Файл стал знанием — можно спрашивать."


def test_the_same_file_sent_twice_is_accepted_once(pipeline):
    """Мутация: убрать запасной ключ по содержимому — тест краснеет.

    Мост строил `source_ref` из `update_id`, уникального у КАЖДОЙ отправки,
    поэтому один и тот же файл не совпадал сам с собой никогда. Замерено: одна и
    та же строка байт под двумя ключами дала два Raw Object с одинаковым
    content_hash, два элемента Inbox и два одинаковых Knowledge Object. Файл на
    диске один — задваивались очередь разбора и корпус.
    """
    payload = "Договор №7 от 1 августа 2026 года. Стороны: ООО «Заря» и ИП Кузнецов.".encode()
    first = asyncio.run(
        pipeline.ingest_file(
            "alice",
            None,
            payload,
            filename="dogovor.txt",
            mime_type="text/plain",
            source_ref="telegram-file:AAA:111",
        )
    )
    second = asyncio.run(
        pipeline.ingest_file(
            "alice",
            None,
            payload,
            filename="dogovor.txt",
            mime_type="text/plain",
            source_ref="telegram-file:BBB:222",  # другая отправка того же файла
        )
    )
    assert second.get("idempotent_replay") is True, "тот же файл принят вторым объектом"
    assert second["raw_object_id"] == first["raw_object_id"]

    rows = pipeline.storage.execute(
        "SELECT COUNT(*) AS c FROM raw_objects WHERE user_id='alice' AND content_type='file'"
    ).fetchone()["c"]
    assert rows == 1, f"в архиве {rows} записи об одном файле"


def test_a_different_file_is_still_a_different_file(pipeline):
    """Контроль: дедуп по содержимому не склеивает разные документы."""
    one = asyncio.run(
        pipeline.ingest_file(
            "alice",
            None,
            b"first document body",
            filename="a.txt",
            mime_type="text/plain",
            source_ref="telegram-file:AAA:1",
        )
    )
    two = asyncio.run(
        pipeline.ingest_file(
            "alice",
            None,
            b"second document body",
            filename="b.txt",
            mime_type="text/plain",
            source_ref="telegram-file:BBB:2",
        )
    )
    assert two.get("idempotent_replay") is not True
    assert one["raw_object_id"] != two["raw_object_id"]


def test_one_persons_file_is_not_another_persons_replay(pipeline):
    """Дедуп по содержимому — в границах одного человека, а не по всей базе."""
    payload = b"shared text that two people happen to send"
    mine = asyncio.run(
        pipeline.ingest_file(
            "alice", None, payload, filename="общий.txt", mime_type="text/plain", source_ref="a:1"
        )
    )
    theirs = asyncio.run(
        pipeline.ingest_file(
            "bob", None, payload, filename="общий.txt", mime_type="text/plain", source_ref="b:1"
        )
    )
    assert theirs.get("idempotent_replay") is not True, "чужой файл воспроизведён как свой"
    assert theirs["raw_object_id"] != mine["raw_object_id"]


def _docx(paragraphs: list[str], *, header: str = "", table: list[list[str]] | None = None) -> bytes:
    import io

    from docx import Document

    document = Document()
    if header:
        document.sections[0].header.paragraphs[0].text = header
    for text in paragraphs:
        document.add_paragraph(text)
    if table:
        grid = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for column_index, value in enumerate(row):
                grid.cell(row_index, column_index).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _xlsx(rows: list[list[str]], *, sheet: str = "Лист") -> bytes:
    import io

    from openpyxl import Workbook

    book = Workbook()
    page = book.active
    page.title = sheet
    for row in rows:
        page.append(row)
    buffer = io.BytesIO()
    book.save(buffer)
    return buffer.getvalue()


def test_a_word_document_gives_up_its_paragraphs_and_its_table(pipeline):
    """Таблица — часть документа, а не украшение: в приказах в ней самое важное."""
    payload = _docx(
        ["С 1 августа назначить ответственным Проскурина В.А.", "Срок — 15 августа 2026 года."],
        table=[["Должность", "ФИО"], ["Начальник", "Проскурин В.А."]],
    )
    result = _ingest(pipeline, payload, "prikaz.docx", "")
    raw = pipeline.storage.get_raw_object(result["raw_object_id"], "alice")
    body = " ".join(str(raw.get("raw_content") or "").split())
    assert "Проскурина В.А." in body
    assert "15 августа 2026" in body
    assert "Начальник" in body and "ФИО" in body, "таблица не попала в текст"


def test_a_spreadsheet_gives_up_every_row(pipeline):
    payload = _xlsx(
        [
            ["Дата", "ФИО", "Тип"],
            ["2026-07-11", "Зайцев М.", "увольнение"],
            ["2026-08-07", "Целио Ю.Ю.", "выходные"],
        ],
        sheet="Рапорты",
    )
    result = _ingest(pipeline, payload, "raporty.xlsx", "")
    raw = pipeline.storage.get_raw_object(result["raw_object_id"], "alice")
    body = " ".join(str(raw.get("raw_content") or "").split())
    assert "Рапорты" in body
    assert "Зайцев М." in body and "Целио Ю.Ю." in body, "часть строк потеряна"
    assert "2026-08-07" in body


def test_an_archive_gives_up_what_is_inside(pipeline):
    import io
    import zipfile

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("inner.txt", "Смета на поверку весов: 42 000 рублей.")
    result = _ingest(pipeline, buffer.getvalue(), "arhiv.zip", "")
    raw = pipeline.storage.get_raw_object(result["raw_object_id"], "alice")
    body = " ".join(str(raw.get("raw_content") or "").split())
    assert "42 000 рублей" in body
    assert "inner.txt" in body


def test_a_document_whose_text_lives_in_the_header_says_so(pipeline):
    """Мутация: убрать `chars` из ответа приёма — тест краснеет.

    Разбор без ошибки — ещё не текст. Пустой .txt и .docx, где всё написанное
    лежит в колонтитуле, приходят с `success=True` и нулём знаков; человеку
    говорили просто «ждёт разбора», и он не знал, что содержимого не видно.
    """
    result = _ingest(pipeline, _docx([], header="Только колонтитул"), "header.docx", "")
    assert result["extraction"]["success"] is True, "разбор должен был пройти без ошибки"
    assert result["extraction"]["chars"] == 0
    assert result["promoted"] is False
    line = _file_fate_line(result)
    assert "Текста в файле не оказалось" in line
    # Это НЕ то же самое, что «разобрать не удалось»: там файл нечитаем вовсе.
    assert "Текст извлечь не удалось" not in line


@pytest.mark.parametrize(
    "name,payload",
    [
        ("broken.docx", b"PK\x03\x04" + b"\x00" * 200),
        ("fake.doc", b"not really a word file"),
        ("qr.png", b"\x89PNG\r\n\x1a\n" + b"\x00" * 64),
    ],
)
def test_an_unreadable_file_says_the_content_is_invisible(pipeline, name, payload):
    result = _ingest(pipeline, payload, name, "")
    assert result["promoted"] is False
    assert "Текст извлечь не удалось" in _file_fate_line(result)


def test_a_document_too_long_to_hold_says_only_the_beginning_was_taken(pipeline):
    """Мутация: брать `text_truncated` не из метаданных разборщика — тест краснеет.

    Замерено: документ на 3.75 млн знаков принимался как 2 млн, и человеку
    говорилось ровно «✅ Файл стал знанием — можно спрашивать». Половина текста
    отброшена, спрашивать по ней бесполезно, и узнать об этом было неоткуда —
    тот же класс, что немой обрыв голоса и разбор, оборванный по сроку.

    Первая редакция правки вычисляла признак на приёмной стороне и всегда
    получала False: разборщик к тому моменту уже обрезал текст своим потолком,
    и мерилось обрезанное.
    """
    huge = ("Строка отчёта. " * 250_000).encode()
    result = _ingest(pipeline, huge, "huge.txt", "text/plain")
    extraction = result["extraction"]
    assert extraction["text_truncated"] is True
    assert extraction["chars"] < len(huge.decode())
    line = _file_fate_line(result)
    assert "принято начало" in line
    assert "по концу файла спрашивать бесполезно" in line


def test_an_ordinary_document_carries_no_truncation_notice(pipeline):
    result = _ingest(pipeline, "Приказ №5 от 1 августа.".encode(), "small.txt", "text/plain")
    assert result["extraction"]["text_truncated"] is False
    assert _file_fate_line(result) == "✅ Файл стал знанием — можно спрашивать."
