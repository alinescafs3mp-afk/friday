"""DOCX completeness treats only proven content-free paragraphs as neutral."""

from __future__ import annotations

import io
from typing import Any

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from friday.agent_runtime._office_attachments import (
    OFFICE_STRUCTURE_KEY,
    build_office_prompt_bundle,
    trusted_office_attachment,
)
from friday.documents import DocumentExtractor, validate_office_structure_index


def _docx_bytes(document: Document) -> bytes:
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _add_table(document: Document, *, rows: int, row_offset: int = 0) -> None:
    table = document.add_table(rows=rows, cols=4)
    for row_number, row in enumerate(table.rows, start=row_offset + 1):
        for column_number, cell in enumerate(row.cells, start=1):
            cell.text = f"R{row_number:02d}C{column_number}"


def _neutral_marker_paragraph(document: Document) -> Any:
    paragraph = document.add_paragraph()
    # The audit did not retain the live OOXML.  ``w:ins`` in paragraph-mark
    # properties is the synthetic mutation that proves the old blanket
    # unsupported-container scan; proofing/permission markers cover common
    # content-free siblings without claiming that exact live shape.
    paragraph._p.get_or_add_pPr().append(parse_xml(f'<w:rPr {nsdecls("w")}><w:ins w:id="7"/></w:rPr>'))
    paragraph._p.append(parse_xml(f'<w:proofErr {nsdecls("w")} w:type="spellStart"/>'))
    paragraph._p.append(parse_xml(f'<w:permStart {nsdecls("w")} w:id="23" w:edGrp="everyone"/>'))
    paragraph._p.append(parse_xml(f'<w:permEnd {nsdecls("w")} w:id="23"/>'))
    assert paragraph.text == ""
    assert paragraph.runs == []
    return paragraph


def _extract(content: bytes) -> tuple[str, dict[str, Any]]:
    result = DocumentExtractor().extract(content, "synthetic-empty-paragraph.docx")
    index = result.office_structure_index
    assert result.success is True
    assert isinstance(index, dict)
    assert validate_office_structure_index(index, result.text) == index
    return result.text, index


def _assert_exact_table_inventory(index: dict[str, Any], *, blocks: int) -> None:
    coverage = index["coverage"]
    assert coverage["blocks_seen"] == coverage["blocks_indexed"] == blocks
    assert coverage["rows_seen"] == coverage["rows_indexed"] == 12
    assert coverage["cells_seen"] == coverage["cells_indexed"] == 48


def test_live_like_trailing_structural_paragraph_is_neutral() -> None:
    document = Document()
    _add_table(document, rows=12)
    _neutral_marker_paragraph(document)

    text, index = _extract(_docx_bytes(document))

    assert index["complete"] is True
    assert index["coverage"]["reasons"] == []
    assert [block["kind"] for block in index["blocks"]] == ["table"]
    assert [block["source_order"] for block in index["blocks"]] == [0]
    _assert_exact_table_inventory(index, blocks=1)

    bundle = build_office_prompt_bundle(
        [
            trusted_office_attachment(
                {
                    "filename": "synthetic-empty-paragraph.docx",
                    "transient_text": text,
                    "extraction_success": True,
                    OFFICE_STRUCTURE_KEY: index,
                }
            )
        ],
        max_chars=100_000,
    )
    assert bundle is not None
    assert bundle.views[0]["index_complete"] is True
    assert bundle.views[0]["prompt_complete"] is True
    assert bundle.views[0]["coverage_reasons"] == []
    assert len(bundle.views[0]["atoms"]) == 12


def test_structurally_empty_paragraph_between_tables_is_not_a_body_atom() -> None:
    document = Document()
    _add_table(document, rows=6)
    _neutral_marker_paragraph(document)
    _add_table(document, rows=6, row_offset=6)

    _text, index = _extract(_docx_bytes(document))

    assert index["complete"] is True
    assert index["coverage"]["reasons"] == []
    assert [block["kind"] for block in index["blocks"]] == ["table", "table"]
    assert [block["source_order"] for block in index["blocks"]] == [0, 1]
    _assert_exact_table_inventory(index, blocks=2)


def _meaningful_blank_fragment(kind: str) -> str:
    word_namespace = nsdecls("w")
    if kind == "drawing":
        return f"<w:r {word_namespace}><w:drawing/></w:r>"
    if kind == "field":
        return f'<w:fldSimple {word_namespace} w:instr="DATE"/>'
    if kind == "object":
        return f"<w:r {word_namespace}><w:object/></w:r>"
    if kind == "bookmark":
        return (
            f'<w:bookmarkStart {word_namespace} w:id="0" w:name="SYNTHETIC"/>'
            '<w:bookmarkEnd xmlns:w="http://schemas.openxmlformats.org/'
            'wordprocessingml/2006/main" w:id="0"/>'
        )
    if kind == "hyperlink":
        return f'<w:hyperlink {nsdecls("w", "r")} r:id="rId999"/>'
    if kind == "empty_run":
        return f"<w:r {word_namespace}/>"
    if kind == "nested_sdt":
        return f"<w:rPr {word_namespace}><w:sdt/></w:rPr>"
    if kind == "unknown_control":
        return '<synthetic:control xmlns:synthetic="urn:friday:synthetic:unknown"/>'
    if kind == "unknown_word_tag":
        return f"<w:syntheticControl {word_namespace}/>"
    if kind == "property_relationship":
        return f'<w:rPr {nsdecls("w", "r")} r:id="rId999"/>'
    raise AssertionError(f"unknown synthetic kind: {kind}")


@pytest.mark.parametrize(
    "kind",
    [
        "drawing",
        "field",
        "object",
        "bookmark",
        "hyperlink",
        "empty_run",
        "nested_sdt",
        "unknown_control",
        "unknown_word_tag",
        "property_relationship",
    ],
)
def test_blank_paragraph_with_meaningful_structure_remains_incomplete(kind: str) -> None:
    document = Document()
    _add_table(document, rows=6)
    paragraph = document.add_paragraph()
    fragment = _meaningful_blank_fragment(kind)
    if kind == "bookmark":
        namespace_end = fragment.index("<w:bookmarkEnd")
        paragraph._p.append(parse_xml(fragment[:namespace_end]))
        paragraph._p.append(parse_xml(fragment[namespace_end:]))
    elif kind in {"nested_sdt", "unknown_control", "unknown_word_tag", "property_relationship"}:
        paragraph._p.get_or_add_pPr().append(parse_xml(fragment))
    else:
        paragraph._p.append(parse_xml(fragment))
    _add_table(document, rows=6, row_offset=6)

    _text, index = _extract(_docx_bytes(document))

    assert index["complete"] is False
    assert index["coverage"]["reasons"] == ["unsupported_body_content"]
    assert [block["kind"] for block in index["blocks"]] == ["table", "paragraph", "table"]
    assert [block["source_order"] for block in index["blocks"]] == [0, 1, 2]
    _assert_exact_table_inventory(index, blocks=3)
