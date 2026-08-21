"""Synthetic end-to-end regressions for the frozen JBL file contracts.

No fixture in this module reads a live conversation or invokes a model/network
provider.  The fakes expose only the bounded capabilities required by each
contract and fail on every unexpected call.
"""

from __future__ import annotations

import copy
import hashlib
import io
import json
from dataclasses import replace
from datetime import UTC, datetime, timedelta
from types import SimpleNamespace
from typing import Any

import pytest
from docx import Document

from friday.agent_runtime import (
    _FALSE_CURRENT_MODEL_OUTAGE,
    _PERSON_DOCUMENT_INVENTORY,
    _WEB_EVIDENCE_MISSING,
    _WEB_ISOLATION_DEICTIC,
    AgentContext,
    AgentRuntime,
    _attachment_evidence_chunks,
    _attachment_web_fact_targets,
    _attachment_web_literals_are_grounded,
    _multi_attachment_summary_count,
    _OwnedAttachment,
    _project_attachments_for_request,
    _reconcile_attachment_web_literals,
    _valid_person_document_inventory_data,
    asks_for_the_web,
)
from friday.agent_runtime._office_attachments import (
    OFFICE_EXACT_UNAVAILABLE_MESSAGE,
    OFFICE_STRUCTURE_KEY,
    trusted_office_attachment,
    validate_runtime_office_index,
)
from friday.documents import DocumentExtractor
from friday.execution_kernel import ToolResult
from friday.permissions import ActorContext
from friday.server import _current_turn_file_attachment
from friday.storage.models import RawObject, new_id
from friday.telegram_bridge._markup import to_telegram_html


def _actor() -> ActorContext:
    return ActorContext(user_id="alice", preset_key="owner", source="test")


def _transient_attachment(*, filename: str, text: str) -> dict[str, Any]:
    """Project current no-save bytes through the server-owned carrier."""

    return _current_turn_file_attachment(
        filename=filename,
        file_ingestion={
            "extraction": {
                "success": True,
                "text_success": True,
                "chars": len(text),
            }
        },
        raw={
            "raw_content": text,
            "metadata_json": {
                "filename": filename,
                "uploaded_by": "alice",
                "extraction_success": True,
                "text_extraction_success": True,
            },
        },
    )


class _AllowAll:
    def authorize(self, actor, capability, **kwargs):  # noqa: ANN001, ARG002
        return SimpleNamespace(allowed=True)


def _tool(name: str) -> dict[str, Any]:
    return {
        "type": "function",
        "function": {
            "name": name,
            "description": "synthetic bounded read",
            "parameters": {"type": "object"},
        },
    }


async def _prepare_without_retrieval(
    user_id: str,
    _message: str,
    conversation_id: str,
    **kwargs: Any,
) -> AgentContext:
    history = list(kwargs.get("prior_history") or [])
    previous_user = next(
        (
            str(item.get("content") or "")
            for item in reversed(history)
            if str(item.get("role") or "") == "user"
        ),
        "",
    )
    return AgentContext(
        conversation_id=conversation_id,
        user_id=user_id,
        person_id=str(kwargs.get("person_id") or user_id),
        conversation_history=history,
        previous_user_turn=previous_user,
    )


class _NeverModel:
    enabled = True
    total_budget_sec = 1.0

    async def chat(self, messages, tools=None, **kwargs):  # noqa: ANN001, ARG002
        raise AssertionError("the code-owned turn reached a model")


class _InventoryKernel:
    authorization = _AllowAll()

    def __init__(
        self,
        *,
        available: bool = True,
        malformed: bool = False,
        person_name: str = "JBL",
    ) -> None:
        self.available = available
        self.malformed = malformed
        self.person_name = person_name
        self.calls: list[dict[str, Any]] = []

    def get_tool_definitions(self, actor, topic=""):  # noqa: ANN001, ARG002
        return [_tool("user_activity")] if self.available else []

    async def execute(self, tool, params, actor=None):  # noqa: ANN001, ARG002
        assert tool == "user_activity"
        assert params["documents_only"] is True
        assert params["offset"] == 0
        self.calls.append(dict(params))
        if self.malformed:
            return ToolResult(
                tool,
                True,
                {
                    "человек": self.person_name,
                    "период": {"с": params["since"], "по": params["until"]},
                    "документов с подтверждённым автором": 0,
                    "документов без отметки автора": 0,
                    "документы": [],
                    "пагинация": {},
                },
            )
        return ToolResult(
            tool,
            True,
            {
                "человек": self.person_name,
                "период": {"с": params["since"], "по": params["until"]},
                "документов с подтверждённым автором": 2,
                "документов без отметки автора": 0,
                "документы": [
                    {"что": "alpha.pdf", "когда": params["since"]},
                    {"что": "beta.docx", "когда": params["until"]},
                ],
                "пагинация": {
                    "смещение": 0,
                    "показано": 2,
                    "из подтверждённых": 2,
                    "следующее смещение": None,
                    "подтверждённый перечень показан полностью": True,
                },
            },
        )


class _InventoryWithMessageSearchKernel(_InventoryKernel):
    """Expose the tempting history route, but reject any attempt to use it."""

    def get_tool_definitions(self, actor, topic=""):  # noqa: ANN001, ARG002
        return [_tool("user_activity"), _tool("message_search")]

    async def execute(self, tool, params, actor=None):  # noqa: ANN001, ARG002
        if tool == "message_search":
            raise AssertionError("the compact upload inventory reached message_search")
        return await super().execute(tool, params, actor=actor)


@pytest.mark.asyncio
async def test_named_day_inventory_and_its_completeness_followup_are_code_owned(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner", display_name="Owner")
    storage.ensure_user("jbl", preset_key="user", display_name="JBL")
    kernel = _InventoryKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_NeverModel(),
        kernel=kernel,
    )
    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)

    first = await runtime.chat(
        "alice",
        "Какие документы вчера загружал JBL?",
        actor=_actor(),
    )
    second = await runtime.chat(
        "alice",
        "И всё?",
        actor=_actor(),
        conversation_id=first["conversation_id"],
    )

    assert len(kernel.calls) == 2
    assert kernel.calls[0] == kernel.calls[1]
    for reply in (first, second):
        assert "alpha.pdf" in reply["message"] and "beta.docx" in reply["message"]
        assert "2 из 2" in reply["message"]
        assert reply["tools_used"] == ["user_activity"]
    assert "Проверила выборку повторно" in second["message"]
    stored = storage.get_message(str(second["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    assert metadata["structural"]["person_document_inventory"] is True


@pytest.mark.asyncio
async def test_named_inventory_without_a_date_is_all_time_but_temporal_cues_stay_closed(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner", display_name="Owner")
    storage.ensure_user("jbl", preset_key="user", display_name="JBL")
    kernel = _InventoryKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_NeverModel(),
        kernel=kernel,
    )
    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)
    catalog_calls: list[tuple[str, str, int]] = []

    def body_free_catalog(user_id: str, uploaded_by: str, *, limit: int = 5_000):  # noqa: ANN202
        catalog_calls.append((user_id, uploaded_by, limit))
        return [
            {"id": "raw_alpha", "filename": "alpha.pdf", "content_type": "file"},
            {"id": "raw_beta", "filename": "beta.docx", "content_type": "file"},
        ]

    monkeypatch.setattr(storage, "list_owned_file_catalog", body_free_catalog)

    all_time = await runtime.chat(
        "alice",
        "Какие файлы загружал JBL",
        actor=_actor(),
    )
    explicit_all_time = await runtime.chat(
        "alice",
        "Какие документы за всё время загружал JBL?",
        actor=_actor(),
    )
    calendar_month = await runtime.chat(
        "alice",
        "Какие файлы загружал JBL в июле?",
        actor=_actor(),
    )

    assert len(kernel.calls) == 1
    assert kernel.calls[0]["person"] == "jbl"
    assert catalog_calls == [("jbl", "jbl", 5_001), ("jbl", "jbl", 5_001)]
    for reply in (all_time, explicit_all_time):
        assert "за всё время" in reply["message"].casefold()
        assert "alpha.pdf" in reply["message"] and "beta.docx" in reply["message"]
        assert reply["tools_used"] == []
    assert "alpha.pdf" in calendar_month["message"] and "2 из 2" in calendar_month["message"]
    assert calendar_month["tools_used"] == ["user_activity"]


@pytest.mark.asyncio
async def test_named_inventory_relative_day_followup_stays_code_owned(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner", display_name="Owner")
    storage.ensure_user("jbl", preset_key="user", display_name="JBL")
    kernel = _InventoryKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_NeverModel(),
        kernel=kernel,
    )
    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)

    yesterday = await runtime.chat(
        "alice",
        "Какие документы вчера загружал JBL?",
        actor=_actor(),
    )
    day_before = await runtime.chat(
        "alice",
        "А позавчера?",
        actor=_actor(),
        conversation_id=yesterday["conversation_id"],
    )

    assert len(kernel.calls) == 2
    assert all(call["person"] == "jbl" for call in kernel.calls)
    assert all(call["documents_only"] is True for call in kernel.calls)
    first_since = datetime.fromisoformat(str(kernel.calls[0]["since"]))
    second_since = datetime.fromisoformat(str(kernel.calls[1]["since"]))
    assert first_since - second_since == timedelta(days=1)
    assert yesterday["tools_used"] == ["user_activity"]
    assert day_before["tools_used"] == ["user_activity"]
    assert "границы дня не определены" not in day_before["message"].casefold()


@pytest.mark.asyncio
async def test_self_document_inventory_needs_neither_a_name_day_nor_admin_tool_schema(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="user", display_name="Алиса")
    kernel = _InventoryKernel(available=False, person_name="Алиса")
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_NeverModel(),
        kernel=kernel,
    )
    actor = ActorContext(user_id="alice", preset_key="user", source="test")
    catalog_calls: list[tuple[str, str, int]] = []

    def body_free_catalog(user_id: str, uploaded_by: str, *, limit: int = 5_000):  # noqa: ANN202
        catalog_calls.append((user_id, uploaded_by, limit))
        return [
            {"id": "raw_alpha", "filename": "alpha.pdf", "content_type": "file"},
            {"id": "raw_beta", "filename": "beta.docx", "content_type": "file"},
        ]

    monkeypatch.setattr(storage, "list_owned_file_catalog", body_free_catalog)

    first = await runtime.chat(
        "alice",
        "Какие я тебе документы скидывал?",
        actor=actor,
    )
    repeated = await runtime.chat(
        "alice",
        "И всё?",
        actor=actor,
        conversation_id=first["conversation_id"],
    )
    natural_self = await runtime.chat(
        "alice",
        "какие у меня есть загруженные файлы?",
        actor=actor,
    )
    all_documents = await runtime.chat(
        "alice",
        "выведи все документы которые я загружал",
        actor=actor,
    )
    reversed_range = await runtime.chat(
        "alice",
        "19-12 число",
        actor=actor,
        conversation_id=all_documents["conversation_id"],
    )
    last_week = await runtime.chat(
        "alice",
        "выведи все документы которые я загружал за последнюю неделю",
        actor=actor,
    )

    assert len(kernel.calls) == 2
    assert all(call["person"] == "alice" for call in kernel.calls)
    assert all(call["documents_only"] is True for call in kernel.calls)
    assert all(call["since"] is not None and call["until"] is not None for call in kernel.calls)
    reversed_since = datetime.fromisoformat(str(kernel.calls[0]["since"]))
    reversed_until = datetime.fromisoformat(str(kernel.calls[0]["until"]))
    assert timedelta(days=7) < reversed_until - reversed_since <= timedelta(days=8)
    assert catalog_calls == [("alice", "alice", 5_001)] * 4
    for reply in (first, repeated, natural_self, all_documents):
        assert "участник не определён" not in reply["message"].casefold()
        assert "за всё время" in reply["message"].casefold()
        assert "alpha.pdf" in reply["message"] and "beta.docx" in reply["message"]
        assert reply["tools_used"] == []
    assert "alpha.pdf" in last_week["message"] and "beta.docx" in last_week["message"]
    assert "2 из 2" in last_week["message"]
    assert last_week["tools_used"] == ["user_activity"]
    assert "alpha.pdf" in reversed_range["message"] and "beta.docx" in reversed_range["message"]
    assert reversed_range["tools_used"] == ["user_activity"]
    stored_all = storage.get_message(str(all_documents["message_id"]), "alice")
    assert stored_all is not None
    all_metadata = json.loads(str(stored_all["metadata_json"] or "{}"))
    assert all_metadata["structural"]["person_document_inventory_self"] is True
    assert "Проверила выборку повторно" in repeated["message"]

    july = await runtime.chat("alice", "Какие документы я скидывал в июле?", actor=actor)
    assert len(kernel.calls) == 3
    assert "alpha.pdf" in july["message"] and "2 из 2" in july["message"]
    assert july["tools_used"] == ["user_activity"]

    for temporal_request in (
        "Какие документы я скидывал в 2025?",
        "Какие документы я скидывал в первом квартале?",
    ):
        scoped = await runtime.chat("alice", temporal_request, actor=actor)
        assert len(kernel.calls) == 3, f"{temporal_request!r} unexpectedly called activity"
        assert len(catalog_calls) == 4, f"{temporal_request!r} was silently widened to all time"
        assert "неизвест" in scoped["message"].casefold()


def test_all_time_inventory_requires_explicit_period_keys() -> None:
    assert not _valid_person_document_inventory_data(
        {
            "период": {},
            "документов с подтверждённым автором": 0,
            "документов без отметки автора": 0,
            "документы": [],
            "пагинация": {
                "смещение": 0,
                "показано": 0,
                "из подтверждённых": 0,
                "подтверждённый перечень показан полностью": True,
                "следующее смещение": None,
            },
        },
        expected_since=None,
        expected_until=None,
    )


@pytest.mark.asyncio
@pytest.mark.parametrize("mode", ["unavailable", "malformed", "ambiguous"])
async def test_exact_inventory_fails_closed_before_generic_generation(
    settings,
    storage,
    monkeypatch,
    mode: str,
) -> None:
    storage.ensure_user("alice", preset_key="owner", display_name="Owner")
    if mode == "ambiguous":
        storage.ensure_user("jbl-one", display_name="JBL")
        storage.ensure_user("jbl-two", display_name="JBL")
    else:
        storage.ensure_user("jbl", display_name="JBL")
    kernel = _InventoryKernel(
        available=mode != "unavailable",
        malformed=mode == "malformed",
    )
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_NeverModel(),
        kernel=kernel,
    )
    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)

    reply = await runtime.chat(
        "alice",
        "Какие документы вчера загружал JBL?",
        actor=_actor(),
    )

    folded = reply["message"].casefold()
    assert "неизвест" in folded
    assert "это всё" in folded
    assert "alpha.pdf" not in reply["message"]
    if mode == "malformed":
        assert len(kernel.calls) == 1
    else:
        assert kernel.calls == []


class _PlainAnswerModel:
    enabled = True
    total_budget_sec = 1.0

    def __init__(self, answer: str) -> None:
        self.answer = answer
        self.calls: list[list[dict[str, Any]]] = []

    async def chat(self, messages, tools=None, **kwargs):  # noqa: ANN001, ARG002
        self.calls.append(list(messages))
        return {"content": self.answer, "tool_calls": None, "_queue_wait_sec": 0.0}


@pytest.mark.asyncio
async def test_unrelated_completeness_question_is_not_hijacked_by_inventory(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    conversation = storage.create_conversation("alice", title="ordinary synthetic chat")
    storage.store_message(str(conversation["id"]), "alice", "user", "Назови один цвет")
    storage.store_message(str(conversation["id"]), "alice", "assistant", "Синий")
    kernel = _InventoryKernel()
    model = _PlainAnswerModel("Да, для текущего вопроса это полный ответ.")
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=kernel,
    )
    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)

    reply = await runtime.chat(
        "alice",
        "Это всё?",
        actor=_actor(),
        conversation_id=str(conversation["id"]),
    )

    assert kernel.calls == []
    assert "документ" not in reply["message"].casefold()
    assert "неизвест" not in reply["message"].casefold()


def test_inventory_intent_does_not_hijack_a_document_content_question() -> None:
    assert _PERSON_DOCUMENT_INVENTORY.search("Какие документы сегодня загружал JBL?")
    assert not _PERSON_DOCUMENT_INVENTORY.search("Что в документе, который JBL загрузил сегодня?")
    assert not _PERSON_DOCUMENT_INVENTORY.search("Что написал JBL в документе, который загрузил сегодня?")


def test_today_inventory_window_stops_at_local_now(settings, storage, monkeypatch) -> None:
    runtime = AgentRuntime(settings, storage)
    fixed = datetime(2026, 8, 9, 12, 34, 56)
    monkeypatch.setattr(runtime, "_local_now", lambda: fixed)
    monkeypatch.setattr(runtime, "_local_today", lambda: fixed.date())

    since, until, label, complete = runtime._closed_document_day_window("сегодня")  # noqa: SLF001

    assert since.startswith("2026-08-08T21:00:00")
    assert until.startswith("2026-08-09T09:34:56")
    assert label == "2026-08-09 по состоянию на 12:34"
    assert complete is False


def test_bare_day_of_month_inventory_resolves_the_latest_past_local_day(
    settings,
    storage,
    monkeypatch,
) -> None:
    runtime = AgentRuntime(settings, storage)
    fixed = datetime(2026, 8, 17, 16, 25, 54)
    monkeypatch.setattr(runtime, "_local_now", lambda: fixed)
    monkeypatch.setattr(runtime, "_local_today", lambda: fixed.date())

    since, until, label, complete = runtime._closed_document_day_window(  # noqa: SLF001
        "Какие я документы 13 числа скидывал?"
    )

    assert since.startswith("2026-08-12T21:00:00")
    assert until.startswith("2026-08-13T20:59:59.999999")
    assert label == "2026-08-13"
    assert complete is True


@pytest.mark.asyncio
async def test_live_wording_for_my_files_on_the_thirteenth_never_widens_to_all_time(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner", display_name="Owner")
    kernel = _InventoryKernel(person_name="Owner")
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_NeverModel(),
        kernel=kernel,
    )
    fixed = datetime(2026, 8, 17, 16, 25, 54)
    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)
    monkeypatch.setattr(runtime, "_local_now", lambda: fixed)
    monkeypatch.setattr(runtime, "_local_today", lambda: fixed.date())

    reply = await runtime.chat(
        "alice",
        "Какие я документы 13 числа скидывал?",
        actor=_actor(),
    )

    assert len(kernel.calls) == 1
    assert kernel.calls[0]["since"].startswith("2026-08-12T21:00:00")
    assert kernel.calls[0]["until"].startswith("2026-08-13T20:59:59.999999")
    assert "за всё время" not in reply["message"].casefold()
    assert "2026-08-13" in reply["message"]


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "message",
    [
        "а 14 числа что я загружал?",
        "14 числа что я загружал?",
    ],
)
async def test_live_compact_self_upload_inventory_uses_exact_moscow_day_without_history_or_model(
    settings,
    storage,
    monkeypatch,
    message: str,
) -> None:
    storage.ensure_user("alice", preset_key="owner", display_name="Owner")
    kernel = _InventoryWithMessageSearchKernel(person_name="Owner")
    runtime = AgentRuntime(
        replace(settings, verify_answers=False, local_timezone="Europe/Moscow"),
        storage,
        llm=_NeverModel(),
        kernel=kernel,
    )
    fixed = datetime(2026, 8, 17, 16, 25, 54)
    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)
    monkeypatch.setattr(runtime, "_local_now", lambda: fixed)
    monkeypatch.setattr(runtime, "_local_today", lambda: fixed.date())

    def forbidden_message_search(*args: Any, **kwargs: Any) -> Any:
        del args, kwargs
        raise AssertionError("the compact upload inventory searched stored messages")

    monkeypatch.setattr(storage, "search_messages", forbidden_message_search)

    reply = await runtime.chat("alice", message, actor=_actor())

    assert len(kernel.calls) == 1
    call = kernel.calls[0]
    assert call["person"] == "alice"
    assert call["documents_only"] is True
    assert call["since"].startswith("2026-08-13T21:00:00")
    assert call["until"].startswith("2026-08-14T20:59:59.999999")
    assert reply["tools_used"] == ["user_activity"]
    assert "2026-08-14" in reply["message"]
    assert "alpha.pdf" in reply["message"] and "beta.docx" in reply["message"]
    assert "за всё время" not in reply["message"].casefold()


@pytest.mark.parametrize(
    "message",
    [
        "а 14 числа что я писал?",
        "а 14 числа что я скачивал?",
        "а 14 числа кто загружал документы?",
        "а 14 числа что было загружено?",
    ],
)
def test_compact_self_upload_inventory_does_not_claim_adjacent_questions(message: str) -> None:
    assert _PERSON_DOCUMENT_INVENTORY.search(message) is None


@pytest.mark.asyncio
async def test_latest_named_uploader_files_are_newest_first_without_a_day_parser(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner", display_name="Owner")
    storage.ensure_user("jbl", preset_key="user", display_name="JBL")
    kernel = _InventoryKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_NeverModel(),
        kernel=kernel,
    )
    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)
    catalog_calls: list[tuple[str, str, int]] = []

    def body_free_catalog(user_id: str, uploaded_by: str, *, limit: int = 5_000):  # noqa: ANN202
        catalog_calls.append((user_id, uploaded_by, limit))
        return [
            {"id": "raw_old", "filename": "old.pdf", "content_type": "file"},
            {"id": "raw_middle", "filename": "middle.pdf", "content_type": "file"},
            {"id": "raw_new", "filename": "new.pdf", "content_type": "file"},
        ]

    monkeypatch.setattr(storage, "list_owned_file_catalog", body_free_catalog)

    reply = await runtime.chat(
        "alice",
        "Какие последние файлы присылал JBL?",
        actor=_actor(),
    )

    assert kernel.calls == []
    assert catalog_calls == [("jbl", "jbl", 5_001)]
    assert "границы дня" not in reply["message"].casefold()
    assert "новые первыми" in reply["message"].casefold()
    assert reply["message"].index("new.pdf") < reply["message"].index("middle.pdf")
    assert reply["message"].index("middle.pdf") < reply["message"].index("old.pdf")


class _NoToolsKernel:
    authorization = _AllowAll()

    def get_tool_definitions(self, actor, topic=""):  # noqa: ANN001, ARG002
        return []

    async def execute(self, tool, params, actor=None):  # noqa: ANN001, ARG002
        raise AssertionError(f"unexpected tool {tool}: {params}")


@pytest.mark.asyncio
async def test_successful_generation_cannot_replay_a_false_model_outage(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    stale = "К сожалению, не могу связаться с моделью — она не отвечает."
    model = _PlainAnswerModel(stale)
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=_NoToolsKernel(),
    )
    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)

    reply = await runtime.chat("alice", "Повтори ответ по существу", actor=_actor())

    assert "модель" not in reply["message"].casefold()
    assert "повторите запрос" in reply["message"].casefold()
    stored = storage.get_message(str(reply["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    assert metadata["structural"]["output_guards"]["false_model_outage_replaced"] is True


class _FailingModel:
    enabled = True
    total_budget_sec = 1.0

    async def chat(self, messages, tools=None, **kwargs):  # noqa: ANN001, ARG002
        raise RuntimeError("synthetic transport failure")


class _DisabledModel:
    enabled = False
    total_budget_sec = 1.0

    async def chat(self, messages, tools=None, **kwargs):  # noqa: ANN001, ARG002
        raise AssertionError("a disabled model was called")


@pytest.mark.asyncio
async def test_a_real_model_failure_keeps_the_truthful_outage_diagnosis(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_FailingModel(),
        kernel=_NoToolsKernel(),
    )
    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)

    reply = await runtime.chat("alice", "Ответь по существу", actor=_actor())

    assert reply["context"]["llm_failed"] is True
    assert "не могу связаться с моделью" in reply["message"].casefold()


@pytest.mark.asyncio
async def test_an_intentionally_disabled_model_is_not_rewritten_as_available(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_DisabledModel(),
        kernel=_NoToolsKernel(),
    )
    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)

    reply = await runtime.chat("alice", "Ответь по существу", actor=_actor())

    assert reply["context"]["llm_failed"] is False
    assert "модель недоступна" in reply["message"].casefold()
    assert "повторите запрос ещё раз" not in reply["message"].casefold()


@pytest.mark.asyncio
async def test_a_repair_cannot_reintroduce_a_false_model_outage_at_the_final_boundary(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    runtime = AgentRuntime(
        replace(settings, verify_answers=True, verify_min_answer_chars=1),
        storage,
        llm=_NeverModel(),
    )
    verification_calls = 0

    async def generate(context, message, attachments):  # noqa: ANN001, ARG001
        return {
            "content": "Первичный содержательный ответ по синтетическому документу.",
            "tools_used": [],
            "_model_generated": True,
        }

    async def verify(query, response, context, *, tool_evidence=None):  # noqa: ANN001, ARG001
        nonlocal verification_calls
        verification_calls += 1
        if verification_calls == 1:
            return {"status": "failed", "ok": False, "score": 0.0, "issues": ["synthetic"]}
        return {"status": "passed", "ok": True, "score": 1.0, "issues": []}

    async def repair(*args, **kwargs):  # noqa: ANN002, ANN003
        return "К сожалению, не могу связаться с моделью — она не отвечает."

    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)
    monkeypatch.setattr(runtime, "_generate_response", generate)
    monkeypatch.setattr(runtime, "_verify_response", verify)
    monkeypatch.setattr(runtime, "_repair_once", repair)

    reply = await runtime.chat(
        "alice",
        "Что указано в приложенном документе?",
        actor=_actor(),
        attachments=[
            _transient_attachment(
                filename="synthetic.txt",
                text="Проверяемый синтетический факт.",
            )
        ],
        enable_tools=False,
    )

    assert verification_calls == 2
    assert "модель" not in reply["message"].casefold()
    assert "повторите запрос" in reply["message"].casefold()
    assert reply["verification_status"] == "unknown"


def test_false_outage_matcher_ignores_conditional_and_historical_discussion() -> None:
    assert _FALSE_CURRENT_MODEL_OUTAGE.search("К сожалению, не могу связаться с моделью — она не отвечает.")
    assert not _FALSE_CURRENT_MODEL_OUTAGE.search("Если модель недоступна, сообщи оператору.")
    assert not _FALSE_CURRENT_MODEL_OUTAGE.search("Вчера модель была недоступна десять минут.")
    assert not _FALSE_CURRENT_MODEL_OUTAGE.search("Он написал: «Сейчас модель недоступна».")


def _trusted_synthetic_docx(*, incomplete: bool) -> dict[str, Any]:
    document = Document()
    document.add_heading("Synthetic status", level=1)
    table = document.add_table(rows=3, cols=2)
    for row, values in zip(
        table.rows,
        (("Item", "State"), ("Alpha", "Ready"), ("Beta", "Review")),
        strict=True,
    ):
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = value
    payload = io.BytesIO()
    document.save(payload)
    result = DocumentExtractor(secret_values=()).extract(
        payload.getvalue(),
        "synthetic-summary.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert result.success is True and isinstance(result.office_structure_index, dict)
    index = copy.deepcopy(result.office_structure_index)
    if incomplete:
        index["complete"] = False
        index["coverage"]["reasons"] = ["text_budget"]
    assert validate_runtime_office_index(index, result.text) == index
    return trusted_office_attachment(
        {
            "filename": "synthetic-summary.docx",
            "transient_text": result.text,
            "extraction_success": True,
            "verification_eligible": True,
            OFFICE_STRUCTURE_KEY: index,
        }
    )


@pytest.mark.asyncio
@pytest.mark.parametrize("incomplete", [False, True])
async def test_bare_docx_summary_is_not_misclassified_as_an_exact_inventory(
    settings,
    storage,
    monkeypatch,
    incomplete: bool,
) -> None:
    """A bare readable upload gets the ordinary attachment review, not a receipt."""

    storage.ensure_user("alice", preset_key="owner")
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_NeverModel(),
    )
    attachment = _trusted_synthetic_docx(incomplete=incomplete)
    draft = (
        "Сводка документа: указаны две позиции — Alpha и Beta. "
        "Это полный обзор всех двух позиций в прочитанной структуре."
    )

    generate_calls = {"n": 0}

    async def generate(context, message, attachments):  # noqa: ANN001, ARG001
        del context, message
        generate_calls["n"] += 1
        assert attachments
        projected = attachments[0]
        if incomplete:
            assert projected.get("_office_structured") is not True
            assert projected.get("_office_full_text_fit") is True
            assert projected.get("_source_text_complete") is True
            assert projected.get("_prompt_projection_complete") is True
            assert "Synthetic status" in str(projected.get("transient_text") or "")
        else:
            assert projected.get("_office_structured") is True
        return {"content": draft, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Загружен документ: synthetic-summary.docx",
        actor=_actor(),
        attachments=[attachment],
        enable_tools=False,
        synthetic_document_notice=True,
    )

    assert generate_calls["n"] == 1
    assert reply["message"] == draft
    assert reply["message"] != OFFICE_EXACT_UNAVAILABLE_MESSAGE
    assert "пришлите" not in reply["message"].casefold()
    assert reply["message_format"] == "markdown"
    assert reply["tools_used"] == []
    assert reply["attachment_context_readable_count"] == 1
    # The bounded Office index may be incomplete while the fully extracted
    # source text still fits the ordinary review prompt. Coverage here is the
    # authenticated source body, not exact-Office inventory completeness.
    assert reply["attachment_coverage_complete"] is True
    assert reply["verification_status"] != "unknown"


def test_stale_web_isolation_rejects_reference_only_requests() -> None:
    requests = (
        "Найди в интернете то же самое",
        "Найди в интернете по тому вопросу",
        "Найди в интернете сведения об этом",
        "Найди в интернете оттуда",
        "Найди в интернете дополнительную информацию о нём",
        "Найди в интернете по ранее присланным данным",
    )
    assert all(_WEB_ISOLATION_DEICTIC.search(request) for request in requests)


def test_natural_multi_document_count_is_bounded_and_negation_safe() -> None:
    assert _multi_attachment_summary_count("Обобщи последние три загруженных документа") == 3
    assert _multi_attachment_summary_count("Сделай общую сводку по трём последним документам") == 3
    assert _multi_attachment_summary_count("Подготовь сводку по 4 последним файлам") == 4
    assert _multi_attachment_summary_count("Не составляй сводку трёх документов") is None
    assert _multi_attachment_summary_count("Повтори фразу «обобщи три документа»") is None


@pytest.mark.parametrize(
    ("question", "predicate"),
    [
        ("Какую должность занимает иванов в документе?", "занимает"),
        ("Кем работает иванов в документе?", "работает"),
    ],
)
def test_a_rare_lowercase_surname_outranks_a_repeated_role_predicate(
    settings,
    storage,
    question: str,
    predicate: str,
) -> None:
    source = (
        (f"{predicate} общую позицию " + "A" * 40 + "\n") * 900
        + "X" * 30_000
        + "\nИванов\nДолжность: главный инженер\n"
    )
    attachment = _OwnedAttachment(
        {
            "filename": "lowercase-surname.txt",
            "transient_text": source,
            "extraction_success": True,
            "verification_eligible": True,
        }
    )

    projected, state = _project_attachments_for_request(question, [attachment])
    body = str(projected[0].get("transient_text") or "")
    evidence = "".join(item["output"].split("\n", 1)[1] for item in _attachment_evidence_chunks(projected))
    runtime = AgentRuntime(settings, storage)
    synthesis = "\n".join(
        str(item.get("content") or "")
        for item in runtime._build_initial_messages(  # noqa: SLF001
            AgentContext(conversation_id="conv", user_id="alice"),
            question,
            projected,
            tool_enabled=False,
        )
    )

    assert state.status == "matched" and state.scan_complete is True
    assert "Иванов\nДолжность: главный инженер" in body
    assert body in evidence
    assert body in synthesis


@pytest.mark.parametrize("truncated", [False, True])
@pytest.mark.asyncio
async def test_missing_required_surname_is_closed_not_predicate_matched(
    settings,
    storage,
    truncated: bool,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    source = ("занимает должность Петров " + "A" * 40 + "\n") * 1200
    attachment = _OwnedAttachment(
        {
            "filename": "surname-absent.txt",
            "transient_text": source,
            "extraction_success": True,
            "verification_eligible": True,
            "text_truncated": truncated,
        }
    )
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_NeverModel(),
    )

    reply = await runtime.chat(
        "alice",
        "Какую должность занимает иванов в документе?",
        actor=_actor(),
        attachments=[attachment],
        enable_tools=False,
    )

    assert reply["attachment_query_status"] == ("unknown" if truncated else "not_found")
    assert reply["attachment_query_files_matched"] == 0
    if truncated:
        assert "доказательно проверить" in reply["message"].casefold()
    else:
        assert "не найден" in reply["message"].casefold()


def test_repeated_surname_retains_the_tail_occurrence_beside_its_position(
    settings,
    storage,
) -> None:
    prefix = "Иванов упомянут в списке.\n" * 1200
    source = prefix + "X" * max(0, 72_000 - len(prefix)) + "\nИванов\nДолжность: главный инженер\n"
    question = "Какая должность у иванова в документе?"
    attachment = _OwnedAttachment(
        {
            "filename": "repeated-surname.txt",
            "transient_text": source,
            "extraction_success": True,
            "verification_eligible": True,
        }
    )

    projected, state = _project_attachments_for_request(question, [attachment])
    body = str(projected[0].get("transient_text") or "")
    evidence = "".join(item["output"].split("\n", 1)[1] for item in _attachment_evidence_chunks(projected))
    runtime = AgentRuntime(settings, storage)
    synthesis = "\n".join(
        str(item.get("content") or "")
        for item in runtime._build_initial_messages(  # noqa: SLF001
            AgentContext(conversation_id="conv", user_id="alice"),
            question,
            projected,
            tool_enabled=False,
        )
    )

    assert state.status == "matched" and state.scan_complete is True
    assert "Иванов\nДолжность: главный инженер" in body
    assert body in evidence
    assert body in synthesis


@pytest.mark.parametrize(
    "question",
    [
        "Подскажи, какая должность у иванова в документе?",
        "Кем в документе работает иванов?",
        "Что за должность у иванова в документе?",
        "Укажи должность иванова в документе",
        "Что указано про иванова в документе?",
        "Можешь сказать, какая должность у иванова в документе?",
        "Какова должность иванова в документе?",
        "Должность иванова в документе?",
        "Определи должность иванова сейчас по данным документа",
        "Какая должность у Иванова, по-твоему, в документе?",
        "Какую роль занимает бизнес-аналитик Иванов в документе?",
    ],
)
def test_natural_surname_lookup_order_and_politeness_keep_the_strong_anchor(
    question: str,
) -> None:
    source = "X" * 72_000 + "\nИванов\nДолжность: главный инженер\n"
    attachment = _OwnedAttachment(
        {
            "filename": "natural-surname.txt",
            "transient_text": source,
            "extraction_success": True,
            "verification_eligible": True,
        }
    )

    projected, state = _project_attachments_for_request(question, [attachment])

    assert state.status == "matched" and state.scan_complete is True
    assert "Иванов\nДолжность: главный инженер" in str(projected[0]["transient_text"])


@pytest.mark.parametrize(
    "question",
    [
        "Что думаешь об этом документе?",
        "Скажи, что думаешь об этом документе?",
        "Подскажи, что ты думаешь об этом документе?",
        "Скажи кратко, о чём документ.",
        "Какая основная мысль документа?",
        "Какой главный вывод документа?",
        "Где здесь слабые места документа?",
        "Каково твоё мнение о документе?",
        "Кто, по-твоему, автор этого документа?",
    ],
)
def test_open_document_synthesis_cannot_be_hijacked_by_lookup_words(question: str) -> None:
    attachment = _OwnedAttachment(
        {
            "filename": "context-only.txt",
            "transient_text": "Полностью прочитанный синтетический текст.",
            "extraction_success": True,
            "verification_eligible": True,
        }
    )

    _projected, projection = _project_attachments_for_request(question, [attachment])

    assert projection.applied is False


def test_a_weak_factual_lookup_cannot_claim_complete_absence() -> None:
    attachment = _OwnedAttachment(
        {
            "filename": "weak-lookup.txt",
            "transient_text": "Полностью прочитанный синтетический текст.",
            "extraction_success": True,
            "verification_eligible": True,
        }
    )

    _projected, weak_lookup = _project_attachments_for_request(
        "Какая должность у инженера в документе?", [attachment]
    )

    assert weak_lookup.status == "unknown"


@pytest.mark.parametrize(
    "literal",
    [
        "SYNTHETIC-NODE-42",
        "owner42@example.invalid",
        "https://document.invalid/Case-42",
    ],
)
def test_unquoted_machine_literal_is_a_strong_document_anchor(literal: str) -> None:
    attachment = _OwnedAttachment(
        {
            "filename": "machine-literal.txt",
            "transient_text": "X" * 72_000 + f"\nExact literal: {literal}\n",
            "extraction_success": True,
            "verification_eligible": True,
        }
    )

    projected, state = _project_attachments_for_request(f"Найди {literal} в документе.", [attachment])

    assert state.status == "matched" and state.scan_complete is True
    assert literal in str(projected[0]["transient_text"])


def test_document_urls_are_exact_inert_evidence_not_web_provenance() -> None:
    exact = "https://Document.Invalid/CasePath?Q=AbC"
    allowed = _attachment_web_fact_targets([{"tool": "attachment", "output": f"Endpoint literal: {exact}"}])

    domain_answer, domain_changed = _reconcile_attachment_web_literals(
        "В документе указан домен document.invalid.",
        allowed=allowed,
    )
    assert domain_changed is True
    assert "document.invalid" in domain_answer
    assert _attachment_web_literals_are_grounded(domain_answer, allowed)

    mutated, _ = _reconcile_attachment_web_literals(
        "В документе указан https://Document.Invalid/casepath?Q=AbC и https://invented.invalid/x.",
        allowed=allowed,
    )
    assert "casepath" not in mutated
    assert "invented.invalid" not in mutated

    provenance, _ = _reconcile_attachment_web_literals(
        f"По данным интернета: {exact}",
        allowed=allowed,
    )
    assert "По данным интернета" not in provenance
    assert "В документе" in provenance
    assert _attachment_web_literals_are_grounded(provenance, allowed)


def test_document_source_word_without_a_web_literal_is_byte_preserved() -> None:
    answer = "  Контрольный код из источника: LINEAGE-TARGET-1.  "

    assert _reconcile_attachment_web_literals(answer, allowed=frozenset()) == (answer, False)


def test_explicit_web_provenance_without_a_literal_remains_guarded() -> None:
    answer = "По данным интернета, контрольный код LINEAGE-TARGET-1."
    reconciled, changed = _reconcile_attachment_web_literals(
        answer,
        allowed=frozenset(),
    )

    assert changed is False
    assert reconciled == answer


@pytest.mark.parametrize("separator", [", ", "! ", "\n"])
def test_attachment_web_reconciliation_does_not_drop_an_unrelated_source_clause(
    separator: str,
) -> None:
    answer = (
        f"Контрольный код из источника: LINEAGE-TARGET-1{separator}"
        "Неподтверждённая ссылка: https://invented.invalid/path."
    )

    reconciled, changed = _reconcile_attachment_web_literals(answer, allowed=frozenset())

    assert changed is True
    assert "LINEAGE-TARGET-1" in reconciled
    assert "invented.invalid" not in reconciled


@pytest.mark.parametrize("separator", [", ", "; ", "! ", "\n"])
def test_benign_document_source_fragment_does_not_poison_an_exact_url(
    separator: str,
) -> None:
    exact = "https://Document.Invalid/CasePath?Q=AbC"
    answer = f"Источник документа: D02{separator}endpoint: {exact}"
    allowed = _attachment_web_fact_targets([{"tool": "attachment", "output": answer}])

    reconciled, changed = _reconcile_attachment_web_literals(answer, allowed=allowed)

    assert changed is True
    assert "Источник документа: D02" in reconciled
    assert exact in reconciled
    assert _attachment_web_literals_are_grounded(reconciled, allowed) is True


@pytest.mark.parametrize(
    "answer_template",
    [
        "Источник документа {url} содержит код LINEAGE-TARGET-1.",
        "Источник документа, {url} содержит код LINEAGE-TARGET-1.",
        "Источник документа: {url}; код LINEAGE-TARGET-1.",
        "Источник документа!\n{url} содержит код LINEAGE-TARGET-1.",
    ],
)
def test_exact_url_inerting_never_erases_a_same_clause_attachment_fact(
    answer_template: str,
) -> None:
    exact = "https://Document.Invalid/CasePath?Q=AbC"
    answer = answer_template.format(url=exact)
    allowed = _attachment_web_fact_targets([{"tool": "attachment", "output": answer}])

    reconciled, changed = _reconcile_attachment_web_literals(answer, allowed=allowed)

    assert changed is True
    assert exact in reconciled
    assert "LINEAGE-TARGET-1" in reconciled
    assert _attachment_web_literals_are_grounded(reconciled, allowed) is True


@pytest.mark.parametrize(
    "answer_template",
    [
        "По данным интернета: {url} содержит код LINEAGE-TARGET-1.",
        "Информация LINEAGE-TARGET-1 была найдена в интернете: {url}.",
    ],
)
def test_fabricated_web_label_is_removed_without_erasing_the_grounded_fact(
    answer_template: str,
) -> None:
    exact = "https://Document.Invalid/CasePath?Q=AbC"
    answer = answer_template.format(url=exact)
    allowed = _attachment_web_fact_targets([{"tool": "attachment", "output": answer}])

    reconciled, changed = _reconcile_attachment_web_literals(answer, allowed=allowed)

    assert changed is True
    assert "данным интернета" not in reconciled.casefold()
    assert exact in reconciled
    assert "LINEAGE-TARGET-1" in reconciled


def test_unsupported_web_clause_cannot_launder_residual_as_a_document_fact() -> None:
    answer = "По данным интернета: https://invented.invalid/not-in-document код FABRICATED-42."

    reconciled, changed = _reconcile_attachment_web_literals(answer, allowed=frozenset())

    assert changed is True
    assert "invented.invalid" not in reconciled
    assert "по данным интернета" in reconciled.casefold()
    assert "в документе указано" not in reconciled.casefold()


@pytest.mark.asyncio
async def test_attachment_source_word_cannot_erase_the_model_answer(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_NeverModel(),
        kernel=_NoToolsKernel(),
    )
    answer = "Контрольный код из источника процитированного ответа: LINEAGE-TARGET-1."

    async def generate(context, message, attachments):  # noqa: ANN001, ARG001
        return {
            "content": answer,
            "tools_used": [],
            "_model_generated": True,
        }

    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)
    monkeypatch.setattr(runtime, "_generate_response", generate)

    reply = await runtime.chat(
        "alice",
        "Повтори контрольный код из источника процитированного ответа.",
        actor=_actor(),
        attachments=[
            _transient_attachment(
                filename="older-source.odt",
                text="Контрольный код: LINEAGE-TARGET-1.",
            )
        ],
        enable_tools=False,
        quoted_attachment_reference=True,
    )

    assert reply["message"] == answer
    assert reply["verification_status"] != "unknown"
    stored = storage.get_message(str(reply["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    assert metadata["attachment_context_used"] is True
    assert metadata["attachment_context_expected_count"] == 1
    assert metadata["attachment_context_readable_count"] == 1
    assert metadata["attachment_coverage_complete"] is True
    assert metadata["verification"]["issues"] != ["unsupported_attachment_web_literal_removed"]


@pytest.mark.asyncio
async def test_exact_attachment_url_survives_the_final_guard_beside_a_document_source_label(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_NeverModel(),
        kernel=_NoToolsKernel(),
    )
    exact = "https://Document.Invalid/CasePath?Q=AbC"
    invented = "https://invented.invalid/not-in-document"
    answer = f"Источник документа: D02; endpoint: {exact}; лишняя ссылка: {invented}."

    async def generate(context, message, attachments):  # noqa: ANN001, ARG001
        return {"content": answer, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)
    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Покажи источник документа и точный endpoint.",
        actor=_actor(),
        attachments=[
            _transient_attachment(
                filename="source.odt",
                text=f"Источник документа: D02; endpoint: {exact}",
            )
        ],
        enable_tools=False,
    )

    assert reply["message"] != _WEB_EVIDENCE_MISSING
    assert "Источник документа: D02" in reply["message"]
    assert exact in reply["message"]
    assert invented not in reply["message"]


@pytest.mark.asyncio
async def test_same_clause_attachment_fact_survives_url_reconciliation_and_final_guard(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_NeverModel(),
        kernel=_NoToolsKernel(),
    )
    exact = "https://Document.Invalid/CasePath?Q=AbC"
    invented = "https://invented.invalid/not-in-document"
    grounded = f"Источник документа {exact} содержит код LINEAGE-TARGET-1."
    answer = f"{grounded} Неподтверждённая ссылка: {invented}."

    async def generate(context, message, attachments):  # noqa: ANN001, ARG001
        return {"content": answer, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)
    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Покажи источник документа, точный endpoint и контрольный код.",
        actor=_actor(),
        attachments=[_transient_attachment(filename="source.odt", text=grounded)],
        enable_tools=False,
    )

    assert reply["message"] != _WEB_EVIDENCE_MISSING
    assert exact in reply["message"]
    assert "LINEAGE-TARGET-1" in reply["message"]
    assert invented not in reply["message"]


@pytest.mark.parametrize(
    "model_answer",
    [
        "Проверила онлайн: контрольный код LINEAGE-TARGET-1.",
        "Интернет-источник сообщает: контрольный код LINEAGE-TARGET-1.",
    ],
)
@pytest.mark.asyncio
async def test_explicit_online_claim_without_a_url_reaches_the_web_hard_guard(
    settings,
    storage,
    monkeypatch,
    model_answer: str,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_NeverModel(),
        kernel=_NoToolsKernel(),
    )

    async def generate(context, message, attachments):  # noqa: ANN001, ARG001
        return {
            "content": model_answer,
            "tools_used": [],
            "_model_generated": True,
        }

    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)
    monkeypatch.setattr(runtime, "_generate_response", generate)

    reply = await runtime.chat(
        "alice",
        "Повтори контрольный код из приложенного документа.",
        actor=_actor(),
        attachments=[
            _transient_attachment(
                filename="source.odt",
                text="Контрольный код: LINEAGE-TARGET-1.",
            )
        ],
        enable_tools=False,
    )

    assert reply["message"] == _WEB_EVIDENCE_MISSING
    assert "LINEAGE-TARGET-1" not in reply["message"]
    stored = storage.get_message(str(reply["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    assert metadata["structural"]["output_guards"]["web_evidence_replaced"] is True


@pytest.mark.asyncio
async def test_unsupported_web_clause_cannot_reach_the_final_attachment_answer(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_NeverModel(),
        kernel=_NoToolsKernel(),
    )
    model_answer = "По данным интернета: https://invented.invalid/not-in-document код FABRICATED-42."

    async def generate(context, message, attachments):  # noqa: ANN001, ARG001
        return {
            "content": model_answer,
            "tools_used": [],
            "_model_generated": True,
        }

    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)
    monkeypatch.setattr(runtime, "_generate_response", generate)

    reply = await runtime.chat(
        "alice",
        "Повтори контрольный код из приложенного документа.",
        actor=_actor(),
        attachments=[
            _transient_attachment(
                filename="source.odt",
                text="Контрольный код: LEGITIMATE-42.",
            )
        ],
        enable_tools=False,
    )

    assert reply["message"] == _WEB_EVIDENCE_MISSING
    assert "FABRICATED-42" not in reply["message"]


def test_hostile_document_url_syntax_stays_visible_without_a_telegram_link() -> None:
    private_url = "http://127.0.0.1/private"
    allowed = _attachment_web_fact_targets([{"tool": "attachment", "output": f"literal {private_url}"}])
    reconciled, _ = _reconcile_attachment_web_literals(
        f"В документе указано `[x]({private_url})`.",
        allowed=allowed,
    )
    rendered = to_telegram_html(reconciled)

    assert "127.0.0.1/private" in rendered
    assert "href=" not in rendered.casefold()


class _WebKernel:
    authorization = _AllowAll()

    def __init__(self) -> None:
        self.calls: list[tuple[str, dict[str, Any]]] = []

    def get_tool_definitions(self, actor, topic=""):  # noqa: ANN001, ARG002
        return [_tool("web_research")]

    async def execute(self, tool, params, actor=None):  # noqa: ANN001, ARG002
        assert tool == "web_research"
        self.calls.append((str(tool), dict(params)))
        public_text = "At normal pressure the synthetic boiling point is 100 C."
        return ToolResult(
            tool,
            True,
            {
                "outbound_attempted": True,
                "sources": [
                    {
                        "url": "https://public.synthetic.example.com/fact",
                        "title": "Synthetic public source",
                        "text": public_text,
                        "text_length": len(public_text),
                        "status_code": 200,
                        "error": "",
                        "truncated": False,
                    }
                ],
                "requested_sources": 1,
                "completed_sources": 1,
                "failed_sources": 0,
                "timed_out_sources": 0,
                "search_timed_out": False,
            },
        )


class _WebAnswerModel:
    enabled = True
    total_budget_sec = 1.0

    def __init__(self) -> None:
        self.calls: list[list[dict[str, Any]]] = []

    async def chat(self, messages, tools=None, **kwargs):  # noqa: ANN001, ARG002
        self.calls.append(list(messages))
        return {
            "content": (
                "Синтетический публичный факт подтверждён: https://public.synthetic.example.com/fact"
            ),
            "tool_calls": None,
            "_queue_wait_sec": 0.0,
        }


def _private_attachment_history(storage, *, old: bool) -> str:
    conversation = storage.create_conversation("alice", title="private synthetic lineage")
    conversation_id = str(conversation["id"])
    storage.store_message(
        conversation_id,
        "alice",
        "user",
        "PRIVATE-HISTORY-CANARY-DO-NOT-SEND",
        metadata={
            "had_attachments": True,
            "attachment_count": 1,
            "private_context_lineage": True,
        },
    )
    storage.store_message(
        conversation_id,
        "alice",
        "assistant",
        "PRIVATE-ANSWER-CANARY-DO-NOT-SEND",
        metadata={"attachment_context_used": True, "private_context_lineage": True},
    )
    if old:
        stale = (datetime.now(UTC) - timedelta(hours=7)).isoformat()
        with storage.transaction() as conn:
            conn.execute(
                "UPDATE messages SET created_at=? WHERE conversation_id=? AND user_id=?",
                (stale, conversation_id, "alice"),
            )
    return conversation_id


def test_fresh_public_news_is_an_explicit_web_request_but_local_news_is_not() -> None:
    assert asks_for_the_web("Покажешь свежие новости за прошедшие сутки?")
    assert asks_for_the_web("Свежие новости за прошедшие сутки покажешь?")
    assert asks_for_the_web("Расскажи последние новости за вчера")
    assert not asks_for_the_web("Покажи новости в документе за вчера")
    assert not asks_for_the_web("В документе сохранены вчерашние новости")


@pytest.mark.asyncio
async def test_recent_private_file_then_fresh_news_denies_web_and_history_use(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    conversation_id = _private_attachment_history(storage, old=False)
    kernel = _WebKernel()
    model = _WebAnswerModel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=kernel,
    )

    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)
    request = "Покажешь свежие новости за прошедшие сутки?"
    reply = await runtime.chat(
        "alice",
        request,
        actor=_actor(),
        conversation_id=conversation_id,
    )

    assert kernel.calls == []
    assert reply["tools_used"] == []
    assert reply["web_evidence_status"] == "none"
    exposed = json.dumps(model.calls, ensure_ascii=False)
    assert "PRIVATE-HISTORY-CANARY" not in exposed
    assert "PRIVATE-ANSWER-CANARY" not in exposed
    assert "приватные вложения" in reply["message"].casefold()


@pytest.mark.asyncio
async def test_old_private_lineage_web_turn_is_denied_before_provider(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    conversation_id = _private_attachment_history(storage, old=True)
    kernel = _WebKernel()
    model = _WebAnswerModel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=kernel,
    )

    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)
    request = "Найди в интернете температуру кипения воды при нормальном давлении"
    reply = await runtime.chat(
        "alice",
        request,
        actor=_actor(),
        conversation_id=conversation_id,
    )

    assert kernel.calls == []
    assert reply["tools_used"] == []
    assert reply["web_evidence_status"] == "none"
    exposed = json.dumps(model.calls, ensure_ascii=False)
    assert "PRIVATE-HISTORY-CANARY" not in exposed
    assert "PRIVATE-ANSWER-CANARY" not in exposed
    assert "приватные вложения" in reply["message"].casefold()
    stored = storage.get_message(str(reply["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    assert metadata["private_context_lineage"] is True


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "query_text",
    [
        "Найди в интернете то же самое",
        "Найди в интернете по тому вопросу",
        "Найди в интернете сведения об этом",
        "Найди в интернете оттуда",
        "Найди в интернете дополнительную информацию о нём",
        "Найди в интернете по ранее присланным данным",
    ],
)
async def test_old_private_lineage_denies_explicit_reference_only_web_requests(
    settings,
    storage,
    monkeypatch,
    query_text: str,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    conversation_id = _private_attachment_history(storage, old=True)
    kernel = _WebKernel()
    model = _WebAnswerModel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=kernel,
    )

    async def prepared(user_id, message, current_conversation_id, **kwargs):  # noqa: ANN001
        del kwargs
        return AgentContext(
            conversation_id=current_conversation_id,
            user_id=user_id,
            person_id=user_id,
            conversation_history=[],
            outward_verdict=("интернет", runtime.web_query_from(message)),
        )

    monkeypatch.setattr(runtime, "_prepare_context", prepared)

    reply = await runtime.chat(
        "alice",
        query_text,
        actor=_actor(),
        conversation_id=conversation_id,
    )

    assert kernel.calls == []
    assert model.calls == []
    assert reply["tools_used"] == []
    assert reply["web_evidence_status"] == "none"
    assert "приватные вложения" in reply["message"].casefold()
    stored = storage.get_message(str(reply["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    assert metadata["structural"]["private_web_search_blocked"] is True


@pytest.mark.asyncio
async def test_recent_attachment_web_turn_denies_even_when_old_literals_leave_prompt_tail(
    settings,
    storage,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    conversation_id = _private_attachment_history(storage, old=False)
    for index in range(21):
        storage.store_message(
            conversation_id,
            "alice",
            "user" if index % 2 == 0 else "assistant",
            f"neutral-{index}",
            metadata={"private_context_lineage": True},
        )
    kernel = _WebKernel()
    model = _WebAnswerModel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=kernel,
    )

    request = "Найди в интернете температуру кипения воды"
    reply = await runtime.chat(
        "alice",
        request,
        actor=_actor(),
        conversation_id=conversation_id,
    )

    assert kernel.calls == []
    assert reply["tools_used"] == []
    assert reply["web_evidence_status"] == "none"
    exposed = json.dumps(model.calls, ensure_ascii=False)
    assert "neutral-20" not in exposed
    assert "PRIVATE-HISTORY-CANARY" not in exposed
    assert "PRIVATE-ANSWER-CANARY" not in exposed
    assert "приватные вложения" in reply["message"].casefold()


def _stored_file(
    storage,
    *,
    filename: str,
    text: str,
    tenant: str = "alice",
    uploader: str = "alice",
) -> RawObject:
    storage.ensure_user(tenant)
    if uploader != tenant:
        storage.ensure_user(uploader)
    raw_id = new_id("raw")
    body = text.encode("utf-8")
    digest = hashlib.sha256(body).hexdigest()
    relative = f"{tenant}/{digest[:2]}/{raw_id}.txt"
    stored = storage.settings.files_dir / relative
    stored.parent.mkdir(parents=True, exist_ok=True)
    stored.write_bytes(body)
    raw = RawObject(
        id=raw_id,
        user_id=tenant,
        source="upload",
        source_ref=new_id("source"),
        raw_content=text,
        content_type="file",
        content_hash=digest,
        metadata_json={
            "filename": filename,
            "uploaded_by": uploader,
            "extraction_success": True,
            "text_extraction_success": True,
            "stored_path": relative,
            "sha256": digest,
            "size_bytes": len(body),
        },
    )
    storage.store_raw_object(raw)
    return raw


def _current_attachment(storage, raw: RawObject) -> dict[str, Any]:  # noqa: ANN001
    metadata = raw.metadata_json if isinstance(raw.metadata_json, dict) else {}
    stored = storage.get_raw_object(raw.id, raw.user_id)
    assert isinstance(stored, dict)
    return _current_turn_file_attachment(
        filename=str(metadata.get("filename") or "attachment"),
        file_ingestion={
            "raw_object_id": raw.id,
            "extraction": {
                "success": True,
                "text_success": True,
                "chars": len(raw.raw_content),
            },
        },
        raw=stored,
        storage=storage,
    )


def _patch_attachment_synthesis(runtime, monkeypatch):  # noqa: ANN001
    seen: list[tuple[str, list[dict[str, Any]]]] = []

    async def generate(context, message, attachments):  # noqa: ANN001, ARG001
        snapshot = [dict(item) for item in (attachments or [])]
        seen.append((str(message), snapshot))
        names = [str(item.get("filename") or "attachment") for item in snapshot]
        return {"content": "Синтетическая сводка: " + ", ".join(names), "tools_used": []}

    monkeypatch.setattr(runtime, "_prepare_context", _prepare_without_retrieval)
    monkeypatch.setattr(runtime, "_generate_response", generate)
    return seen


@pytest.mark.asyncio
async def test_three_separate_upload_turns_restore_one_exact_complete_active_set(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    files = [
        _stored_file(storage, filename=f"doc-{index}.txt", text=f"DOC-{index}|" + chr(64 + index) * 14_994)
        for index in range(1, 4)
    ]
    runtime = AgentRuntime(replace(settings, verify_answers=False), storage, llm=_NeverModel())
    seen = _patch_attachment_synthesis(runtime, monkeypatch)
    conversation_id: str | None = None
    for index, raw in enumerate(files, start=1):
        uploaded = await runtime.chat(
            "alice",
            f"Это документ {index}",
            actor=_actor(),
            conversation_id=conversation_id,
            attachments=[_current_attachment(storage, raw)],
            enable_tools=False,
        )
        conversation_id = str(uploaded["conversation_id"])

    summary = await runtime.chat(
        "alice",
        "Обобщи эти три документа",
        actor=_actor(),
        conversation_id=conversation_id,
        attachments=[],
        enable_tools=False,
    )

    final_attachments = seen[-1][1]
    assert [item["raw_object_id"] for item in final_attachments] == [raw.id for raw in files]
    assert [item["filename"] for item in final_attachments] == [f"doc-{index}.txt" for index in range(1, 4)]
    assert all(len(str(item.get("transient_text") or "")) == 15_000 for item in final_attachments)
    assert summary["restored_attachment_count"] == 3
    assert summary["attachment_context_expected_count"] == 3
    assert summary["attachment_context_readable_count"] == 3
    assert summary["attachment_coverage_complete"] is True
    assert "повторно" not in summary["message"].casefold()


@pytest.mark.asyncio
async def test_current_third_upload_caption_combines_it_with_two_prior_upload_origins(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    files = [
        _stored_file(storage, filename=f"caption-{index}.txt", text=f"CAPTION-{index}")
        for index in range(1, 4)
    ]
    runtime = AgentRuntime(replace(settings, verify_answers=False), storage, llm=_NeverModel())
    seen = _patch_attachment_synthesis(runtime, monkeypatch)
    conversation_id: str | None = None
    for index, raw in enumerate(files[:2], start=1):
        uploaded = await runtime.chat(
            "alice",
            f"Файл {index}",
            actor=_actor(),
            conversation_id=conversation_id,
            attachments=[_current_attachment(storage, raw)],
            enable_tools=False,
        )
        conversation_id = str(uploaded["conversation_id"])

    summary = await runtime.chat(
        "alice",
        "Сделай общую сводку по трём последним документам",
        actor=_actor(),
        conversation_id=conversation_id,
        attachments=[_current_attachment(storage, files[2])],
        enable_tools=False,
    )

    assert [item["raw_object_id"] for item in seen[-1][1]] == [raw.id for raw in files]
    assert summary["restored_attachment_count"] == 2
    assert summary["attachment_context_expected_count"] == 3
    assert summary["attachment_context_readable_count"] == 3
    assert summary["attachment_context_available"] is True


@pytest.mark.asyncio
async def test_an_authoritatively_empty_document_is_an_available_set_member(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    files = [
        _stored_file(storage, filename="nonempty-a.txt", text="ALPHA"),
        _stored_file(storage, filename="empty.txt", text=""),
        _stored_file(storage, filename="nonempty-b.txt", text="BETA"),
    ]
    runtime = AgentRuntime(replace(settings, verify_answers=False), storage, llm=_NeverModel())
    seen = _patch_attachment_synthesis(runtime, monkeypatch)
    conversation_id: str | None = None
    for index, raw in enumerate(files, start=1):
        uploaded = await runtime.chat(
            "alice",
            f"Материал {index}",
            actor=_actor(),
            conversation_id=conversation_id,
            attachments=[_current_attachment(storage, raw)],
            enable_tools=False,
        )
        conversation_id = str(uploaded["conversation_id"])

    summary = await runtime.chat(
        "alice",
        "Обобщи все три документа",
        actor=_actor(),
        conversation_id=conversation_id,
        enable_tools=False,
    )

    assert len(seen[-1][1]) == 3
    assert seen[-1][1][1]["empty_text"] is True
    assert summary["attachment_context_readable_count"] == 3
    assert summary["attachment_context_available"] is True
    assert "недоста" not in summary["message"].casefold()


@pytest.mark.asyncio
async def test_explicit_four_file_summary_with_only_three_uploads_is_honestly_incomplete(
    settings,
    storage,
    monkeypatch,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    files = [
        _stored_file(storage, filename=f"only-{index}.txt", text=f"ONLY-{index}") for index in range(1, 4)
    ]
    runtime = AgentRuntime(replace(settings, verify_answers=False), storage, llm=_NeverModel())
    seen = _patch_attachment_synthesis(runtime, monkeypatch)
    conversation_id: str | None = None
    for index, raw in enumerate(files, start=1):
        uploaded = await runtime.chat(
            "alice",
            f"Загрузка {index}",
            actor=_actor(),
            conversation_id=conversation_id,
            attachments=[_current_attachment(storage, raw)],
            enable_tools=False,
        )
        conversation_id = str(uploaded["conversation_id"])
    generated_before = len(seen)

    summary = await runtime.chat(
        "alice",
        "Подготовь сводку по 4 последним файлам",
        actor=_actor(),
        conversation_id=conversation_id,
        enable_tools=False,
    )

    assert len(seen) == generated_before, "the model was asked to fill a missing fourth file"
    assert summary["attachment_context_expected_count"] == 4
    # Evidence publication is all-or-none: three rows were restored, but none
    # may be presented as the requested four-file set or reach synthesis.
    assert summary["restored_attachment_count"] == 3
    assert summary["attachment_context_readable_count"] == 0
    assert "0 из 4" in summary["message"]
    assert "неизвест" in summary["message"].casefold()


def test_multi_restore_never_backfills_an_unowned_upload_slot(settings, storage) -> None:
    storage.ensure_user("alice", preset_key="owner")
    owned = [
        _stored_file(storage, filename=f"owned-{index}.txt", text=f"OWNED-{index}") for index in range(2)
    ]
    foreign = _stored_file(
        storage,
        filename="foreign.txt",
        text="FOREIGN-CANARY",
        tenant="foreign-tenant",
        uploader="foreign-tenant",
    )
    now = datetime.now(UTC).isoformat()
    history = [
        {
            "role": "user",
            "content": f"upload-{index}",
            "created_at": now,
            "metadata_json": json.dumps(
                {
                    "had_attachments": True,
                    "attachment_count": 1,
                    "attachment_origin": "upload",
                    "conversation_attachment_raw_ids": [raw.id],
                }
            ),
        }
        for index, raw in enumerate([*owned, foreign], start=1)
    ]
    runtime = AgentRuntime(settings, storage)

    restored, expected = runtime._restore_conversation_attachments(  # noqa: SLF001
        "Обобщи эти три документа",
        history,
        tenant_id="alice",
        person_id="alice",
        allow_file_read=True,
    )

    assert [item["raw_object_id"] for item in restored] == [raw.id for raw in owned]
    assert expected == 3
    assert "FOREIGN-CANARY" not in json.dumps(restored, ensure_ascii=False)
