"""Independent mutation checks for the content-free Office structure contract.

The fixtures are wholly synthetic.  They pin the two properties that make the
index safe to use for exhaustive answers: it is bound to the exact legacy text,
and its durable form contains locations rather than copied cell values.
"""

from __future__ import annotations

import copy
import hashlib
import io
import json

from docx import Document
from openpyxl import Workbook

from friday.documents import DocumentExtractor, validate_office_structure_index


def _docx_bytes(document: Document) -> bytes:
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _xlsx_bytes(workbook: Workbook) -> bytes:
    stream = io.BytesIO()
    workbook.save(stream)
    workbook.close()
    return stream.getvalue()


def _index(result) -> dict:
    index = result.office_structure_index
    assert isinstance(index, dict)
    assert validate_office_structure_index(index, result.text) == index
    return index


def test_docx_keeps_legacy_text_but_indexes_true_interleaved_order_and_split_runs():
    document = Document()
    document.add_paragraph("PARAGRAPH-BEFORE")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "ФИО"
    table.cell(0, 1).text = "Роль"
    row = table.add_row()
    name_paragraph = row.cells[0].paragraphs[0]
    name_paragraph.add_run("PERSON-SPLIT-")
    name_paragraph.add_run("VALUE")
    row.cells[1].text = "ROLE-SENTINEL"
    document.add_paragraph("PARAGRAPH-AFTER")

    result = DocumentExtractor().extract(_docx_bytes(document), "synthetic.docx")
    index = _index(result)

    assert result.text == (
        "PARAGRAPH-BEFORE\nPARAGRAPH-AFTER\nФИО | Роль\nPERSON-SPLIT-VALUE | ROLE-SENTINEL"
    )
    assert [block["kind"] for block in index["blocks"]] == ["paragraph", "table", "paragraph"]
    assert index["text_sha256"] == hashlib.sha256(result.text.encode("utf-8")).hexdigest()
    assert len(index["record_sets"]) == 1
    assert index["record_sets"][0]["records_total"] == 1
    assert len(index["candidate_refs"]) == 1
    start, end = index["candidate_refs"][0]["text_span"]
    assert result.text[start:end] == "PERSON-SPLIT-VALUE"

    encoded = json.dumps(index, ensure_ascii=False, sort_keys=True)
    for literal in (
        "PARAGRAPH-BEFORE",
        "PARAGRAPH-AFTER",
        "PERSON-SPLIT-VALUE",
        "ROLE-SENTINEL",
    ):
        assert literal not in encoded


def test_docx_merged_header_is_one_header_and_never_a_record():
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "ФИО"
    row = table.add_row()
    row.cells[0].text = "MERGED-HEADER-PERSON"
    row.cells[1].text = "Описание"

    result = DocumentExtractor().extract(_docx_bytes(document), "merged.docx")
    index = _index(result)

    table_rows = next(block["rows"] for block in index["blocks"] if block["kind"] == "table")
    assert table_rows[0]["role"] == "header"
    assert table_rows[1]["role"] != "record"
    # A header merged across two logical columns does not identify one exact
    # person column.  V1 must fail closed rather than multiply it into records.
    assert index["record_sets"] == []
    assert index["candidate_refs"] == []


def test_xlsx_formula_without_a_cached_result_removes_completeness_without_changing_text():
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "SHEET-PRIVATE-SENTINEL"
    sheet.append(["ФИО", "Роль"])
    sheet.append(["XLSX-PERSON", "Оператор"])
    sheet.append(["FORMULA-ROW", "=1+1"])

    result = DocumentExtractor().extract(_xlsx_bytes(workbook), "synthetic.xlsx")
    index = _index(result)

    assert result.text.startswith("--- Sheet: SHEET-PRIVATE-SENTINEL ---\n")
    assert "ФИО | Роль\nXLSX-PERSON | Оператор\nFORMULA-ROW | " in result.text
    assert index["complete"] is False
    assert any("formula" in reason for reason in index["coverage"]["reasons"])
    encoded = json.dumps(index, ensure_ascii=False, sort_keys=True)
    assert "SHEET-PRIVATE-SENTINEL" not in encoded
    assert "XLSX-PERSON" not in encoded


def test_docx_unread_headers_and_nested_tables_are_declared_incomplete():
    document = Document()
    document.sections[0].header.paragraphs[0].text = "PRIVATE-HEADER-SENTINEL"
    outer = document.add_table(rows=1, cols=1)
    outer.cell(0, 0).text = "OUTER-CELL"
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "PRIVATE-NESTED-SENTINEL"

    result = DocumentExtractor().extract(_docx_bytes(document), "unsupported.docx")
    index = _index(result)

    assert index["complete"] is False
    reasons = " ".join(index["coverage"]["reasons"])
    assert "header" in reasons
    assert "nested" in reasons
    encoded = json.dumps(index, ensure_ascii=False, sort_keys=True)
    assert "PRIVATE-HEADER-SENTINEL" not in encoded
    assert "PRIVATE-NESTED-SENTINEL" not in encoded


def test_validator_rejects_hash_spans_references_and_schema_extensions():
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["ФИО", "Роль"])
    sheet.append(["VALIDATOR-PERSON", "Инженер"])
    result = DocumentExtractor().extract(_xlsx_bytes(workbook), "validator.xlsx")
    original = _index(result)

    wrong_text = result.text.replace("VALIDATOR-PERSON", "OTHER-PERSON")
    assert validate_office_structure_index(original, wrong_text) is None

    bad_span = copy.deepcopy(original)
    bad_span["candidate_refs"][0]["text_span"] = [0, len(result.text) + 1]
    assert validate_office_structure_index(bad_span, result.text) is None

    dangling = copy.deepcopy(original)
    dangling["candidate_refs"][0]["record_id"] = "row-does-not-exist"
    assert validate_office_structure_index(dangling, result.text) is None

    literal_extension = copy.deepcopy(original)
    literal_extension["candidate_refs"][0]["value"] = "VALIDATOR-PERSON"
    assert validate_office_structure_index(literal_extension, result.text) is None

    missing_candidate = copy.deepcopy(original)
    missing_candidate["candidate_refs"] = []
    assert validate_office_structure_index(missing_candidate, result.text) is None

    record_id = original["record_sets"][0]["record_ids"][0]
    coordinated_omission = copy.deepcopy(original)
    coordinated_omission["record_sets"][0]["record_ids"] = []
    coordinated_omission["record_sets"][0]["records_total"] = 0
    coordinated_omission["candidate_refs"] = []
    omitted_row = next(
        row
        for block in coordinated_omission["blocks"]
        for row in block.get("rows", [])
        if row["id"] == record_id
    )
    omitted_row["role"] = "unknown"
    assert validate_office_structure_index(coordinated_omission, result.text) is None

    wrong_role = copy.deepcopy(original)
    target_row = next(
        row for block in wrong_role["blocks"] for row in block.get("rows", []) if row["id"] == record_id
    )
    target_row["role"] = "header"
    assert validate_office_structure_index(wrong_role, result.text) is None

    missing_source_order = copy.deepcopy(original)
    missing_source_order["blocks"][0].pop("source_order")
    assert validate_office_structure_index(missing_source_order, result.text) is None

    dishonest_complete = copy.deepcopy(original)
    dishonest_complete["coverage"]["reasons"] = ["index_budget"]
    dishonest_complete["complete"] = True
    assert validate_office_structure_index(dishonest_complete, result.text) is None


def test_office_text_budget_is_visible_as_incomplete_structure():
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["ФИО", "Описание"])
    for number in range(40):
        sheet.append([f"PERSON-{number:03d}", "X" * 600])

    result = DocumentExtractor(max_text_chars=10_000).extract(
        _xlsx_bytes(workbook),
        "bounded.xlsx",
    )
    index = _index(result)

    assert len(result.text) == 10_000
    assert result.metadata.get("text_truncated") is True
    assert index["complete"] is False
    assert index["coverage"]["reasons"]


def test_validator_rejects_a_coordinated_source_order_rewrite():
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["ФИО", "Роль"])
    sheet.append(["FIRST-PERSON", "Первая"])
    sheet.append(["SECOND-PERSON", "Вторая"])
    result = DocumentExtractor().extract(_xlsx_bytes(workbook), "ordered.xlsx")
    original = _index(result)

    reordered = copy.deepcopy(original)
    reordered["record_sets"][0]["record_ids"].reverse()
    reordered["candidate_refs"].reverse()
    assert validate_office_structure_index(reordered, result.text) is None


def test_preheader_data_and_person_column_footer_are_not_silently_omitted():
    preheader = Workbook()
    sheet = preheader.active
    sheet.append(["PREHEADER-PERSON", "Роль"])
    sheet.append(["ФИО", "Роль"])
    sheet.append(["AFTER-HEADER-PERSON", "Роль"])
    preheader_result = DocumentExtractor().extract(
        _xlsx_bytes(preheader),
        "preheader.xlsx",
    )
    preheader_index = _index(preheader_result)
    assert preheader_index["complete"] is True
    assert preheader_index["record_sets"] == []
    assert preheader_index["candidate_refs"] == []

    role_word = Workbook()
    sheet = role_word.active
    sheet.append(["ФИО", "Роль"])
    sheet.append(["FIRST-PERSON", "Оператор"])
    sheet.append(["SECOND-PERSON", "Total Quality Manager"])
    role_result = DocumentExtractor().extract(_xlsx_bytes(role_word), "role-word.xlsx")
    role_index = _index(role_result)
    assert role_index["record_sets"][0]["records_total"] == 2

    footer_word = Workbook()
    sheet = footer_word.active
    sheet.append(["ФИО", "Роль"])
    sheet.append(["FIRST-PERSON", "Оператор"])
    sheet.append(["Итого PERSON-LIKE", "Оператор"])
    footer_result = DocumentExtractor().extract(_xlsx_bytes(footer_word), "footer-word.xlsx")
    footer_index = _index(footer_result)
    assert footer_index["record_sets"] == []
    assert footer_index["candidate_refs"] == []
