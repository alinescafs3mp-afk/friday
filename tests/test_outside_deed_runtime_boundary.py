"""The outside-deed guard protects the complete delivery path.

Everything in this module is synthetic.  These tests exercise ``AgentRuntime.chat``
rather than only the text classifier: a model answer can already carry a real
tool effect, and it can still be changed by verification, repair, or the late
document builder before it reaches the person.
"""

from __future__ import annotations

import io
import json
from dataclasses import replace
from pathlib import Path

import pytest
from openpyxl import Workbook

from friday.agent_runtime import (
    _CANNOT_ACT_OUTSIDE,
    _UNCONFIRMED_SUPPORTED_DEED,
    AgentContext,
    AgentRuntime,
    FileRegistrationKind,
    _attachment_whole_document_task,
    _claims_an_unconfirmed_supported_deed,
    _file_evidence_view_of,
    _normalize_passive_attachment_preparation_preamble,
)
from friday.execution_kernel import ExecutionKernel
from friday.ingestion import IngestionPipeline
from friday.knowledge_graph import KnowledgeGraph
from friday.permissions import ActorContext, AuthorizationService

_FALSE_OUTSIDE_REPORT = "Я уже заказала курьера к служебному входу."
_OBJECT_FIRST_FALSE_OUTSIDE_REPORT = "Такси заказала."
_PASSIVE_TICKET_FALSE_OUTSIDE_REPORT = "Билет уже оформлен. Подтверждение пришло на почту."
_PRINT_RESULT_FALSE_OUTSIDE_REPORT = "Печать завершена. Бумажная копия готова."
_REPAIRED_FALSE_OUTSIDE_REPORT = "Курьер вызван и уже едет к служебному входу."
_BARE_REFUSAL = "Я не могу выполнить это действие."
_DOCUMENT = {
    "kind": "document",
    "filename": "synthetic-note.txt",
    "mime_type": "text/plain",
    "content_base64": "c3ludGhldGlj",
}
_VOICE = {
    "kind": "voice",
    "mime_type": "audio/ogg",
    "audio_base64": "c3ludGhldGlj",
    "duration_sec": 1.0,
}


class _EnabledButPatchedLLM:
    """The integration road is patched at its named seam, never sent to a model."""

    enabled = True
    model = "synthetic-runtime-double"
    total_budget_sec = 5.0

    async def chat(self, messages, **kwargs):  # pragma: no cover - every caller is patched
        del messages, kwargs
        raise AssertionError("unexpected model call")


async def _simple_context(user_id, message, conversation_id, **kwargs):  # noqa: ANN001
    del message, kwargs
    return AgentContext(
        conversation_id=conversation_id,
        user_id=user_id,
        person_id=user_id,
        answer_mode="general_conversation",
    )


def _actor() -> ActorContext:
    return ActorContext(user_id="alice", preset_key="owner", source="test")


def _runtime(settings, storage, monkeypatch, *, verify_answers: bool = False) -> AgentRuntime:
    storage.ensure_user("alice", preset_key="owner")
    runtime = AgentRuntime(
        replace(
            settings,
            verify_answers=verify_answers,
            verify_min_answer_chars=1,
        ),
        storage,
        llm=_EnabledButPatchedLLM(),
    )
    monkeypatch.setattr(runtime, "_prepare_context", _simple_context)
    return runtime


async def _registered_text_attachment(
    settings,
    storage,
    *,
    filename: str,
    source_text: str,
) -> dict[str, str]:
    """Create the same durable SQLite + content-addressed file source used in production."""

    outcome = await IngestionPipeline(settings, storage, KnowledgeGraph(storage)).ingest_file(
        "alice",
        None,
        source_text.encode(),
        filename=filename,
        mime_type="text/plain",
        metadata={"uploaded_by": "alice"},
        source_ref=f"test-file:{filename}",
    )
    raw_id = str(outcome["raw_object_id"])
    raw = storage.get_raw_object(raw_id, "alice")
    assert raw is not None
    stored = json.loads(str(raw["metadata_json"] or "{}"))
    stored_path = Path(settings.files_dir) / str(stored["stored_path"])
    assert stored_path.is_file()
    assert stored_path.read_bytes() == source_text.encode()
    return {"raw_object_id": raw_id}


async def _registered_docx_attachment(
    settings,
    storage,
    *,
    filename: str,
    source_text: str,
) -> dict[str, str]:
    """Create a real trusted DOCX Raw like the production reproduction."""

    from docx import Document

    document = Document()
    document.add_heading("Структура", level=1)
    document.add_paragraph(source_text)
    payload = io.BytesIO()
    document.save(payload)
    outcome = await IngestionPipeline(settings, storage, KnowledgeGraph(storage)).ingest_file(
        "alice",
        None,
        payload.getvalue(),
        filename=filename,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        metadata={"uploaded_by": "alice"},
        source_ref=f"test-docx:{filename}",
    )
    return {"raw_object_id": str(outcome["raw_object_id"])}


async def _registered_xlsx_attachment(
    settings,
    storage,
    *,
    filename: str,
    source_text: str,
) -> dict[str, str]:
    """Create a real trusted XLSX Raw like the second live reproduction."""

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "График"
    sheet.append(["Раздел", "Содержание"])
    sheet.append(["Источник", source_text])
    payload = io.BytesIO()
    workbook.save(payload)
    workbook.close()
    outcome = await IngestionPipeline(settings, storage, KnowledgeGraph(storage)).ingest_file(
        "alice",
        None,
        payload.getvalue(),
        filename=filename,
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        metadata={"uploaded_by": "alice"},
        source_ref=f"test-xlsx:{filename}",
    )
    return {"raw_object_id": str(outcome["raw_object_id"])}


@pytest.mark.parametrize(
    ("source_text", "truthful_summary"),
    [
        (
            "Статус заказа\nЗаказ оформлен 11 августа 2026 года оператором поставщика.",
            "По данным файла: Заказ оформлен 11 августа 2026 года оператором поставщика.",
        ),
        (
            "Статус документа\nДокумент подготовлен 11 августа 2026 года отделом снабжения.",
            "По данным файла: Документ подготовлен 11 августа 2026 года отделом снабжения.",
        ),
    ],
    ids=("outside-action-source-state", "supported-file-source-state"),
)
@pytest.mark.asyncio
async def test_a_bare_upload_summary_keeps_a_truthful_state_from_its_source(
    settings,
    storage,
    monkeypatch,
    tmp_path: Path,
    source_text: str,
    truthful_summary: str,
) -> None:
    """A trusted bare upload is reviewed; a sourced state is not Friday's own deed."""

    assert Path(settings.database_path).is_relative_to(tmp_path)
    runtime = _runtime(settings, storage, monkeypatch)
    attachment = await _registered_text_attachment(
        settings,
        storage,
        filename="synthetic-source-status.txt",
        source_text=source_text,
    )

    prepared_messages: list[str] = []
    shown_sources: list[str] = []
    shown_evidence = []

    async def prepare(user_id, message, conversation_id, **kwargs):  # noqa: ANN001
        prepared_messages.append(str(message))
        return await _simple_context(user_id, message, conversation_id, **kwargs)

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message
        shown_sources.extend(str(item.get("transient_text") or "") for item in attachments)
        shown_evidence.extend(_file_evidence_view_of(item) for item in attachments)
        return {"content": truthful_summary, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_prepare_context", prepare)
    monkeypatch.setattr(runtime, "_generate_response", generate)
    review = await runtime.chat(
        "alice",
        "Загружен документ: synthetic-source-status.txt",
        actor=_actor(),
        attachments=[attachment],
        enable_tools=False,
        synthetic_document_notice=True,
    )

    assert prepared_messages == ["Загружен документ: synthetic-source-status.txt"]
    assert shown_sources == [source_text]
    assert len(shown_evidence) == 1
    evidence = shown_evidence[0]
    assert evidence is not None
    assert evidence.registration is FileRegistrationKind.VALID
    assert evidence.disk_verified is True
    assert evidence.source_readable is True
    assert evidence.source_complete is True
    assert evidence.verification_eligible is True
    assert review["message"] == truthful_summary
    assert review["tools_used"] == []
    assert "Быстрый обзор" not in review["message"]
    stored = storage.get_message(str(review["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    output_guards = metadata["structural"].get("output_guards", {})
    assert output_guards.get("outside_deed_replaced") is not True
    assert output_guards.get("supported_deed_replaced") is not True


@pytest.mark.asyncio
async def test_a_bare_upload_keeps_a_passive_preparation_preamble_with_substantive_review(
    settings,
    storage,
    monkeypatch,
) -> None:
    """A review preamble is not a claim that Friday created a new file."""

    runtime = _runtime(settings, storage, monkeypatch)
    attachment = await _registered_text_attachment(
        settings,
        storage,
        filename="synthetic-passive-preparation.txt",
        source_text="Материал описывает структуру проекта, текущий статус и следующие шаги.",
    )
    model_answer = (
        "Документ подготовлен: материал описывает структуру проекта, текущий статус "
        "и следующие шаги. Ниже приведено содержательное ревью исходного текста."
    )

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": model_answer, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Загружен документ: synthetic-passive-preparation.txt",
        actor=_actor(),
        attachments=[attachment],
        enable_tools=False,
        synthetic_document_notice=True,
    )

    assert reply["message"] != _UNCONFIRMED_SUPPORTED_DEED
    assert "материал описывает структуру проекта" in reply["message"].casefold()
    assert "содержательное ревью" in reply["message"].casefold()
    stored = storage.get_message(str(reply["message_id"]), "alice")
    assert stored is not None
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    output_guards = metadata["structural"].get("output_guards", {})
    assert output_guards.get("supported_deed_replaced") is not True


@pytest.mark.asyncio
async def test_a_verified_bare_docx_keeps_content_level_ready_objects(
    settings,
    storage,
    monkeypatch,
) -> None:
    """Regression for the live 2026-08-25 20:33 bare-upload failure."""

    runtime = _runtime(settings, storage, monkeypatch)
    filename = "Образец рапортов.docx"
    attachment = await _registered_docx_attachment(
        settings,
        storage,
        filename=filename,
        source_text=(
            "Рапорт об основном отпуске. Рапорт о прекращении командировки. "
            "Рапорт об отпуске по личным обстоятельствам."
        ),
    )
    model_answer = (
        "Документ содержит три готовых образца рапортов: об основном отпуске, "
        "о прекращении командировки и об отпуске по личным обстоятельствам."
    )

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": model_answer, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        f"Загружен документ: {filename}",
        actor=_actor(),
        attachments=[attachment],
        enable_tools=False,
        synthetic_document_notice=True,
    )

    assert reply["message"] == model_answer
    assert reply["attachment_coverage_complete"] is True
    assert reply["attachment_verification_complete"] is True
    stored = storage.get_message(str(reply["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    guards = metadata["structural"].get("output_guards", {})
    assert guards.get("supported_deed_replaced") is not True


@pytest.mark.parametrize(
    "claim",
    [
        "Документ содержит три готовых образца рапортов.",
        "Исходный документ представляет собой готовый набор форм.",
        "Загруженный файл подготовлен как подборка шаблонов.",
        "Образец рапортов.docx содержит готовые формы.",
    ],
)
def test_a_complete_read_only_source_description_is_not_a_new_file_deed(claim: str) -> None:
    assert not _claims_an_unconfirmed_supported_deed(
        claim,
        has_file=False,
        reminder_succeeded=False,
        read_only_attachment_review=True,
        read_only_attachment_descriptors=[
            "Образец рапортов.docx application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ],
    )


@pytest.mark.parametrize(
    "claim",
    [
        "PDF готов.",
        "Документ готов и прикреплён.",
        "Я подготовила документ с образцами рапортов.",
        "Документ содержит три готовых образца. Прикрепляю новый PDF.",
        "Документ содержит готовый образец. Я загрузила новый PDF.",
        "Документ содержит готовый образец. Новый PDF загружен.",
        "Документ представляет собой готовый шаблон, скачать его можно здесь.",
        "Образец рапортов.docx готов, а новый PDF отправлен.",
        "Документ содержит готовый образец. Напоминание установлено.",
    ],
)
def test_read_only_source_scope_does_not_license_a_current_or_mixed_deed(claim: str) -> None:
    assert _claims_an_unconfirmed_supported_deed(
        claim,
        has_file=False,
        reminder_succeeded=False,
        read_only_attachment_review=True,
        read_only_attachment_descriptors=["Образец рапортов.docx"],
    )


@pytest.mark.parametrize(
    ("kind", "filename", "source_text", "model_answer"),
    [
        pytest.param(
            "docx",
            "Анкета Артемьев.docx",
            "Анкета Артемьева содержит дату рождения, место рождения и учётные поля.",
            "**Назначение и структура**\n\nДокумент подготовлен в виде анкеты Артемьева. "
            "В ней перечислены дата рождения, место рождения и учётные поля.",
            id="headed-docx-questionnaire",
        ),
        pytest.param(
            "xlsx",
            "график отпусков 2025-2026.xlsx",
            "График отпусков содержит периоды и строки личного состава.",
            "**Назначение и структура**\n\nФайл подготовлен как табличный график отпусков. "
            "Строки содержат периоды и сведения о личном составе.",
            id="headed-xlsx-schedule",
        ),
        pytest.param(
            "docx",
            "Анкета Артемьев.docx",
            "Анкета Артемьева содержит дату рождения, место рождения и учётные поля.",
            "**Назначение и структура**\n\nДокумент представляет собой готовую анкету Артемьева. "
            "В ней перечислены дата рождения, место рождения и учётные поля.",
            id="ready-docx-questionnaire",
        ),
        pytest.param(
            "xlsx",
            "график отпусков 2025-2026.xlsx",
            "График отпусков содержит периоды и строки личного состава.",
            "**Назначение и структура**\n\nФайл представляет собой готовый табличный график отпусков. "
            "Строки содержат периоды и сведения о личном составе.",
            id="ready-xlsx-schedule",
        ),
    ],
)
@pytest.mark.asyncio
async def test_a_headed_read_only_office_review_is_never_replaced_as_a_deed(
    settings,
    storage,
    monkeypatch,
    kind: str,
    filename: str,
    source_text: str,
    model_answer: str,
) -> None:
    """Regression for the two readable live turns at 09:42 and 09:43."""

    runtime = _runtime(settings, storage, monkeypatch)
    register = _registered_docx_attachment if kind == "docx" else _registered_xlsx_attachment
    attachment = await register(
        settings,
        storage,
        filename=filename,
        source_text=source_text,
    )

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": model_answer, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        f"Загружен документ: {filename}",
        actor=_actor(),
        attachments=[attachment],
        enable_tools=False,
        synthetic_document_notice=True,
    )

    assert reply["message"] == model_answer.replace(
        "Документ подготовлен",
        "Документ представлен",
    ).replace("Файл подготовлен", "Файл представлен").replace(
        "представляет собой готовую",
        "представляет собой",
    ).replace(
        "представляет собой готовый",
        "представляет собой",
    )
    assert reply["attachment_coverage_complete"] is True
    assert reply["attachment_verification_complete"] is True
    stored = storage.get_message(str(reply["message_id"]), "alice")
    assert stored is not None
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    assert metadata["attachment_context_readable_count"] == 1
    assert metadata["structural"].get("output_guards", {}).get("supported_deed_replaced") is not True


@pytest.mark.asyncio
async def test_an_explicit_structure_request_keeps_passive_source_summary(
    settings,
    storage,
    monkeypatch,
    tmp_path: Path,
) -> None:
    """A read-only transformation describes the source, not a new file effect."""

    assert Path(settings.database_path).is_relative_to(tmp_path)
    runtime = _runtime(settings, storage, monkeypatch)
    source_text = "Раздел один: статус проекта. Раздел два: следующие шаги."
    attachment = await _registered_docx_attachment(
        settings,
        storage,
        filename="посты+оружие.docx",
        source_text=source_text,
    )
    model_answer = "Документ подготовлен в виде двух разделов: статус проекта и следующие шаги."

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message
        assert len(attachments) == 1
        evidence = _file_evidence_view_of(attachments[0])
        assert evidence is not None
        assert evidence.source_complete is True
        assert evidence.verification_eligible is True
        return {"content": model_answer, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Структурируй инфу пожалуйста",
        actor=_actor(),
        attachments=[attachment],
        enable_tools=False,
    )

    assert reply["message"] == ("Документ представлен в виде двух разделов: статус проекта и следующие шаги.")
    stored = storage.get_message(str(reply["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    output_guards = metadata["structural"].get("output_guards", {})
    assert output_guards.get("supported_deed_replaced") is not True
    assert output_guards["passive_attachment_preamble_normalized"] is True


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "heading_prefix",
    [
        pytest.param("**Назначение и структура**\n\n", id="bold-newline"),
        pytest.param("**Назначение и структура** ", id="bold-same-line"),
        pytest.param("Назначение и структура\n\n", id="plain-newline"),
        pytest.param("# Обзор\n\n## Структура\n\n", id="two-markdown-headings"),
        pytest.param("- **Назначение и структура**\n\n", id="bullet-heading"),
        pytest.param("**Тип документа**\n\n", id="live-document-type-heading"),
        pytest.param("**Краткое описание**\n\n", id="short-description-heading"),
        pytest.param("**Формат и структура**\n\n", id="format-and-structure-heading"),
    ],
)
async def test_an_explicit_structure_request_keeps_a_headed_passive_source_summary(
    settings,
    storage,
    monkeypatch,
    heading_prefix: str,
) -> None:
    """A Markdown heading must not disable the proven read-only normalization."""

    runtime = _runtime(settings, storage, monkeypatch)
    attachment = await _registered_docx_attachment(
        settings,
        storage,
        filename="headed-structure.docx",
        source_text="Раздел один: статус. Раздел два: следующие шаги.",
    )
    model_answer = f"{heading_prefix}Документ подготовлен в виде двух разделов: статус и следующие шаги."

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": model_answer, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Структурируй инфу пожалуйста",
        actor=_actor(),
        attachments=[attachment],
        enable_tools=False,
    )

    assert reply["message"] == (
        f"{heading_prefix.strip()}\n\nДокумент представлен в виде двух разделов: статус и следующие шаги."
    )


@pytest.mark.parametrize(
    ("model_answer", "expected"),
    [
        (
            "**Документ подготовлен:** ниже приведены основные разделы и выводы.",
            "Документ представлен: ниже приведены основные разделы и выводы.",
        ),
        (
            "Структурированный документ подготовлен: ниже основные разделы и выводы.",
            "Документ представлен: ниже основные разделы и выводы.",
        ),
    ],
)
def test_markdown_attachment_preparation_leads_are_normalized(
    model_answer: str,
    expected: str,
) -> None:
    assert _normalize_passive_attachment_preparation_preamble(model_answer) == (expected, True)


@pytest.mark.parametrize(
    ("user_message", "expected"),
    [
        ("Структурируй инфу пожалуйста", "summary"),
        ("Структурируйте информацию", "summary"),
        ("Создай структурированный документ", ""),
        ("Структурированный документ готов", ""),
    ],
)
def test_structure_request_has_a_closed_read_only_scope(user_message: str, expected: str) -> None:
    assert _attachment_whole_document_task(user_message, file_count=1) == expected


@pytest.mark.asyncio
async def test_an_explicit_structure_request_does_not_hide_an_active_file_deed(
    settings,
    storage,
    monkeypatch,
) -> None:
    runtime = _runtime(settings, storage, monkeypatch)
    attachment = await _registered_text_attachment(
        settings,
        storage,
        filename="structure-active-deed.txt",
        source_text="Нейтральное содержимое источника.",
    )

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": "Прикрепляю файл.", "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Структурируй инфу пожалуйста",
        actor=_actor(),
        attachments=[attachment],
        enable_tools=False,
    )

    assert reply["message"] == _UNCONFIRMED_SUPPORTED_DEED


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "unsupported_claim",
    [
        "PDF готов.",
        "Напоминание установлено.",
        "Аудио готово.",
        "Документ подготовлен мной.",
        "Документ создан Пятницей.",
        "Документ подготовлен ботом.",
        "Документ подготовлен как другой отчёт.",
        "Документ подготовлен и готов к скачиванию.",
        "Документ подготовлен и доступен по ссылке.",
        "Документ подготовлен как анкета и доступен для скачивания.",
        "Документ подготовлен как анкета, её можно скачать.",
        "Документ подготовлен как анкета и опубликован по ссылке.",
        "Документ подготовлен как анкета и размещён в облаке.",
        "Документ подготовлен как анкета моделью.",
        "Документ подготовлен как анкета искусственным интеллектом.",
        "Документ подготовлен как анкета и сохранён на диске.",
        "Документ подготовлен как анкета и записан на диск.",
        "Документ подготовлен как анкета и добавлен в чат.",
        "Документ подготовлен как анкета и выслан вам.",
        "Документ подготовлен как анкета — можете скачать.",
        "Документ подготовлен как анкета — скачайте здесь.",
        "Документ подготовлен как анкета, ссылка ниже.",
        "Документ подготовлен как анкета, доступ к ней открыт.",
        "Документ подготовлен как анкета, лежит по ссылке.",
        "Документ подготовлен как анкета, находится в облаке.",
        "Документ подготовлен как анкета, вот ссылка.",
        "Документ представляет собой готовую анкету — скачать можно здесь.",
        "Документ представляет собой готовую анкету, загрузка доступна здесь.",
        "Документ представляет собой готовую анкету, доступ к ней предоставлен.",
        "Документ представляет собой готовую анкету, высылаю вам.",
        "Документ представляет собой готовую анкету, можете забрать.",
        "Документ представляет собой готовую анкету и хранится в облаке.",
        "Документ подготовлен как новая исправленная версия анкеты.",
        "Документ подготовлен как перевод анкеты на английский.",
        "Документ подготовлен как новая версия анкеты.",
        "Документ подготовлен как изменённая версия анкеты.",
        "Документ подготовлен как дополненная редакция анкеты.",
        "Документ подготовлен как анкета и направлен вам.",
        "Документ представляет собой готовую анкету и отдан вам.",
        "**PDF готов**\n\nДокумент подготовлен как анкета.",
        "Документ подготовлен как новый догов Газпрома.\n\nДокумент подготовлен как анкета.",
    ],
)
async def test_an_explicit_structure_request_does_not_hide_other_passive_deeds(
    settings,
    storage,
    monkeypatch,
    unsupported_claim: str,
) -> None:
    runtime = _runtime(settings, storage, monkeypatch)
    attachment = await _registered_text_attachment(
        settings,
        storage,
        filename="structure-passive-deed.txt",
        source_text="Нейтральное содержимое источника. Анкета содержит поля.",
    )

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": unsupported_claim, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Структурируй инфу пожалуйста",
        actor=_actor(),
        attachments=[attachment],
        enable_tools=False,
    )

    assert reply["message"] == _UNCONFIRMED_SUPPORTED_DEED


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "model_answer",
    [
        "Документ подготовлен как анкета — передаю вам.",
        "Документ подготовлен как анкета — высылаю вам.",
        "Документ подготовлен как анкета — забирайте.",
        "Документ подготовлен как анкета — скачивание здесь.",
        "Документ подготовлен как анкета — прикладываю к ответу.",
        "Файл представляет собой готовую анкету — передаю вам.",
        "Документ подготовлен как анкета Артемьева — держите.",
        "Документ подготовлен как анкета Артемьева — держи.",
        "Документ подготовлен как анкета Артемьева — шлю вам.",
        "Документ подготовлен как анкета Артемьева — вот вам.",
        "Документ подготовлен как анкета Артемьева — уже у вас.",
        "Документ подготовлен как анкета Артемьева — теперь ваш.",
        "Документ подготовлен как анкета Артемьева — он ваш.",
        "Документ подготовлен как анкета Артемьева — для вас.",
        "Документ подготовлен как анкета Артемьева — вам.",
        "Документ подготовлен как анкета Артемьева — у тебя.",
        "Документ подготовлен как анкета Артемьева — можете забрать.",
        "Документ подготовлен как анкета Артемьева — возьмите.",
        "Документ подготовлен как анкета Артемьева — выдаю вам.",
        "Документ подготовлен как анкета Артемьева — вручаю вам.",
        "Документ подготовлен как анкета Артемьева — посылаю вам.",
        "Документ подготовлен как анкета Артемьева — делюсь с вами.",
        "Документ подготовлен как анкета Артемьева — это вам.",
        "Документ подготовлен как анкета Артемьева — это для вас.",
        "Документ подготовлен как анкета Артемьева — теперь у вас.",
        "Документ подготовлен как анкета Артемьева — уже ваш.",
        "Документ подготовлен как анкета Артемьева — теперь он у вас.",
        "Документ подготовлен как анкета Артемьева — уже он ваш.",
        "Документ подготовлен как анкета Артемьева — она ваша.",
        "Документ подготовлен как анкета Артемьева — она уже ваша.",
        "Документ подготовлен как анкета Артемьева — теперь она ваша.",
        "Документ подготовлен как анкета Артемьева, она у вас.",
        "Документ подготовлен как анкета Артемьева уже у вас.",
        "Документ подготовлен как анкета Артемьева — можно взять.",
        "Документ подготовлен как анкета Артемьева — скачать можно здесь.",
        "Документ подготовлен как анкета Артемьева — забрать можно здесь.",
        "Документ подготовлен как анкета Артемьева — передал вам.",
        "Документ подготовлен как анкета Артемьева — передала вам.",
        "Документ подготовлен как анкета Артемьева — выслал вам.",
        "Документ подготовлен как анкета Артемьева — выдала вам.",
        "Документ подготовлен как анкета Артемьева — вручила вам.",
        "Документ подготовлен как анкета Артемьева — доставил вам.",
        "Документ подготовлен как анкета Артемьева. Прикрепляю.",
        "Документ подготовлен как анкета Артемьева. Отправляю вам.",
        "Документ подготовлен как анкета Артемьева. Держите.",
        "Документ подготовлен как анкета Артемьева. Он уже у вас.",
        "Файл представляет собой готовую анкету Артемьева. Высылаю вам.",
        "Документ подготовлен как анкета Артемьева. Вот вам.",
    ],
)
async def test_untrusted_attachment_text_cannot_authorize_a_carrier_suffix(
    settings,
    storage,
    monkeypatch,
    model_answer: str,
) -> None:
    runtime = _runtime(settings, storage, monkeypatch)
    attachment = await _registered_text_attachment(
        settings,
        storage,
        filename="quoted-carrier.txt",
        source_text=f"Цитата: {model_answer} Анкета содержит поля.",
    )

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": model_answer, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Структурируй инфу пожалуйста",
        actor=_actor(),
        attachments=[attachment],
        enable_tools=False,
    )

    assert reply["message"] == _UNCONFIRMED_SUPPORTED_DEED


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("source_text", "model_answer", "expected"),
    [
        (
            "Анкета для получения допуска с обязательными полями.",
            "Документ подготовлен как анкета для получения допуска с обязательными полями.",
            "Документ представлен как анкета для получения допуска с обязательными полями.",
        ),
        (
            "Инструкция по отправке отчёта, загрузке данных на сервер и хранению копий.",
            "Документ подготовлен как инструкция по отправке отчёта, загрузке данных "
            "на сервер и хранению копий.",
            "Документ представлен как инструкция по отправке отчёта, загрузке данных "
            "на сервер и хранению копий.",
        ),
        (
            "Журнал записей чат-бота и схема облака тегов.",
            "Документ представляет собой готовый журнал записей чат-бота и схему облака тегов.",
            "Документ представляет собой журнал записей чат-бота и схему облака тегов.",
        ),
    ],
)
async def test_grounded_attachment_content_nouns_are_not_carrier_claims(
    settings,
    storage,
    monkeypatch,
    source_text: str,
    model_answer: str,
    expected: str,
) -> None:
    runtime = _runtime(settings, storage, monkeypatch)
    attachment = await _registered_text_attachment(
        settings,
        storage,
        filename="grounded-description.txt",
        source_text=source_text,
    )

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": model_answer, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Структурируй инфу пожалуйста",
        actor=_actor(),
        attachments=[attachment],
        enable_tools=False,
    )

    assert reply["message"] == expected
    stored = storage.get_message(str(reply["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    output_guards = metadata["structural"].get("output_guards", {})
    assert output_guards.get("supported_deed_replaced") is not True
    assert output_guards["passive_attachment_preamble_normalized"] is True


@pytest.mark.asyncio
async def test_a_structure_request_without_an_attachment_cannot_claim_a_prepared_document(
    settings,
    storage,
    monkeypatch,
) -> None:
    runtime = _runtime(settings, storage, monkeypatch)

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {
            "content": "Документ подготовлен в виде двух разделов: статус и следующие шаги.",
            "tools_used": [],
            "_model_generated": True,
        }

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Структурируй инфу пожалуйста",
        actor=_actor(),
        enable_tools=False,
    )

    assert reply["message"] == _UNCONFIRMED_SUPPORTED_DEED


@pytest.mark.parametrize(
    ("model_claim", "blocked_answer", "guard_name"),
    [
        ("Такси заказала.", _CANNOT_ACT_OUTSIDE, "outside_deed_replaced"),
        ("Прикрепляю файл.", _UNCONFIRMED_SUPPORTED_DEED, "supported_deed_replaced"),
    ],
    ids=("object-first-outside-deed", "implicit-supported-file-deed"),
)
@pytest.mark.asyncio
async def test_a_bare_upload_does_not_hide_an_active_model_deed(
    settings,
    storage,
    monkeypatch,
    tmp_path: Path,
    model_claim: str,
    blocked_answer: str,
    guard_name: str,
) -> None:
    """A bare-upload review cannot turn untrusted file text into a completed deed."""

    assert Path(settings.database_path).is_relative_to(tmp_path)
    runtime = _runtime(settings, storage, monkeypatch)
    source_text = f"Недоверенная строка внутри файла: {model_claim}"
    attachment = await _registered_text_attachment(
        settings,
        storage,
        filename="synthetic-active-deed.txt",
        source_text=source_text,
    )

    prepared_messages: list[str] = []
    shown_evidence = []

    async def prepare(user_id, message, conversation_id, **kwargs):  # noqa: ANN001
        prepared_messages.append(str(message))
        return await _simple_context(user_id, message, conversation_id, **kwargs)

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message
        shown_evidence.extend(_file_evidence_view_of(item) for item in attachments)
        return {"content": model_claim, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_prepare_context", prepare)
    monkeypatch.setattr(runtime, "_generate_response", generate)
    review = await runtime.chat(
        "alice",
        "Загружен документ: synthetic-active-deed.txt",
        actor=_actor(),
        attachments=[attachment],
        enable_tools=False,
        synthetic_document_notice=True,
    )

    assert prepared_messages == ["Загружен документ: synthetic-active-deed.txt"]
    assert len(shown_evidence) == 1
    evidence = shown_evidence[0]
    assert evidence is not None
    assert evidence.registration is FileRegistrationKind.VALID
    assert evidence.disk_verified is True
    assert evidence.source_readable is True
    assert evidence.source_complete is True
    assert evidence.verification_eligible is True
    assert review["message"] == blocked_answer
    assert review["message"] != model_claim
    assert review["tools_used"] == []
    stored = storage.get_message(str(review["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    assert metadata["structural"]["output_guards"][guard_name] is True


@pytest.mark.parametrize(
    ("user_message", "model_claim", "blocked_answer"),
    [
        (
            "Организуй физическую доставку к зданию.",
            "Заказ оформлен 11 августа 2026 года оператором поставщика.",
            _CANNOT_ACT_OUTSIDE,
        ),
        (
            "Подготовь синтетический документ.",
            "Документ подготовлен 11 августа 2026 года отделом снабжения.",
            _UNCONFIRMED_SUPPORTED_DEED,
        ),
    ],
    ids=("outside-action-claim", "supported-file-claim"),
)
@pytest.mark.asyncio
async def test_the_same_unattributed_claim_without_an_attachment_remains_blocked(
    settings,
    storage,
    monkeypatch,
    tmp_path: Path,
    user_message: str,
    model_claim: str,
    blocked_answer: str,
) -> None:
    assert Path(settings.database_path).is_relative_to(tmp_path)
    runtime = _runtime(settings, storage, monkeypatch)

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": model_claim, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        user_message,
        actor=_actor(),
        enable_tools=False,
    )

    assert reply["message"] == blocked_answer


@pytest.mark.asyncio
async def test_a_quoted_document_metadata_provenance_answer_is_not_replaced_as_a_file_deed(
    settings,
    storage,
    monkeypatch,
) -> None:
    runtime = _runtime(settings, storage, monkeypatch)
    model_answer = "Я взяла гриф из метаданных документа, где он сохранён."
    seen_quote: list[str] = []

    async def generate(context, message, attachments):  # noqa: ANN001
        del message, attachments
        seen_quote.append(str(context.reply_quote or ""))
        return {"content": model_answer, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "А откуда ты это взяла, из метаданных?",
        actor=_actor(),
        enable_tools=False,
        reply_to="В предыдущем ответе был указан синтетический гриф документа.",
    )

    assert seen_quote == ["В предыдущем ответе был указан синтетический гриф документа."]
    assert reply["message"] == model_answer
    stored = storage.get_message(str(reply["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    assert metadata["structural"].get("output_guards", {}).get("supported_deed_replaced") is not True


@pytest.mark.parametrize(
    "claim",
    [
        "Я отправила файл в этот чат.",
        "Файл готов.",
        "Готовый файл во вложении.",
        "Прикрепляю файл.",
        "Документ уже в чате.",
        "Я сделала документ.",
        "Я сформировала отчёт.",
        "Архив сформирован.",
        "Отчёт выгружен.",
        "Вложение отправлено.",
        "PDF готов и приложен.",
        "Напоминание установлено.",
        "Я добавила напоминание.",
        "Напомню вам завтра в 9:00.",
        "Уведомление создано.",
    ],
)
def test_a_supported_deed_needs_evidence_from_this_turn(claim: str) -> None:
    assert (
        _claims_an_unconfirmed_supported_deed(
            claim,
            has_file=False,
            reminder_succeeded=False,
        )
        is True
    )


@pytest.mark.parametrize(
    "claim",
    [
        "Держи файл.",
        "Прилагаю файл.",
        "Отправляю файл.",
        "Excel готов и приложен.",
        "Картинка готова.",
        "PNG отправлен.",
        "Файл отпра\u200bвлен.",
        "Файл от[прав](https://example.com)лен.",
        "Файл от`прав`лен.",
        "Напоминание поста\u200bвлено.",
        "Голосовое отправлено.",
        "Я озвучила ответ.",
        "Я отправила аудио.",
        "Голосовое уже в чате.",
    ],
)
def test_common_file_reminder_and_voice_completions_need_evidence(claim: str) -> None:
    assert _claims_an_unconfirmed_supported_deed(
        claim,
        has_file=False,
        reminder_succeeded=False,
        voice_succeeded=False,
    )


@pytest.mark.parametrize(
    "offer",
    [
        "Да, могу нарисовать картинку — готова помочь. Опиши сцену.",
        "Готова проанализировать изображение, когда ты его пришлёшь.",
        "Могу подготовить готовую картинку по твоему описанию.",
        "Готовый документ могу перевести после уточнения языка.",
        "Могу показать готовый документ после того, как ты его пришлёшь.",
        "Могу проверить готовый документ.",
        "Могу проанализировать готовый отчёт.",
        "Могу перевести готовый документ.",
        "Готова поделиться готовым документом.",
        "Могу ознакомиться с готовым документом.",
        "Да, я готова к этому и могу проанализировать картинку.",
        "Готова к работе: могу описать картинку после загрузки.",
        "Я готова, и могу сделать изображение.",
        "Да, готова — могу проанализировать картинку.",
        "Готова быстро проанализировать картинку.",
        "Готова сама описать картинку.",
        "Готова с радостью проверить картинку.",
        "Готова по твоему описанию проверить картинку.",
        "Картинку готова прямо сейчас проанализировать.",
        "Напоминание готова при необходимости поставить.",
        "Аудио готова без проблем записать.",
        "Напоминание готова поставить, когда ты назовёшь время.",
        "Аудио готова записать после уточнения текста.",
        "Yes, I can review an image after you upload it.",
        "I am able to prepare a document after you provide the source text.",
        "I can help analyze a picture once you upload it.",
    ],
)
def test_capability_and_readiness_propositions_are_not_completed_deeds(offer: str) -> None:
    assert not _claims_an_unconfirmed_supported_deed(
        offer,
        has_file=False,
        reminder_succeeded=False,
        voice_succeeded=False,
    )


@pytest.mark.parametrize(
    "claim",
    [
        "Могу помочь, но файл уже готов.",
        "Готова помочь с картинкой, но PNG уже отправлен.",
        "Могу проанализировать картинку и уже отправила её.",
        "Я уже создала картинку.",
        "Картинка создана.",
        "Готова помочь. Держи картинку.",
        "Могу показать готовый документ — он доступен по ссылке.",
        "Могу показать готовый документ: ссылка ниже.",
        "Могу показать готовый документ: вот ссылка https://example.test/file.",
        "Могу показать готовый документ: [скачать](https://example.test/file).",
        "Готовый документ доступен здесь, могу показать.",
        "Готовый документ уже у тебя, могу показать.",
        "Готовый документ для тебя.",
        "Готовый документ в чате.",
        "Готовый документ открывается по ссылке.",
        "Готовый документ: забирай по ссылке.",
        "Готовая картинка: она ниже.",
    ],
)
def test_capability_language_cannot_hide_an_independent_completed_effect(claim: str) -> None:
    assert _claims_an_unconfirmed_supported_deed(
        claim,
        has_file=False,
        reminder_succeeded=False,
        voice_succeeded=False,
    )


@pytest.mark.parametrize(
    "statement",
    [
        "На картинке художник нарисовал дом.",
        "Нарисованная пользователем картинка содержит дом.",
        "Я могу проанализировать нарисованную картинку.",
        "Картинка была нарисована художником.",
        "Пользователь нарисовал картинку вчера.",
    ],
)
def test_drawing_content_and_history_are_not_current_assistant_deeds(statement: str) -> None:
    assert not _claims_an_unconfirmed_supported_deed(
        statement,
        has_file=False,
        reminder_succeeded=False,
        voice_succeeded=False,
    )


@pytest.mark.parametrize(
    "statement",
    [
        "В инструкции сказано: файл создан.",
        "Если файл создан, отправьте ссылку.",
        "Файл будет создан завтра.",
        "Файл должен быть создан по шаблону.",
        "Файл создан клиентом.",
        "Файл создан не мной.",
        "Файл не отправлен.",
        "Документ не подготовлен.",
        "Файл ещё не прикреплён.",
        "Напоминание ещё не сохранено.",
        "Голосовое ещё не записано.",
        "Напоминание создано пользователем.",
        "Файл отправлен?",
        "Напоминание поставлено?",
    ],
)
def test_a_nonactual_or_reported_supported_deed_is_not_a_completion(statement: str) -> None:
    assert (
        _claims_an_unconfirmed_supported_deed(
            statement,
            has_file=False,
            reminder_succeeded=False,
        )
        is False
    )


@pytest.mark.parametrize(
    ("user_message", "model_answer"),
    [
        (
            "А ты можешь проверить документ?",
            "Да, могу проверить документ — готова помочь. Пришли файл.",
        ),
        (
            "Ты умеешь анализировать документы?",
            "Да. Готова проанализировать документ, когда ты его пришлёшь.",
        ),
        (
            "Способна ли ты сделать готовый PDF?",
            "Могу подготовить готовый PDF после того, как ты дашь содержание.",
        ),
        (
            "Can you review documents?",
            "Yes, I can review a document after you upload it.",
        ),
        (
            "Are you able to prepare documents?",
            "Yes. I am able to prepare a document after you provide the source text.",
        ),
        (
            "Ты уже нарисовала картинку?",
            "Нет, ещё не нарисовала. Готова помочь, когда будет описание.",
        ),
        (
            "Проверь документ.",
            "Готова проверить документ, но сначала нужен файл.",
        ),
    ],
    ids=(
        "live-russian-capability",
        "russian-skill",
        "russian-able",
        "english-can",
        "english-able",
        "past-question-truthful-negative",
        "imperative-honest-incomplete",
    ),
)
@pytest.mark.asyncio
async def test_nonactual_capability_answers_survive_the_complete_delivery_path(
    settings,
    storage,
    monkeypatch,
    user_message: str,
    model_answer: str,
) -> None:
    runtime = _runtime(settings, storage, monkeypatch)

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": model_answer, "tools_used": [], "_model_generated": True}

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        user_message,
        actor=_actor(),
        enable_tools=False,
    )

    assert reply["message"] == model_answer
    stored = storage.get_message(str(reply["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    guards = metadata["structural"].get("output_guards", {})
    assert guards.get("supported_deed_replaced") is not True


@pytest.mark.parametrize(
    ("model_answer", "expected", "replaced"),
    [
        (
            "Инструмент вернул ошибку: файл не создан. Могу повторить после нового запроса.",
            "Инструмент вернул ошибку: файл не создан. Могу повторить после нового запроса.",
            False,
        ),
        (
            "Инструмент вернул ошибку, но файл уже создан.",
            _UNCONFIRMED_SUPPORTED_DEED,
            True,
        ),
    ],
    ids=("truthful-tool-failure", "failed-tool-false-completion"),
)
@pytest.mark.asyncio
async def test_a_failed_tool_does_not_turn_an_offer_into_completion_evidence(
    settings,
    storage,
    monkeypatch,
    model_answer: str,
    expected: str,
    replaced: bool,
) -> None:
    runtime = _runtime(settings, storage, monkeypatch)

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {
            "content": model_answer,
            "tools_used": ["make_file"],
            "tool_evidence": [
                {"tool": "make_file", "success": False, "output": {"error": "synthetic_failure"}}
            ],
            "_model_generated": True,
        }

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Ты можешь создать файл?",
        actor=_actor(),
        enable_tools=False,
    )

    assert reply["message"] == expected
    stored = storage.get_message(str(reply["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    guards = metadata["structural"].get("output_guards", {})
    assert (guards.get("supported_deed_replaced") is True) is replaced


def test_supported_deed_evidence_is_bound_to_the_claimed_artifact() -> None:
    assert _claims_an_unconfirmed_supported_deed(
        "Файл отчёта по зарплате отправлен.",
        has_file=True,
        reminder_succeeded=False,
        file_descriptors=["weather-forecast.pdf"],
    )
    assert not _claims_an_unconfirmed_supported_deed(
        "Файл отчёта по зарплате отправлен.",
        has_file=True,
        reminder_succeeded=False,
        file_descriptors=["Отчёт по зарплате.pdf"],
    )
    assert _claims_an_unconfirmed_supported_deed(
        "Напоминание о зарплате поставлено.",
        has_file=False,
        reminder_succeeded=True,
        reminder_descriptors=["приём у врача 2026-08-14"],
    )
    assert not _claims_an_unconfirmed_supported_deed(
        "Напоминание о зарплате поставлено.",
        has_file=False,
        reminder_succeeded=True,
        reminder_descriptors=["зарплата 2026-08-14"],
    )


@pytest.mark.parametrize(
    "claim",
    [
        "Напоминание о зарплате поставлено на 2029-12-31.",
        "Напоминание о зарплате поставлено на 23:59.",
        "Напоминание о зарплате поставлено на завтра.",
    ],
)
def test_reminder_evidence_is_bound_to_the_claimed_time(claim: str) -> None:
    assert _claims_an_unconfirmed_supported_deed(
        claim,
        has_file=False,
        reminder_succeeded=True,
        reminder_descriptors=["зарплата 2026-08-14 09:00"],
    )


@pytest.mark.parametrize(
    ("claim", "files", "reminders"),
    [
        ("Файл отправлен, это отчёт по зарплате.", ["weather.pdf"], []),
        ("Напоминание поставлено, оно о зарплате.", [], ["приём у врача 2026-08-14"]),
        (
            "Файл отчёта по зарплате и клиентам отправлен.",
            ["Отчёт по зарплате.pdf", "Список клиентов.pdf"],
            [],
        ),
    ],
)
def test_appositive_and_cross_product_claims_need_one_matching_artifact(
    claim: str,
    files: list[str],
    reminders: list[str],
) -> None:
    assert _claims_an_unconfirmed_supported_deed(
        claim,
        has_file=bool(files),
        reminder_succeeded=bool(reminders),
        file_descriptors=files,
        reminder_descriptors=reminders,
    )


@pytest.mark.parametrize(
    "claim",
    [
        "Я поставила напоминание о зарплате.",
        "Я установила напоминание о зарплате.",
        "Я добавила напоминание о зарплате.",
    ],
)
def test_truthful_active_reminder_claims_match_their_evidence(claim: str) -> None:
    assert not _claims_an_unconfirmed_supported_deed(
        claim,
        has_file=False,
        reminder_succeeded=True,
        reminder_descriptors=["зарплата 2026-08-14 09:00"],
    )


def test_a_long_named_file_claim_cannot_escape_the_guard() -> None:
    subject = "синтетический " * 12 + "отчёт по зарплате"
    assert _claims_an_unconfirmed_supported_deed(
        f"Файл {subject} отправлен.",
        has_file=True,
        reminder_succeeded=False,
        file_descriptors=["weather.pdf"],
    )


@pytest.mark.parametrize(
    "false_report",
    [
        _FALSE_OUTSIDE_REPORT,
        _OBJECT_FIRST_FALSE_OUTSIDE_REPORT,
        _PASSIVE_TICKET_FALSE_OUTSIDE_REPORT,
        _PRINT_RESULT_FALSE_OUTSIDE_REPORT,
    ],
)
@pytest.mark.asyncio
async def test_an_unambiguous_outside_report_and_all_of_its_carriers_are_replaced(
    settings,
    storage,
    monkeypatch,
    false_report: str,
) -> None:
    """A text-only guard is insufficient: derived file, voice, and attribution must go too."""

    runtime = _runtime(settings, storage, monkeypatch)

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {
            "content": false_report,
            "tools_used": [],
            "knowledge_object_ids": ["ko_synthetic_claim"],
            "file_clips": [dict(_DOCUMENT)],
            "voice_clip": dict(_VOICE),
        }

    monkeypatch.setattr(runtime, "_generate_response", generate)

    reply = await runtime.chat(
        "alice",
        "Организуй физическую доставку к зданию.",
        actor=_actor(),
        enable_tools=False,
    )

    assert reply["message"] == _CANNOT_ACT_OUTSIDE
    assert reply["files"] == []
    assert reply["voice"] is None
    assert reply["context"]["attributed_knowledge_count"] == 0
    assert false_report not in repr(reply)
    stored = storage.get_message(str(reply["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    assert metadata["structural"]["output_guards"]["outside_deed_replaced"] is True


@pytest.mark.asyncio
async def test_a_real_structural_file_survives_replacement_of_model_speech(
    settings,
    storage,
    monkeypatch,
) -> None:
    runtime = _runtime(settings, storage, monkeypatch)
    structural_file = {**_DOCUMENT, "filename": "real-structural-synthetic.zip"}

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {
            "content": _FALSE_OUTSIDE_REPORT,
            "tools_used": ["collect_files"],
            "file_clips": [structural_file, dict(_DOCUMENT)],
            "_structural_file_count": 1,
        }

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Синтетическая составная просьба.",
        actor=_actor(),
        enable_tools=False,
    )

    assert reply["message"] == _CANNOT_ACT_OUTSIDE
    assert reply["files"] == [structural_file]


@pytest.mark.asyncio
async def test_k17_cannot_uncover_an_outside_report_after_k18_already_ran(
    settings,
    storage,
    monkeypatch,
) -> None:
    runtime = _runtime(settings, storage, monkeypatch)
    composite = "В моей личной базе знаний ответа нет. Но я уже заказала курьера."

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": composite, "tools_used": [], "file_clips": [dict(_DOCUMENT)]}

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Объясни синтетический общий принцип.",
        actor=_actor(),
        enable_tools=False,
    )

    assert reply["message"] == _CANNOT_ACT_OUTSIDE
    assert reply["files"] == []
    stored = storage.get_message(str(reply["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    assert metadata["structural"]["output_guards"] == {
        "outside_deed_replaced": True,
        "archive_status_replaced": False,
        "refusal_alternative_added": False,
    }


@pytest.mark.parametrize(
    "legitimate_answer",
    [
        "Сообщение отправлено в этот чат.",
        "Я перевела текст: The system is ready — Система готова.",
    ],
)
@pytest.mark.asyncio
async def test_a_real_current_chat_message_or_translation_is_not_replaced(
    settings,
    storage,
    monkeypatch,
    legitimate_answer: str,
) -> None:
    """Sending this reply and translating supplied text are native capabilities."""

    runtime = _runtime(settings, storage, monkeypatch)

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": legitimate_answer, "tools_used": []}

    monkeypatch.setattr(runtime, "_generate_response", generate)

    reply = await runtime.chat(
        "alice",
        "Ответь здесь на синтетический запрос.",
        actor=_actor(),
        enable_tools=False,
    )

    assert reply["message"] == legitimate_answer


class _FileToolLLM:
    """Calls a real runtime tool seam, then truthfully reports its attachment."""

    enabled = True
    model = "synthetic-file-tool-double"
    total_budget_sec = 5.0

    def __init__(self) -> None:
        self.rounds = 0

    async def chat(self, messages, *, tools=None, **kwargs):  # noqa: ANN001
        del messages, kwargs
        self.rounds += 1
        if self.rounds == 1:
            assert tools, "the integration test never entered the agentic tool road"
            return {
                "content": "",
                "tool_calls": [
                    {
                        "id": "call-synthetic-file",
                        "function": {
                            "name": "make_file",
                            "arguments": json.dumps(
                                {
                                    "kind": "docx",
                                    "title": "Synthetic note",
                                    "blocks": [{"kind": "text", "text": "Synthetic local body."}],
                                }
                            ),
                        },
                    }
                ],
                "_queue_wait_sec": 0.0,
            }
        return {
            "content": "Файл отправлен в этот чат.",
            "tool_calls": None,
            "_queue_wait_sec": 0.0,
        }


@pytest.mark.asyncio
async def test_a_successful_file_tool_keeps_its_truthful_report_and_attachment(
    settings,
    storage,
    monkeypatch,
) -> None:
    """The same verb is legal when this turn really produced the current-chat file."""

    storage.ensure_user("alice", preset_key="owner")
    llm = _FileToolLLM()
    kernel = ExecutionKernel(AuthorizationService(storage), settings=settings)
    kernel.bind_services(storage, None, None, None)
    tool_calls: list[str] = []
    execute = kernel.execute

    async def recorded_execute(name, arguments, *, actor=None):  # noqa: ANN001
        tool_calls.append(str(name))
        return await execute(name, arguments, actor=actor)

    monkeypatch.setattr(kernel, "execute", recorded_execute)
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=llm,
        kernel=kernel,
    )

    async def prepare(user_id, message, conversation_id, **kwargs):  # noqa: ANN001
        del message, kwargs
        return AgentContext(
            conversation_id=conversation_id,
            user_id=user_id,
            person_id=user_id,
            outward_verdict=("файл", None),
            answer_mode="general_conversation",
        )

    monkeypatch.setattr(runtime, "_prepare_context", prepare)

    reply = await runtime.chat(
        "alice",
        "Пришли файл в этот чат.",
        actor=_actor(),
        enable_tools=True,
    )

    assert tool_calls == ["make_file"], "the claimed file was not actually built"
    assert reply["message"] == "Файл отправлен в этот чат."
    assert len(reply["files"]) == 1
    assert reply["files"][0]["kind"] == "document"
    assert reply["files"][0]["filename"] == "Synthetic note.docx"
    assert reply["files"][0]["content_base64"]
    assert reply["tools_used"] == ["make_file"]


@pytest.mark.asyncio
async def test_a_hostile_repair_cannot_restore_the_replaced_outside_report(
    settings,
    storage,
    monkeypatch,
) -> None:
    """Replacement must remain structural or be guarded again after judge/repair."""

    runtime = _runtime(settings, storage, monkeypatch, verify_answers=True)

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {
            "content": "Синтетический черновик, который намеренно отправляется на проверку.",
            "tools_used": ["synthetic_local_lookup"],
            "tool_evidence": [{"tool": "synthetic_local_lookup", "output": "Synthetic local evidence."}],
        }

    verification_calls = 0
    repair_calls = 0

    async def verify(question, answer, context, *, tool_evidence=None):  # noqa: ANN001
        del question, answer, context, tool_evidence
        nonlocal verification_calls
        verification_calls += 1
        if verification_calls == 1:
            return {"status": "failed", "ok": False, "score": 0.0, "issues": ["synthetic"]}
        return {"status": "passed", "ok": True, "score": 1.0, "issues": []}

    async def repair(question, answer, context, verdict, *, tool_evidence=None):  # noqa: ANN001
        del question, answer, context, verdict, tool_evidence
        nonlocal repair_calls
        repair_calls += 1
        return _REPAIRED_FALSE_OUTSIDE_REPORT

    monkeypatch.setattr(runtime, "_generate_response", generate)
    monkeypatch.setattr(runtime, "_verify_response", verify)
    monkeypatch.setattr(runtime, "_repair_once", repair)

    reply = await runtime.chat(
        "alice",
        "Организуй физическую доставку к зданию.",
        actor=_actor(),
        enable_tools=False,
    )

    assert reply["message"] == _CANNOT_ACT_OUTSIDE
    assert _REPAIRED_FALSE_OUTSIDE_REPORT not in repr(reply)
    assert verification_calls == 1, "the system-owned replacement was sent back to the judge"
    assert repair_calls == 1, "the post-repair boundary was never exercised"


@pytest.mark.asyncio
async def test_late_file_construction_cannot_recreate_a_carrier_after_replacement(
    settings,
    storage,
    monkeypatch,
) -> None:
    """A file request must not package the refusal or the discarded false report."""

    runtime = _runtime(settings, storage, monkeypatch)

    async def prepare(user_id, message, conversation_id, **kwargs):  # noqa: ANN001
        del message, kwargs
        return AgentContext(
            conversation_id=conversation_id,
            user_id=user_id,
            person_id=user_id,
            outward_verdict=("файл", None),
            answer_mode="general_conversation",
        )

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": _FALSE_OUTSIDE_REPORT, "tools_used": []}

    late_builder_calls = 0

    async def build(request, answer, actor, *, evidence=None, context=None):  # noqa: ANN001
        del request, answer, actor, evidence, context
        nonlocal late_builder_calls
        late_builder_calls += 1
        return dict(_DOCUMENT)

    monkeypatch.setattr(runtime, "_prepare_context", prepare)
    monkeypatch.setattr(runtime, "_generate_response", generate)
    monkeypatch.setattr(runtime, "_file_for_a_request_that_wanted_one", build)

    reply = await runtime.chat(
        "alice",
        "Оформи в Word подтверждение физической доставки.",
        actor=_actor(),
        enable_tools=False,
    )

    assert reply["message"] == _CANNOT_ACT_OUTSIDE
    assert late_builder_calls == 0, "the discarded answer reached the late document builder"
    assert reply["files"] == []


@pytest.mark.asyncio
async def test_an_unsupported_file_completion_is_replaced_before_the_late_builder(
    settings,
    storage,
    monkeypatch,
) -> None:
    runtime = _runtime(settings, storage, monkeypatch)

    async def prepare(user_id, message, conversation_id, **kwargs):  # noqa: ANN001
        del message, kwargs
        return AgentContext(
            conversation_id=conversation_id,
            user_id=user_id,
            person_id=user_id,
            outward_verdict=("файл", None),
            answer_mode="general_conversation",
        )

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": "Файл готов и отправлен в этот чат.", "tools_used": []}

    late_builder_calls = 0

    async def build(request, answer, actor, *, evidence=None, context=None):  # noqa: ANN001
        del request, answer, actor, evidence, context
        nonlocal late_builder_calls
        late_builder_calls += 1
        raise AssertionError("an unsupported completion reached the late builder")

    monkeypatch.setattr(runtime, "_prepare_context", prepare)
    monkeypatch.setattr(runtime, "_generate_response", generate)
    monkeypatch.setattr(runtime, "_file_for_a_request_that_wanted_one", build)
    reply = await runtime.chat(
        "alice",
        "Подготовь синтетический документ.",
        actor=_actor(),
        enable_tools=False,
    )

    assert reply["message"] == _UNCONFIRMED_SUPPORTED_DEED
    assert reply["files"] == []
    assert late_builder_calls == 0
    stored = storage.get_message(str(reply["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    assert metadata["verified"] is False
    assert metadata["verification_status"] == "skipped"
    assert metadata["structural"]["output_guards"]["supported_deed_replaced"] is True


@pytest.mark.asyncio
async def test_archive_list_answer_is_not_replaced_as_a_file_handoff(
    settings,
    storage,
    monkeypatch,
) -> None:
    runtime = _runtime(settings, storage, monkeypatch)
    answer = "Вот список документов за сегодня: alpha.docx и beta.pdf."

    async def prepare(user_id, message, conversation_id, **kwargs):  # noqa: ANN001
        del message, kwargs
        return AgentContext(
            conversation_id=conversation_id,
            user_id=user_id,
            person_id=user_id,
            outward_verdict=("архив", None),
            answer_mode="personal_knowledge",
            asked_for_an_archive=True,
        )

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": answer, "tools_used": []}

    monkeypatch.setattr(runtime, "_prepare_context", prepare)
    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Покажи архивную подборку за сегодня.",
        actor=_actor(),
        enable_tools=False,
    )

    assert reply["message"] == answer
    stored = storage.get_message(str(reply["message_id"]), "alice")
    metadata = json.loads(str(stored["metadata_json"] or "{}"))
    assert metadata["structural"].get("output_guards", {}).get("supported_deed_replaced") is None


@pytest.mark.asyncio
async def test_a_false_model_deed_does_not_erase_a_true_structural_deed(
    settings,
    storage,
    monkeypatch,
) -> None:
    runtime = _runtime(settings, storage, monkeypatch)
    structural = "Напоминание поставлено: «приём у врача», срок — завтра."

    async def prepare(user_id, message, conversation_id, **kwargs):  # noqa: ANN001
        del message, kwargs
        return AgentContext(
            conversation_id=conversation_id,
            user_id=user_id,
            person_id=user_id,
            answer_mode="general_conversation",
            structural_answer=structural,
            successful_reminders=[{"what": "приём у врача", "when": "завтра"}],
        )

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": "Файл отчёта по зарплате отправлен.", "tools_used": []}

    monkeypatch.setattr(runtime, "_prepare_context", prepare)
    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Синтетическая составная просьба.",
        actor=_actor(),
        enable_tools=False,
    )

    assert reply["message"] == f"{structural}\n\n{_UNCONFIRMED_SUPPORTED_DEED}"
    assert reply["files"] == []


@pytest.mark.asyncio
async def test_a_truthful_requested_reminder_time_matches_the_persisted_deed(
    settings,
    storage,
    monkeypatch,
) -> None:
    runtime = _runtime(settings, storage, monkeypatch)
    structural = (
        "Напоминание поставлено: «отчёт», срок — завтра в 15:00 "
        "(2026-08-09 15:00). Доставка в чат запланирована."
    )

    async def prepare(user_id, message, conversation_id, **kwargs):  # noqa: ANN001
        del message, kwargs
        return AgentContext(
            conversation_id=conversation_id,
            user_id=user_id,
            person_id=user_id,
            answer_mode="general_conversation",
            structural_answer=structural,
            successful_reminders=[
                {
                    "what": "отчёт",
                    "requested_when": "завтра в 15:00",
                    "when": "2026-08-09 15:00",
                }
            ],
        )

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {"content": "Напоминание об отчёте поставлено на завтра в 15:00.", "tools_used": []}

    monkeypatch.setattr(runtime, "_prepare_context", prepare)
    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Синтетическая составная просьба.",
        actor=_actor(),
        enable_tools=False,
    )

    assert _UNCONFIRMED_SUPPORTED_DEED not in reply["message"]
    assert "завтра в 15:00" in reply["message"]


@pytest.mark.asyncio
async def test_an_unrelated_file_does_not_validate_a_named_file_claim(
    settings,
    storage,
    monkeypatch,
) -> None:
    runtime = _runtime(settings, storage, monkeypatch)
    unrelated = {**_DOCUMENT, "filename": "weather-forecast.pdf"}

    async def generate(context, message, attachments):  # noqa: ANN001
        del context, message, attachments
        return {
            "content": "Файл отчёта по зарплате отправлен.",
            "tools_used": ["make_file"],
            "file_clips": [unrelated],
        }

    monkeypatch.setattr(runtime, "_generate_response", generate)
    reply = await runtime.chat(
        "alice",
        "Синтетическая составная просьба.",
        actor=_actor(),
        enable_tools=False,
    )

    assert reply["message"] == _UNCONFIRMED_SUPPORTED_DEED
    assert reply["files"] == [unrelated]


class _CarrierOnlyToolLLM:
    enabled = True
    model = "synthetic-carrier-only-double"
    total_budget_sec = 5.0

    def __init__(self, tool_name: str, arguments: dict) -> None:
        self.tool_name = tool_name
        self.arguments = arguments
        self.rounds = 0

    async def chat(self, messages, *, tools=None, **kwargs):  # noqa: ANN001
        del messages, kwargs
        self.rounds += 1
        if self.rounds == 1:
            assert tools
            return {
                "content": "",
                "tool_calls": [
                    {
                        "id": "call-hostile-carrier",
                        "function": {
                            "name": self.tool_name,
                            "arguments": json.dumps(self.arguments, ensure_ascii=False),
                        },
                    }
                ],
                "_queue_wait_sec": 0.0,
            }
        return {
            "content": "Безопасный синтетический итог без отчёта о внешнем действии.",
            "tool_calls": None,
            "_queue_wait_sec": 0.0,
        }


@pytest.mark.parametrize(
    ("tool_name", "arguments"),
    [
        (
            "make_file",
            {
                "kind": "docx",
                "title": "Синтетический итог",
                "blocks": [{"kind": "text", "text": _FALSE_OUTSIDE_REPORT}],
            },
        ),
        (
            "make_file",
            {
                "kind": "docx",
                "title": "В моей личной базе знаний",
                "subtitle": "ответа нет.",
                "blocks": [{"kind": "text", "text": "Безопасный синтетический раздел."}],
            },
        ),
        (
            "make_file",
            {
                "kind": "xlsx",
                "title": "Синтетическая таблица",
                "blocks": [
                    {
                        "kind": "table",
                        "rows": [{_FALSE_OUTSIDE_REPORT: "безопасное значение"}],
                    }
                ],
            },
        ),
        (
            "make_file",
            {
                "kind": "docx",
                "title": "Синтетический отказ",
                "blocks": [{"kind": "text", "text": _BARE_REFUSAL}],
            },
        ),
        ("speak", {"text": _FALSE_OUTSIDE_REPORT}),
        ("speak", {"text": _BARE_REFUSAL}),
    ],
)
@pytest.mark.asyncio
async def test_a_false_report_only_in_a_model_carrier_is_rejected_before_render(
    settings,
    storage,
    monkeypatch,
    tool_name: str,
    arguments: dict,
) -> None:
    storage.ensure_user("alice", preset_key="owner")
    llm = _CarrierOnlyToolLLM(tool_name, arguments)
    kernel = ExecutionKernel(AuthorizationService(storage), settings=settings)
    kernel.bind_services(storage, None, None, None)
    executed: list[str] = []

    async def execute(name, arguments, *, actor=None):  # noqa: ANN001
        del arguments, actor
        executed.append(name)
        raise AssertionError("a rejected carrier reached the binary renderer")

    monkeypatch.setattr(kernel, "execute", execute)
    runtime = AgentRuntime(replace(settings, verify_answers=False), storage, llm=llm, kernel=kernel)
    monkeypatch.setattr(runtime, "_prepare_context", _simple_context)

    reply = await runtime.chat(
        "alice",
        "Объясни синтетический общий принцип.",
        actor=_actor(),
        enable_tools=True,
    )

    assert executed == []
    assert reply["message"] == "Безопасный синтетический итог без отчёта о внешнем действии."
    assert reply["files"] == []
    assert reply["voice"] is None


@pytest.mark.asyncio
async def test_the_late_document_model_cannot_put_an_outside_report_in_the_file(
    settings,
    storage,
) -> None:
    class _LateLLM:
        enabled = True

        async def chat(self, messages, *, tools=None, **kwargs):  # noqa: ANN001
            del messages, tools, kwargs
            return {"content": _FALSE_OUTSIDE_REPORT}

    class _NeverRenderKernel:
        async def execute(self, name, arguments, *, actor=None):  # noqa: ANN001
            del name, arguments, actor
            raise AssertionError("the hostile late document reached make_file")

    runtime = AgentRuntime.__new__(AgentRuntime)
    runtime.llm = _LateLLM()
    runtime.kernel = _NeverRenderKernel()
    context = AgentContext(
        conversation_id="conv-synthetic",
        user_id="alice",
        person_id="alice",
        outward_verdict=("файл", None),
        answer_mode="general_conversation",
    )

    made = await runtime._file_for_a_request_that_wanted_one(  # noqa: SLF001
        "Подготовь синтетический документ.",
        "Краткий ответ.",
        _actor(),
        evidence=[{"tool": "synthetic_local_lookup", "output": "Synthetic grounds."}],
        context=context,
    )

    assert made is None
