"""Deidentified regressions for owner file turns followed by public web turns.

Every file, model response, and web result in this module is synthetic.  The
tests exercise the full ``AgentRuntime.chat`` boundary without a live model,
private corpus, network provider, or production account.
"""

from __future__ import annotations

import hashlib
import io
import json
from dataclasses import dataclass, replace
from datetime import UTC, datetime, timedelta
from types import SimpleNamespace
from typing import Any

import pytest
from docx import Document

from friday.agent_runtime import (
    AgentContext,
    AgentRuntime,
    _filename_clue_request,
    _self_contained_explicit_public_web_query,
    _self_contained_public_market_query,
    _self_contained_public_product_spec_query,
    _simple_public_news_freshness,
    asks_for_the_web,
    file_turn_authority,
)
from friday.agent_runtime._office_attachments import (
    OFFICE_EXACT_UNAVAILABLE_MESSAGE,
    OFFICE_STRUCTURE_KEY,
    validate_runtime_office_index,
)
from friday.documents import DocumentExtractor
from friday.execution_kernel import ToolResult
from friday.office_attestation import (
    OFFICE_STRUCTURE_ATTESTATION_METADATA_KEY,
    sign_office_structure_index,
)
from friday.permissions import ActorContext
from friday.storage.models import RawObject, new_id

OWNER = "owner_web_file_regression"
NEWS_REQUEST = "Покажешь свежие новости за прошедшие сутки?"
PUBLIC_URL = "https://public.synthetic.example.com/news"
PUBLIC_FACT = "Синтетическая публичная новость подтверждена источником."
PRIVATE_PREFIX = "PRIVATE-SYNTHETIC-DOC-CANARY"
PRODUCT_URL = "https://public.synthetic.example.com/qnap-tvs-675"
PRODUCT_FACT = "QNAP TVS-675 uses the synthetic PUBLIC-CPU-675 processor."
PRIVATE_HISTORY_CANARY = "PRIVATE-PRODUCT-SPEC-HISTORY-CANARY"
GENERIC_WEB_URL = "https://public.synthetic.example.com/nextcloud"
GENERIC_WEB_FACT = "The synthetic official Nextcloud project page is public."


def test_foreign_news_sites_are_web_authority_without_becoming_a_file_summary() -> None:
    request = "Сделай сводку по новостям СВО на зарубежных сайтах"
    live_calendar_request = "Привет! дай сводку по новостям СВО за 24 число"

    assert asks_for_the_web(request) is True
    authority = file_turn_authority(request)
    assert authority.proved("web")
    assert not authority.proved("local_read")

    world_question = "Какие мировые новости вышли вчера?"
    assert asks_for_the_web(world_question) is True
    assert file_turn_authority(world_question).proved("web")

    for data_only in (
        "В зарубежных СМИ сегодня обсуждают новости.",
        "«Сделай сводку по новостям на зарубежных сайтах»",
    ):
        assert asks_for_the_web(data_only) is False
        assert not file_turn_authority(data_only).proved("web")

    local_request = "Сделай сводку новостей из этого файла о зарубежных СМИ"
    assert asks_for_the_web(local_request) is False
    local_authority = file_turn_authority(local_request)
    assert local_authority.proved("local_read")
    assert not local_authority.proved("web")
    assert AgentRuntime.web_query_from(live_calendar_request) == "сводку по новостям СВО за 24 число"

    for public_request in (
        live_calendar_request,
        "Пожалуйста, подготовь новости экономики за 24.08.2026",
        "Можно сводку новостей за 24 число?",
        "Сводку новостей за 24 число, пожалуйста",
        "Что было в новостях 24 числа?",
        "Можешь дать сводку новостей за 24 число?",
        "Подведи итоги новостей за 24 число",
        "Расскажи последние новости науки",
        "Свежие новости за прошедшие сутки покажешь?",
        "Give me news for August 24",
        "Please give me news for August 24",
        "Tell me the news for August 24",
        "Привет, Пятница! Дай сводку новостей за 24 число",
        "Подскажи новости за 24 число",
        "Скажи новости за 24 число",
        "Что по новостям за 24 число?",
        "Каковы новости за 24 число?",
        "Could you please give me the news for August 24?",
        "Can I get the news for August 24?",
        "Could I get the latest news about OpenAI?",
        "I’d like the latest news about OpenAI",
        "I want the latest news about OpenAI",
        "Do you have the latest news about OpenAI?",
        "Any news on OpenAI today?",
        "What's today's news about science?",
        "What's in the news for August 24?",
        "Give me news for August 24, 2026",
        "Дай новости экономики, науки и спорта за 24 число",
        "Give me news on science, technology, and business for August 24",
        "Дай новости за 2026-08-24",
        "Today, tell me the news about science",
        "For today, give me the news about science",
        "Сегодня, покажи новости науки",
        "За сегодня, покажи новости науки",
        "Расскажи последние новости OpenAI",
        "Не мог бы ты рассказать последние новости OpenAI?",
        "Не могла бы ты показать последние новости OpenAI?",
        "Мне нужны последние новости OpenAI",
        "Хочу узнать последние новости OpenAI",
        "Хотел бы узнать последние новости OpenAI",
        "Есть свежие новости по OpenAI?",
        "Есть ли последние новости OpenAI?",
        "Новости OpenAI сегодня",
        "Каковы главные новости NVIDIA?",
        "OpenAI news today",
        "NVIDIA news today",
        "AI news today",
        "OpenAI latest news today",
        "Главные новости OpenAI сегодня",
        "Дай последние новости про PDF стандарты",
        "Дай последние новости про стандарты PDF",
        "Дай новости PDF сегодня",
        "Give me OpenAI news from the past day",
        "Give me OpenAI news for two weeks",
        "Расскажи последние новости искусственного интеллекта",
        "Расскажи последние новости квантовых вычислений",
        "Show latest news about OpenAI",
        "Tell me the latest news about artificial intelligence",
        "Пятница, дай сводку новостей за 24 число",
        "Каковы главные новости за 24 число?",
        "Дай новости России, Украины и мира за 24 августа",
        "Give me news about science and AI today",
        "Расскажи последние новости NVIDIA",
        "Расскажи последние новости криптовалют",
        "Расскажи последние новости Израиля",
        "Give me latest news about Nvidia",
        "Give me latest Nvidia news",
        "Give me latest quantum computing news",
        "Give me latest news about private equity",
        "Give me latest news about confidential computing",
        "Give me latest news about internal combustion engines",
        "Give me latest news about password managers",
        "Give me latest news about token prices",
        "Give me latest news about code generation",
        "Give me latest news about Project Gutenberg",
        "Find current news about document management",
        "Find latest news about password breach",
        "Find latest news about code review",
        "Найди последние новости о код-ревью",
        "Каковы сегодняшние новости NVIDIA?",
        "За 24 августа дай новости NVIDIA",
        "For August 24 give me news about NVIDIA",
        "Найди новости про файловые системы сегодня",
        "Покажи последние новости о документах Microsoft",
        "Найди свежие новости архива Интернета",
        "Какие сегодня новости науки?",
        "Дай новости сегодня про Газпром",
        "Покажи новости за сегодня по экономике",
        "Give me today's news about science",
        "Tell me today's news",
        "Give me latest news about AI/ML",
        "Give me latest news about S&P 500",
        "Give me latest news about C++",
        "Give me latest news about .NET",
        "Дай новости сегодня про .NET",
        "Give me latest news about Project 2025",
        "Дай последние новости про Проект 2025",
        "Give me latest news about iOS 26.1",
        "Give me latest news about Python 3.13",
        "Give me latest news about Windows 11.2",
        "Find current news about Ubuntu 24.04",
        "Give me news about Ubuntu 24.04 today",
        "Give me latest news about GPT-4.1",
        "Give me latest news about iOS 18.6.2",
        "Give me latest news about Python 3.13.1",
        "Give me news about Linux 6.12.1 today",
        "Latest news X",
        "Latest R news",
        "Найди последние новости о код ревью",
        "Найди последние новости о код-ревью",
        "Give me latest news about Secret Santa",
        "Give me latest news about Secret Garden",
        "Give me latest news about Secret Service",
        "Give me latest news about secret ballot",
        "Give me latest news about secret diplomacy",
        "Give me latest news about code coverage",
        "Give me latest news about code ownership",
        "Give me latest news about code repository",
        "Give me latest news about token unlock",
        "Give me latest news about token airdrop",
        "Give me latest news about password hygiene",
        "Give me latest news about password recovery",
        "Give me latest news about login outage",
        "Give me latest news about login fraud",
        "Give me latest news about credential exposure",
        "Give me latest news about API key compromise",
        "Give me latest news about personal finance",
        "Give me latest news about read receipts",
        "Give me latest news about email marketing",
        "Give me latest news about export controls",
        "Give me latest news about CRM updates",
        "Give me latest news about software updates",
        "Give me latest news about Windows updates",
        "Give me latest news about security updates",
        "Give me latest news about update management",
        "Give me latest news about Wi-Fi 7",
        "Give me latest news about CVE-2026-1234",
        "Give me latest news about Russia–Ukraine war",
        "Give me latest news about Russia—Ukraine war",
        "Give me latest news about Russia‑Ukraine war",
        "Give me latest news about COVID‑19",
        "Give me latest news about AI‑generated content",
        "Give me latest news about U.S. politics",
        "Give me latest news about U.K. economy",
        "Give me latest news about EU–U.S. relations",
        "Give me latest news about OpenAI’s models",
        "Give me latest news about O’Reilly Media",
        "Give me latest news about Biden–Xi talks",
        "Give me latest news about Beyoncé",
        "Give me latest news about Pokémon",
        "Give me news about Black Friday today",
        "Give me news about Wednesday season 2 today",
        "Give me news about March Madness today",
        "Give me news about Monday.com today",
        "Give me news about August smart locks today",
        "Дай новости про Чёрную пятницу сегодня",
        "Дай новости про сериал Среда сегодня",
        "Дай последние новости про управление тегами",
        "Дай последние новости про изменения законодательства",
        "Дай последние новости про изменения климата",
        "Дай последние новости про добавление тегов",
        "Расскажи последние новости генерации кода",
        "Расскажи последние новости токенов криптовалют",
        "Расскажи последние новости проекта Геном человека",
        "Дай новости нефти и газа за 24 число",
        "Give me news about Nvidia and AMD today",
    ):
        assert asks_for_the_web(public_request) is True
        assert file_turn_authority(public_request).actions == frozenset({"web"})

    for titled_topic in (
        "Give me news about Black Friday today",
        "Give me news about Wednesday season 2 today",
        "Give me news about March Madness today",
        "Give me news about Monday.com today",
        "Give me news about August smart locks today",
        "Дай новости про Чёрную пятницу сегодня",
        "Дай новости про сериал Среда сегодня",
    ):
        assert _simple_public_news_freshness(titled_topic) == "day"

    assert _simple_public_news_freshness("Give me OpenAI news from the past day") == "day"
    assert _simple_public_news_freshness("Give me OpenAI news for two weeks") is None

    for non_request in (
        "В новостях за 24 число обсуждали экономику.",
        "В заметке написано: «дай сводку новостей за 24 число»",
        "Дай сводку новостей из этого файла за 24 число",
        "Дай сводку новостей по нашей CRM за 24 число",
        "Покажи новости из моей заметки за вчера",
        "Расскажи актуальные новости о нём",
        "Привет! дай сводку новостей за 24 число «PRIVATE-CARRIER»",
        "Привет! дай сводку новостей за 24 число\nPRIVATE-CARRIER",
        "Дай сводку новостей за 24 число. Проект АЛЬФА имеет код QZ-417",
        "Дай сводку новостей за 24 число; PRIVATE-CARRIER",
        "Дай сводку новостей за 24 число! PRIVATE-CARRIER",
        "Дай сводку новостей за 24 число? PRIVATE-CARRIER",
        "Дай сводку новостей за 24 число, PRIVATE-CARRIER",
        "Дай сводку новостей за 24 число: PRIVATE-CARRIER",
        "Дай сводку новостей за 24 число.PRIVATE-CARRIER",
        "Дай сводку новостей за 24 число． PRIVATE-CARRIER",
        "Дай сводку новостей за 24 число.\u200bPRIVATE-CARRIER",
        "Дай сводку новостей за 24 число\u2028PRIVATE-CARRIER",
        "Дай сводку [новостей за 24 число](https://private.invalid/CARRIER)",
        "Дай сводку новостей за 24 число https://private.invalid/CARRIER",
        "Дай сводку новостей за 24 число Проект АЛЬФА имеет код QZ417 сегодня",
        "Дай сводку новостей за 24 число Проект АЛЬФА имеет код QZ417 на зарубежных сайтах",
        "Give me news for August 24 Project ALPHA has code QZ417 today",
        "Give me news for August 24 Project ALPHA from foreign media",
        "Дай сводку новостей за 24 число Проект АЛЬФА 3 дня",
        "Дай сводку новостей за 24 число Проект АЛЬФА месяц",
        "Дай сводку новостей за 24 число Проект АЛЬФА за 3 дня",
        "Расскажи последние новости науки Проект ALPHA код QZ417",
        "Расскажи последние новости науки внутренний бюджет 42 млн",
        "Дай новости экономики, Проекта АЛЬФА за 24 число",
        "Give me news on science, Project ALPHA for August 24",
        "Give me news for August 24, 2026 Project ALPHA",
        "Дай новости за 2026-08-24 Проект АЛЬФА",
        "Show latest news about OpenAI Project ALPHA",
        "Дай новости науки и внутренний бюджет 42 млн за 24 число",
        "Give me news about science and internal budget 42 today",
        "Show latest news airgapcredentials",
        "Give me latest Project ALPHA news",
        "Расскажи последние новости project-alpha",
        "Give me latest projectalpha news",
        "Give me latest p-r-o-j-e-c-t Alpha news",
        "Give me latest рroject Alpha news",
        "Give me latest news ρroject Alpha",
        "Give me latest news ρ-r-o-j-e-c-t Alpha",
        "Give me latest news ѕecret Alpha",
        "Give me latest news credentіals Alpha",
        "Give me latest news prօject Alpha",
        "Give me latest news cօde Alpha",
        "Give me latest news prοject Alpha",
        "Give me latest news cοde Alpha",
        "Give me latest news about ṕroject ALPHA",
        "Give me latest news about ćode QZ417",
        "Give me latest news about śecret ALPHA",
        "Give me latest news about ínternal budget",
        "Give me latest news about prøject ALPHA",
        "Give me latest news about cøde QZ417",
        "Give me latest news about credentıals Alpha",
        "Give me latest news about secreŧ ALPHA",
        "Give me latest news about cønfidential QZ417",
        "Give me latest news about tøken QZ417",
        "Дай последние новости көд QZ417",
        "Дай последние новости сєкрет ALPHA",
        "Дай последние новости төкен QZ417",
        "Дай последние новости пројкт АЛЬФА",
        "Дай последние новости парөль hunter",
        "Give me latest news about pr0ject ALPHA",
        "Give me latest news about proj3ct ALPHA",
        "Give me latest news about c0de QZ417",
        "Give me latest news about s3cret ALPHA",
        "Give me latest news about 1nternal budget",
        "Give me latest news about credent1als vault",
        "Give me latest news about pr1vate client Acme",
        "Give me latest news about t0ken QZ417",
        "Give me latest news about passw0rd hunter",
        "Дай последние новости про3кт АЛЬФА",
        "Дай последние новости к0д QZ417",
        "Дай последние новости с3крет ALPHA",
        "Give me latest news about projec ALPHA",
        "Give me latest news about projject ALPHA",
        "Give me latest news about project0 ALPHA",
        "Give me latest news about cod QZ417",
        "Give me latest news about codde QZ417",
        "Give me latest news about secre ALPHA",
        "Give me latest news about secrett ALPHA",
        "Дай последние новости проек ALPHA",
        "Дай последние новости коод QZ417",
        "Дай последние новости про проєкм QZ417",
        "Дай последние новости привтный QZ417",
        "Дай последние новости внутреенний QZ417",
        "Дай последние новости внутрений QZ417",
        "Дай последние новости проєкт АЛЬФА",
        "Дай последние новости про внутрішній бюджет",
        "Дай последние новости про конфіденційний QZ417",
        "Дай последние новости про особистий файл",
        "Дай последние новости про мій файл",
        "Give me latest news my-file",
        "Give me latest news our-crm",
        "Give me latest news my-notes",
        "Give me latest news our-archive",
        "Give me latest news c-r-e-d-e-n-t-i-a-l-s",
        "Give me latest news c-o-d-e-QZ417",
        "Give me latest news i-n-t-e-r-n-a-l-budget",
        "Give me latest news about our client Acme",
        "Give me latest news about my client Ivanov",
        "Give me latest news about our customer Acme",
        "Give me latest news about my boss Alice",
        "Give me latest news about our contract QZ417",
        "Give me latest news about my deal Alpha",
        "Give me latest news about our meeting agenda",
        "Give me latest news about my salary details",
        "Give me latest news about our invoice 123",
        "Give me latest news about our support ticket",
        "Give me latest news about mine Alpha",
        "Give me latest news about ours Alpha",
        "Give me latest news about yours Alpha",
        "Give me latest news about private QZ417",
        "Give me latest news about confidential QZ417",
        "Give me latest news about internal QZ417",
        "Give me latest news about private ALPHA",
        "Give me latest news about confidential ALPHA",
        "Give me latest news about internal ALPHA",
        "Расскажи последние новости про приватный QZ417",
        "Расскажи последние новости про конфиденциальный QZ417",
        "Расскажи последние новости про внутренний QZ417",
        "Give me latest news about privateQZ417",
        "Give me latest news about confidentialQZ417",
        "Give me latest news about internalQZ417",
        "Give me latest news about codeALPHA",
        "Give me latest news about tokenALPHA",
        "Give me latest news about passwordALPHA",
        "Give me latest news about credentialALPHA",
        "Give me latest news about secretALPHA",
        "Расскажи последние новости про кодАЛЬФА",
        "Расскажи последние новости про токенАЛЬФА",
        "Расскажи последние новости про парольАЛЬФА",
        "Расскажи последние новости про секретАЛЬФА",
        "Give me latest news from the private channel",
        "Дай последние новости из закрытого канала",
        "Give me latest news about pro ject Alpha",
        "Give me latest news about co de QZ417",
        "Give me latest news about creden tials vault",
        "Give me latest news about password is hunter",
        "Дай последние новости про токен равен abcdef",
        "Give me latest news about password is/hunter",
        "Give me latest news about password is.hunter",
        "Give me latest news about password is,hunter",
        "Give me latest news about password is&hunter",
        "Give me latest news about password/is hunter",
        "Give me latest news about api key is/alpha",
        "Give me latest news about token equals/abcdef",
        "Дай последние новости про пароль равен/hunter",
        "Дай последние новости про токен равен.abcdef",
        "Give me latest news from private/channel",
        "Give me latest news about private.channel",
        "Give me latest news about private,channel",
        "Give me latest news from internal/archive",
        "Give me latest news about confidential&notes",
        "Give me latest news from o/ur notes",
        "Give me latest news from m.y notes",
        "Give me latest news about project/ALPHA",
        "Give me latest news about private/QZ417",
        "Give me latest news about code/QZ417",
        "Дай последние новости про проект/АЛЬФА",
        "Дай последние новости про код/QZ417",
        "Дай последние новости из закрытого/канала",
        "Дай последние новости из приватного.архива",
        "Дай последние новости из внутренней,CRM",
        "Give me today's news about science and remind me tomorrow",
        "Give me today's news about AI/ML and run code",
        "Give me today's news about AI/ML and speak aloud",
        "Give me today's news about AI and delete the file",
        "Give me latest news about Nvidia and read the attached file",
        "Give me latest news about Nvidia and open the file",
        "Give me latest news about Nvidia and summarize the document",
        "Give me latest news about Nvidia and compare the files",
        "Give me latest news about Nvidia and search the archive",
        "Give me latest news about Nvidia and find person Alice",
        "Give me latest news about Nvidia and tag the note",
        "Дай последние новости NVIDIA и прочитай приложенный файл",
        "Дай последние новости NVIDIA и суммируй документ",
        "Дай последние новости NVIDIA и сравни файлы",
        "Дай последние новости NVIDIA и найди человека Алису",
        "Дай последние новости NVIDIA и обнови CRM",
        "Дай последние новости NVIDIA и напиши Алисе",
        "Give me today's news about science/remind me tomorrow",
        "Give me today's news about science&remind me tomorrow",
        "Give me today's news about AI/ML/remind me tomorrow",
        "Give me today's news about S&P500&run code",
        "Give me today's news about .NET, remind me tomorrow",
        "Дай последние новости S&P500&запусти код",
        "Give me today's news about science remind me tomorrow",
        "Give me today's news about science run code",
        "Give me today's news about science speak aloud",
        "Give me today's news about science-remind me tomorrow",
        "Give me today's news about foo.remind me tomorrow",
        "Give me latest news about private client Acme",
        "Give me latest news about NVIDIA and private client Acme",
        "Give me latest news about closed channel Alpha",
        "Give me latest news about NVIDIA and order pizza",
        "Give me latest news about NVIDIA and pay the bill",
        "Give me latest news about NVIDIA and restart the server",
        "Give me latest news about NVIDIA and schedule a meeting",
        "Give me latest news about NVIDIA and book a meeting",
        "Give me latest news about NVIDIA and create a reminder",
        "Give me latest news about NVIDIA and open source",
        "Give me latest news about Google and contact tracing",
        "Give me latest news about X and message sending",
        "Give me latest news about X and search engines",
        "Give me latest news about NVIDIA and read the message",
        "Give me latest news about NVIDIA and summarize the message",
        "Give me latest news about NVIDIA and update the message",
        "Give me latest news about NVIDIA and delete the message",
        "Give me latest news about NVIDIA and open the PDF",
        "Give me latest news about NVIDIA and analyze the PDF",
        "Give me latest news about NVIDIA and review the PDF",
        "Give me latest news about NVIDIA and compare the PDF",
        "Give me latest news about NVIDIA and save the PDF",
        "Give me latest news about NVIDIA and send the attached file",
        "Give me latest news about NVIDIA and email the PDF",
        "Give me latest news about NVIDIA and read PDF",
        "Give me latest news about NVIDIA and delete message",
        "Give me latest news about NVIDIA and summarize PDF",
        "Give me latest news about NVIDIA and analyze report",
        "Give me latest news about NVIDIA and remind tomorrow",
        "Give me latest news about NVIDIA and speak response",
        "Give me latest news about NVIDIA and upload image",
        "Give me latest news about NVIDIA and export CSV",
        "Give me latest news about NVIDIA and email report",
        "Give me latest news about NVIDIA and call Alice",
        "Give me latest news about NVIDIA and set an alarm",
        "Give me latest news about NVIDIA and shut down the server",
        "Give me latest news about NVIDIA and deploy the app",
        "Give me latest news about NVIDIA and run ls",
        "Дай последние новости NVIDIA и выключи свет",
        "Give me latest news about NVIDIA and launch the browser",
        "Give me latest news about NVIDIA and buy stock",
        "Give me latest news about NVIDIA and place an order",
        "Give me latest news about NVIDIA and change password",
        "Дай последние новости NVIDIA и купи акции",
        "Дай последние новости NVIDIA и смени пароль",
        "Give me latest news about NVIDIA from Alice's notes",
        "Give me latest news about NVIDIA based on Alice's document",
        "Give me latest news about NVIDIA from the client notes",
        "Дай последние новости NVIDIA из заметок Алисы",
        "Give me latest news about NVIDIA PIN is 1234",
        "Give me latest news about NVIDIA SSH key is abcdef",
        "Give me latest news about NVIDIA cookie is abcdef",
        "Дай последние новости NVIDIA пин равен 1234",
        "Give me latest news about secret ALPHA",
        "Give me latest news about API key QZ417",
        "Give me latest news about SSH key QZ417",
        "Give me latest news about login USER123",
        "Give me latest news about PIN 1234",
        "Give me latest news about cookie ABC123",
        "Give me latest news about session ABC123",
        "Дай последние новости API ключ AB12",
        "Дай последние новости SSH ключ QZ417",
        "Дай последние новости логин USER123",
        "Дай последние новости ПИН 1234",
        "Дай последние новости куки ABC123",
        "Дай последние новости сессию ABC123",
        "Give me today's news about NVIDIA and email alice",
        "Give me today's news about NVIDIA and call bob",
        "Give me today's news about NVIDIA and turn the lights off",
        "Give me today's news about NVIDIA and reboot the router",
        "Дай последние новости NVIDIA и перезагрузи роутер",
        "Give me latest news from employee Alice",
        "Дай последние новости из канала Ивана",
        "Give me latest news from the attached archive",
        "Give me latest news based on the attached archive",
        "Дай последние новости по списку клиента",
        "Дай последние новости из списка клиента",
        "Дай последние новости про список клиента",
        "Give me today's news about NVIDIA and text alice",
        "Give me today's news about NVIDIA and run pytest",
        "Give me today's news about NVIDIA and wipe the disk",
        "Give me today's news about NVIDIA and transfer funds",
        "Give me today's news about NVIDIA and notify alice",
        "Give me today's news about NVIDIA and dial bob",
        "Give me today's news about NVIDIA and dim lights",
        "Give me today's news about NVIDIA and power off",
        "Give me today's news about NVIDIA and format disk",
        "Give me today's news about NVIDIA and encrypt disk",
        "Give me today's news about NVIDIA and cancel reservation",
        "Give me today's news about NVIDIA and take photo",
        "Дай последние новости NVIDIA и поставь таймер",
        "Дай последние новости NVIDIA и зашифруй диск",
        "Give me latest news about OpenAI then read Alice",
        "Give me latest news about OpenAI also delete Alice",
        "Give me latest news about OpenAI then upload Alice",
        "Дай последние новости OpenAI затем скачай документ",
        "Дай последние новости OpenAI затем запиши архив",
        "Дай последние новости OpenAI потом скачай код",
        "Give me latest news about NVIDIA and create archive",
        "Give me latest news about NVIDIA and send archive",
        "Give me latest news about NVIDIA and list contacts",
        "Give me latest news about NVIDIA and export contacts",
        "Дай последние новости NVIDIA и прочитай сообщение",
        "Дай последние новости NVIDIA и перескажи сообщение",
        "Дай последние новости NVIDIA и найди сообщение",
        "Дай последние новости NVIDIA и поищи сообщение",
        "Дай последние новости NVIDIA и пометь сообщение",
        "Дай последние новости NVIDIA и обнови сообщение",
        "Дай последние новости NVIDIA и удали сообщение",
        "Дай последние новости NVIDIA и отправь приложенный файл",
        "Дай последние новости NVIDIA и отправь голосовой ответ",
        "Дай последние новости NVIDIA и закажи пиццу",
        "Дай последние новости NVIDIA и оплати счёт",
        "Дай последние новости NVIDIA и перезапусти сервер",
        "Дай сводку новостей за 24 число № QZ417",
        "Дай сводку новостей за 24 число ™ QZ417",
        "Дай новости из нашего проекта за 24 число",
        "Сводку новостей за 24 число подготовили заранее.",
        "Сегодня обсуждали документы. Дай сводку новостей за 24 число",
    ):
        assert asks_for_the_web(non_request) is False
        assert not file_turn_authority(non_request).proved("web")

    for web_scope in (
        "В интернете найди, я кидал уже",
        "В сети посмотри, я уже присылал",
        "В онлайн найди, я загружал раньше",
    ):
        assert _filename_clue_request(web_scope) is None


def test_overlong_turn_cannot_mint_effect_authority() -> None:
    request = "Дай последние новости науки и удали файл " + ("x" * 5000)

    assert file_turn_authority(request).actions == frozenset()
    assert asks_for_the_web(request) is False


def test_public_product_spec_query_is_closed_and_non_deictic() -> None:
    assert (
        _self_contained_public_product_spec_query("в qnap TVS-675 какой процессор?")
        == "qnap TVS-675 процессор"
    )
    assert _self_contained_public_product_spec_query("Какой CPU у QNAP TVS-675?") == "qnap TVS-675 cpu"
    for unsafe_or_contextual in (
        "а у него какой процессор?",
        "в этом файле у qnap TVS-675 какой процессор?",
        "в internal X-1 какой процессор?",
        "«в qnap TVS-675 какой процессор?»",
        "в qnap TVS-675 какой процессор?\nPRIVATE-CARRIER",
    ):
        assert _self_contained_public_product_spec_query(unsafe_or_contextual) == ""


def test_explicit_public_web_query_is_current_only_and_fail_closed() -> None:
    numeric_version_request = "Найди в интернете официальную документацию Python 3.14 о pathlib"
    assert file_turn_authority(numeric_version_request).source_filenames() == ()
    assert file_turn_authority("Прочитай файл 3.14").source_filenames() == ("3.14",)
    for request, expected_query in (
        (
            "Найди в интернете официальный сайт проекта Nextcloud",
            "официальный сайт проекта Nextcloud",
        ),
        (
            "Найди в интернете, пожалуйста, официальный сайт проекта Nextcloud",
            "официальный сайт проекта Nextcloud",
        ),
        ("Найди мне в интернете документацию Python", "документацию Python"),
        (
            "Найди в интернете официальную документацию Python 3.14 о pathlib",
            "официальную документацию Python 3.14 о pathlib",
        ),
        (
            "Найди в интернете storage capacity Synology DS923+",
            "storage capacity Synology DS923+",
        ),
        ("Найди в интернете биографию Ады Лавлейс", "биографию Ады Лавлейс"),
        ("Найди в интернете IT-сертификацию ISO 27001", "IT-сертификацию ISO 27001"),
        ("Google it Nextcloud", "Nextcloud"),
    ):
        assert _self_contained_explicit_public_web_query(request) == expected_query
    for unsafe_or_contextual in (
        "Найди в интернете это",
        "Google it",
        "Найди в интернете данные из этого файла",
        "Найди в интернете данные из документа",
        "Найди в интернете сведения из моего архива",
        "Найди в интернете PRIVATE-CANARY из письма",
        "Найди в интернете PRIVATE-CANARY из нашей CRM",
        "Найди в интернете что писал Иванов",
        "Найди в интернете профиль сотрудника Иванова",
        "Найди в интернете где он работает",
        "Найди в интернете Nextcloud и позвони Иванову",
        "Найди в интернете пароль от моего роутера",
        "Найди в интернете па\u200bроль от роутера",
        "Найди в интернете PRIVATE-CANARY-MARKER",
        "Найди в интернете report.pdf",
        "Найди в интернете https://example.test/private",
        "«Найди в интернете официальный сайт проекта Nextcloud»",
        "Найди в интернете официальный сайт проекта Nextcloud\nPRIVATE-CARRIER",
    ):
        assert _self_contained_explicit_public_web_query(unsafe_or_contextual) == ""


def test_public_market_query_is_current_only_and_fail_closed() -> None:
    request = "Как менялась стоимость и сложность монеты monero с весны этого года по данный момент?"
    assert _self_contained_public_market_query(request) == request
    assert _self_contained_public_market_query("Какой курс доллара сегодня?") == (
        "Какой курс доллара сегодня?"
    )
    for unsafe_or_contextual in (
        "Как менялась стоимость monero?",
        "Как менялась стоимость этой монеты с весны этого года?",
        "Как менялась стоимость monero из моего файла с весны этого года?",
        "Как менялась стоимость PRIVATE-CANARY с весны этого года?",
        "«Как менялась стоимость monero с весны этого года?»",
        "Запиши стоимость monero на данный момент",
        "Как менялась стоимость monero с весны этого года?\nPRIVATE-CARRIER",
    ):
        assert _self_contained_public_market_query(unsafe_or_contextual) == ""


def _actor() -> ActorContext:
    return ActorContext(user_id=OWNER, preset_key="owner", source="synthetic-test")


class _AllowAll:
    def authorize(self, actor, capability, **kwargs):  # noqa: ANN001, ARG002
        return SimpleNamespace(allowed=True)


def _tool(name: str) -> dict[str, Any]:
    return {
        "type": "function",
        "function": {
            "name": name,
            "description": "synthetic bounded public read",
            "parameters": {"type": "object"},
        },
    }


class _SyntheticWebKernel:
    authorization = _AllowAll()

    def __init__(
        self,
        *,
        url: str = PUBLIC_URL,
        fact: str = PUBLIC_FACT,
        title: str = "Synthetic public news source",
    ) -> None:
        self.calls: list[tuple[str, dict[str, Any]]] = []
        self.url = url
        self.fact = fact
        self.title = title

    def get_tool_definitions(self, actor, topic=""):  # noqa: ANN001, ARG002
        return [_tool("web_research")]

    @staticmethod
    def get_tool(name: str) -> Any:
        if name == "web_research":
            return SimpleNamespace(
                name="web_research",
                risk="mutate",
                security_id="web.research",
            )
        return None

    async def execute(self, tool, params, actor=None):  # noqa: ANN001, ARG002
        assert tool == "web_research"
        self.calls.append((str(tool), dict(params)))
        freshness = str(params.get("freshness") or "")
        filter_proof = (
            {"freshness": freshness, "applied_search_filters": {"freshness": freshness}} if freshness else {}
        )
        return ToolResult(
            tool,
            True,
            {
                "outbound_attempted": True,
                "query": str(params.get("query") or ""),
                **filter_proof,
                "sources": [
                    {
                        "url": self.url,
                        "title": self.title,
                        "text": self.fact,
                        "text_length": len(self.fact),
                        "status_code": 200,
                        "error": "",
                        "truncated": False,
                    }
                ],
                "requested_sources": 1,
                "completed_sources": 1,
                "failed_sources": 0,
                "timed_out_sources": 0,
                "search_timed_out": False,
            },
        )


class _ScriptedModel:
    enabled = True
    model = "synthetic-owner-regression"
    total_budget_sec = 1.0

    def __init__(
        self,
        answers: dict[str, str],
        *,
        web_fact: str = PUBLIC_FACT,
        web_url: str = PUBLIC_URL,
    ) -> None:
        self.answers = dict(answers)
        self.web_fact = web_fact
        self.web_url = web_url
        self.calls: list[dict[str, Any]] = []

    async def chat(self, messages, tools=None, **kwargs):  # noqa: ANN001, ARG002
        snapshot = {
            "messages": [dict(item) for item in messages],
            "tools": [dict(item) for item in (tools or [])],
        }
        self.calls.append(snapshot)
        rendered = json.dumps(snapshot, ensure_ascii=False)
        if self.web_fact in rendered and self.web_url in rendered:
            return {
                "content": f"{self.web_fact} {self.web_url}",
                "tool_calls": None,
                "_queue_wait_sec": 0.0,
            }
        for marker, answer in self.answers.items():
            if marker in rendered:
                return {
                    "content": answer,
                    "tool_calls": None,
                    "_queue_wait_sec": 0.0,
                }
        raise AssertionError("synthetic model received an uncontracted payload")


@pytest.mark.asyncio
async def test_foreign_source_authority_survives_the_arbiter_query_rewrite(settings, storage) -> None:
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_ScriptedModel({}),
        kernel=kernel,
    )
    request = "Какие мировые новости вышли вчера?"
    context = AgentContext(
        conversation_id="foreign-rewrite",
        user_id=OWNER,
        outward_verdict=("интернет", "elections latest developments"),
    )
    messages: list[dict[str, Any]] = []
    tools_used: list[str] = []
    evidence: list[dict[str, str]] = []

    await runtime._prefetch_the_web_if_asked(  # noqa: SLF001
        request,
        _actor(),
        [_tool("web_research")],
        messages,
        tools_used,
        evidence,
        [],
        context,
    )

    assert kernel.calls == [
        (
            "web_research",
            {
                "query": "elections latest developments",
                "max_sources": 3,
                "source_class": "foreign",
            },
        )
    ]
    assert context.web_evidence_status == "sourced"


class _NeverModel:
    enabled = True
    model = "synthetic-never-called"
    total_budget_sec = 1.0

    async def chat(self, messages, tools=None, **kwargs):  # noqa: ANN001, ARG002
        raise AssertionError("code-owned file/privacy path reached the model")


@dataclass(frozen=True)
class _StoredDocument:
    raw: RawObject
    attachment: dict[str, Any]
    text: str
    names: tuple[str, ...]
    marker: str


def _render_docx(
    *,
    document_number: int,
    padding: int,
    prefix: str = PRIVATE_PREFIX,
) -> tuple[bytes, Any, tuple[str, ...]]:
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "ФИО"
    table.rows[0].cells[1].text = "Должность"
    names = tuple(f"{prefix}-{document_number}-PERSON-{row_number:02d}" for row_number in range(1, 4))
    for row_number, name in enumerate(names, start=1):
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = f"SYNTHETIC-ROLE-{document_number}-{row_number:02d}" + (
            "Z" * padding if row_number == 3 else ""
        )
    stream = io.BytesIO()
    document.save(stream)
    payload = stream.getvalue()
    extracted = DocumentExtractor().extract(payload, f"synthetic-{document_number}.docx")
    assert extracted.success is True
    return payload, extracted, names


def _store_docx(
    settings,
    storage,
    *,
    target_chars: int,
    document_number: int,
    complete: bool = True,
    prefix: str = PRIVATE_PREFIX,
) -> _StoredDocument:
    _, baseline, _ = _render_docx(
        document_number=document_number,
        padding=0,
        prefix=prefix,
    )
    padding = target_chars - len(baseline.text)
    assert padding >= 0
    payload, extracted, names = _render_docx(
        document_number=document_number,
        padding=padding,
        prefix=prefix,
    )
    assert len(extracted.text) == target_chars
    index = extracted.office_structure_index
    assert isinstance(index, dict)
    assert index.get("complete") is True
    assert validate_runtime_office_index(index, extracted.text) == index

    source_hash = hashlib.sha256(payload).hexdigest()
    relative_path = f"{OWNER}/{source_hash[:2]}/{source_hash}.docx"
    stored_path = settings.files_dir / relative_path
    stored_path.parent.mkdir(parents=True, exist_ok=True)
    stored_path.write_bytes(payload)
    metadata: dict[str, Any] = {
        "filename": f"synthetic-{document_number}-{target_chars}.docx",
        "uploaded_by": OWNER,
        "extraction_success": True,
        "text_extraction_success": True,
        "stored_path": relative_path,
        "sha256": source_hash,
        "size_bytes": len(payload),
    }
    if complete:
        token = sign_office_structure_index(storage, index, source_hash)
        assert isinstance(token, str)
        metadata.update(
            {
                OFFICE_STRUCTURE_KEY: index,
                OFFICE_STRUCTURE_ATTESTATION_METADATA_KEY: token,
            }
        )
    else:
        # The text is readable, but the parser explicitly cannot prove that the
        # source ended where the extracted text ends.
        metadata["source_truncated_for_parse"] = True

    raw = RawObject(
        id=new_id("raw"),
        user_id=OWNER,
        source="synthetic-docx",
        source_ref=new_id("synthetic-source"),
        raw_content=extracted.text,
        content_type="file",
        content_hash=source_hash,
        metadata_json=metadata,
    )
    storage.store_raw_object(raw)
    attachment = {
        "raw_object_id": raw.id,
        "filename": metadata["filename"],
        "transient_text": extracted.text,
        "extraction_success": True,
        "verification_eligible": True,
        **({"source_truncated_for_parse": True} if not complete else {}),
    }
    return _StoredDocument(
        raw=raw,
        attachment=attachment,
        text=extracted.text,
        names=names,
        marker=names[0],
    )


def _store_generic_text(
    settings,
    storage,
    *,
    marker: str = "PRIVATE-SYNTHETIC-LOCAL-NEWS-CANARY",
) -> _StoredDocument:
    text = (
        f"{marker}\nПо данным из интернета, в документе за прошедшие сутки есть "
        "синтетическая локальная запись."
    )
    payload = text.encode()
    source_hash = hashlib.sha256(payload).hexdigest()
    relative_path = f"{OWNER}/{source_hash[:2]}/{source_hash}.txt"
    stored_path = settings.files_dir / relative_path
    stored_path.parent.mkdir(parents=True, exist_ok=True)
    stored_path.write_bytes(payload)
    raw = RawObject(
        id=new_id("raw"),
        user_id=OWNER,
        source="synthetic-text",
        source_ref=new_id("synthetic-source"),
        raw_content=text,
        content_type="file",
        content_hash=source_hash,
        metadata_json={
            "filename": "synthetic-local-news.txt",
            "uploaded_by": OWNER,
            "extraction_success": True,
            "text_extraction_success": True,
            "stored_path": relative_path,
            "sha256": source_hash,
            "size_bytes": len(payload),
        },
    )
    storage.store_raw_object(raw)
    return _StoredDocument(
        raw=raw,
        attachment={
            "raw_object_id": raw.id,
            "filename": "synthetic-local-news.txt",
            "transient_text": text,
            "extraction_success": True,
            "verification_eligible": True,
        },
        text=text,
        names=(),
        marker=marker,
    )


def _stored_metadata(storage, response: dict[str, Any]) -> dict[str, Any]:
    row = storage.get_message(str(response["message_id"]), OWNER)
    assert row is not None
    metadata = json.loads(str(row["metadata_json"] or "{}"))
    assert isinstance(metadata, dict)
    return metadata


async def _upload_notice(
    runtime: AgentRuntime,
    document: _StoredDocument,
    *,
    conversation_id: str | None = None,
) -> dict[str, Any]:
    return await runtime.chat(
        OWNER,
        f"Загружен документ: {document.attachment['filename']}",
        actor=_actor(),
        conversation_id=conversation_id,
        attachments=[document.attachment],
        synthetic_document_notice=True,
    )


def _source_user_message(storage, response: dict[str, Any]) -> dict[str, Any]:
    rows = storage.get_conversation_messages(str(response["conversation_id"]), user_id=OWNER)
    source = next(item for item in reversed(rows) if item["role"] == "user")
    return source


@pytest.mark.asyncio
async def test_three_complete_bare_docx_summaries_then_isolated_news_excludes_history(
    settings,
    storage,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    documents = [
        _store_docx(settings, storage, target_chars=target, document_number=number)
        for number, target in enumerate((832, 1416, 1590), start=1)
    ]
    expected_summaries = {
        document.marker: (
            f"Синтетическая сводка {number}: в документе указаны ровно 3 позиции — "
            f"{', '.join(document.names)}."
        )
        for number, document in enumerate(documents, start=1)
    }
    model = _ScriptedModel(expected_summaries)
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=kernel,
    )

    conversation_id: str | None = None
    for document in documents:
        response = await _upload_notice(runtime, document, conversation_id=conversation_id)
        conversation_id = str(response["conversation_id"])

        assert response["message"] == expected_summaries[document.marker]
        assert "Быстрый обзор" not in response["message"]
        assert response["message"] != OFFICE_EXACT_UNAVAILABLE_MESSAGE
        assert response["tools_used"] == []
        assert response["attachment_context_expected_count"] == 1
        assert response["attachment_context_readable_count"] == 1
        assert response["attachment_coverage_complete"] is True
        assert response["attachment_verification_complete"] is True
        assert response["restored_attachment_count"] == 0
        metadata = _stored_metadata(storage, response)
        assert metadata["structural"]["model_spoke"] is True
        assert metadata["structural"]["verdict_kind"] == ""
        assert metadata["attachment_coverage_complete"] is True

    assert len(model.calls) == 3
    assert kernel.calls == []
    isolated_request = "Сделай сводку по новостям на зарубежных сайтах"
    assert asks_for_the_web(isolated_request) is True
    assert file_turn_authority(isolated_request).actions == frozenset({"web"})

    news = await runtime.chat(
        OWNER,
        isolated_request,
        actor=_actor(),
        conversation_id=conversation_id,
    )

    assert kernel.calls == [
        (
            "web_research",
            {
                "query": runtime.web_query_from(isolated_request),
                "max_sources": 3,
                "freshness": "day",
                "source_class": "foreign",
            },
        )
    ]
    assert news["tools_used"] == ["web_research"]
    assert news["web_evidence_status"] == "sourced"
    assert news["web_sources"] == [{"url": PUBLIC_URL, "title": "Synthetic public news source"}]
    assert news["attachment_context_expected_count"] == 0
    assert news["attachment_context_readable_count"] == 0
    assert news["restored_attachment_count"] == 0
    metadata = _stored_metadata(storage, news)
    assert metadata["private_context_lineage"] is True
    assert metadata["structural"].get("private_web_search_blocked") is not True

    assert len(model.calls) > 3
    public_payload = json.dumps(model.calls[-1], ensure_ascii=False)
    outbound_payload = json.dumps(kernel.calls, ensure_ascii=False)
    assert isolated_request in public_payload
    assert PUBLIC_FACT in public_payload
    assert PRIVATE_PREFIX not in public_payload
    assert "Загружен документ" not in public_payload
    assert PRIVATE_PREFIX not in outbound_payload


@pytest.mark.asyncio
async def test_qnap_product_spec_uses_isolated_web_after_private_history(
    settings,
    storage,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    conversation = storage.create_conversation(OWNER, title="synthetic private history")
    conversation_id = str(conversation["id"])
    storage.store_message(
        conversation_id,
        OWNER,
        "user",
        f"Загружен документ: {PRIVATE_HISTORY_CANARY}.docx",
        metadata={"had_attachments": True, "private_context_lineage": True},
    )
    storage.store_message(
        conversation_id,
        OWNER,
        "assistant",
        f"Приватная сводка: {PRIVATE_HISTORY_CANARY}",
        metadata={"attachment_context_used": True, "private_context_lineage": True},
    )

    kernel = _SyntheticWebKernel(
        url=PRODUCT_URL,
        fact=PRODUCT_FACT,
        title="Synthetic QNAP TVS-675 specification",
    )
    model = _ScriptedModel({}, web_fact=PRODUCT_FACT, web_url=PRODUCT_URL)
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=kernel,
    )
    request = "в qnap TVS-675 какой процессор?"

    response = await runtime.chat(
        OWNER,
        request,
        actor=_actor(),
        conversation_id=conversation_id,
    )

    assert kernel.calls == [
        (
            "web_research",
            {"query": "qnap TVS-675 процессор", "max_sources": 3},
        )
    ]
    assert response["tools_used"] == ["web_research"]
    assert response["web_evidence_status"] == "sourced"
    assert response["web_sources"] == [{"url": PRODUCT_URL, "title": "Synthetic QNAP TVS-675 specification"}]
    assert PRODUCT_FACT in response["message"]
    metadata = _stored_metadata(storage, response)
    assert metadata["private_context_lineage"] is True
    assert metadata["structural"].get("private_web_search_blocked") is not True

    model_payload = json.dumps(model.calls, ensure_ascii=False)
    outbound_payload = json.dumps(kernel.calls, ensure_ascii=False)
    assert request in model_payload
    assert PRODUCT_FACT in model_payload
    assert PRIVATE_HISTORY_CANARY not in model_payload
    assert PRIVATE_HISTORY_CANARY not in outbound_payload
    assert "Загружен документ" not in model_payload


@pytest.mark.asyncio
async def test_current_monero_market_history_uses_isolated_web_after_private_history(
    settings,
    storage,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    conversation = storage.create_conversation(OWNER, title="synthetic private history")
    conversation_id = str(conversation["id"])
    storage.store_message(
        conversation_id,
        OWNER,
        "user",
        f"Загружен документ: {PRIVATE_HISTORY_CANARY}.docx",
        metadata={"had_attachments": True, "private_context_lineage": True},
    )
    storage.store_message(
        conversation_id,
        OWNER,
        "assistant",
        f"Приватная сводка: {PRIVATE_HISTORY_CANARY}",
        metadata={"attachment_context_used": True, "private_context_lineage": True},
    )
    fact = "Synthetic public Monero price and network-difficulty history is source-backed."
    url = "https://public.synthetic.example.com/monero-history"
    kernel = _SyntheticWebKernel(url=url, fact=fact, title="Synthetic Monero market history")
    model = _ScriptedModel({}, web_fact=fact, web_url=url)
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=kernel,
    )
    request = "Как менялась стоимость и сложность монеты monero с весны этого года по данный момент?"

    response = await runtime.chat(
        OWNER,
        request,
        actor=_actor(),
        conversation_id=conversation_id,
    )

    assert kernel.calls == [("web_research", {"query": request, "max_sources": 3})]
    assert response["tools_used"] == ["web_research"]
    assert response["web_evidence_status"] == "sourced"
    assert fact in response["message"]
    metadata = _stored_metadata(storage, response)
    assert metadata["private_context_lineage"] is True
    assert metadata["structural"].get("private_web_search_blocked") is not True
    exposed = json.dumps([model.calls, kernel.calls], ensure_ascii=False)
    assert PRIVATE_HISTORY_CANARY not in exposed
    assert "Загружен документ" not in exposed


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("query_text", "expected_query"),
    [
        (
            "Найди в интернете официальный сайт проекта Nextcloud",
            "официальный сайт проекта Nextcloud",
        ),
        ("Найди мне в интернете документацию Python", "документацию Python"),
        (
            "Найди в интернете официальную документацию Python 3.14 о pathlib",
            "официальную документацию Python 3.14 о pathlib",
        ),
        (
            "Найди в интернете storage capacity Synology DS923+",
            "storage capacity Synology DS923+",
        ),
        ("Найди в интернете биографию Ады Лавлейс", "биографию Ады Лавлейс"),
    ],
)
async def test_explicit_public_web_uses_history_free_lane_after_private_lineage(
    settings,
    storage,
    query_text: str,
    expected_query: str,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    conversation = storage.create_conversation(OWNER, title="synthetic private history")
    conversation_id = str(conversation["id"])
    storage.store_message(
        conversation_id,
        OWNER,
        "user",
        f"Загружен документ: {PRIVATE_HISTORY_CANARY}.docx",
        metadata={"had_attachments": True, "private_context_lineage": True},
    )
    storage.store_message(
        conversation_id,
        OWNER,
        "assistant",
        f"Приватная сводка: {PRIVATE_HISTORY_CANARY}",
        metadata={"attachment_context_used": True, "private_context_lineage": True},
    )

    kernel = _SyntheticWebKernel(
        url=GENERIC_WEB_URL,
        fact=GENERIC_WEB_FACT,
        title="Synthetic official Nextcloud project page",
    )
    model = _ScriptedModel({}, web_fact=GENERIC_WEB_FACT, web_url=GENERIC_WEB_URL)
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=kernel,
    )
    response = await runtime.chat(
        OWNER,
        query_text,
        actor=_actor(),
        conversation_id=conversation_id,
    )

    assert kernel.calls == [
        (
            "web_research",
            {"query": expected_query, "max_sources": 3},
        )
    ]
    assert response["tools_used"] == ["web_research"]
    assert response["web_evidence_status"] == "sourced"
    assert response["web_sources"] == [
        {"url": GENERIC_WEB_URL, "title": "Synthetic official Nextcloud project page"}
    ]
    assert GENERIC_WEB_FACT in response["message"]
    metadata = _stored_metadata(storage, response)
    assert metadata["private_context_lineage"] is True
    assert metadata["structural"].get("private_web_search_blocked") is not True

    model_payload = json.dumps(model.calls, ensure_ascii=False)
    outbound_payload = json.dumps(kernel.calls, ensure_ascii=False)
    assert query_text in model_payload
    assert GENERIC_WEB_FACT in model_payload
    assert PRIVATE_HISTORY_CANARY not in model_payload
    assert PRIVATE_HISTORY_CANARY not in outbound_payload
    assert "Загружен документ" not in model_payload


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "query_text",
    [
        "Найди в интернете это",
        "Найди в интернете данные из этого файла",
        "Найди в интернете данные из документа",
        "Найди в интернете сведения из моего архива",
        "Найди в интернете PRIVATE-CANARY из письма",
        "Найди в интернете PRIVATE-CANARY из нашей CRM",
        "Найди в интернете что писал Иванов",
        "Найди в интернете профиль сотрудника Иванова",
        "Найди в интернете где он работает",
        "Найди в интернете Nextcloud и позвони Иванову",
        "Найди в интернете пароль от моего роутера",
        "Найди в интернете па\u200bроль от роутера",
        "Найди в интернете PRIVATE-CANARY-MARKER",
    ],
)
async def test_explicit_public_web_private_contours_never_leave_private_conversation(
    settings,
    storage,
    query_text: str,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    conversation = storage.create_conversation(OWNER, title="synthetic private history")
    conversation_id = str(conversation["id"])
    storage.store_message(
        conversation_id,
        OWNER,
        "user",
        f"Загружен документ: {PRIVATE_HISTORY_CANARY}.docx",
        metadata={"had_attachments": True, "private_context_lineage": True},
    )
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_NeverModel(),
        kernel=kernel,
    )

    assert asks_for_the_web(query_text) is True
    response = await runtime.chat(
        OWNER,
        query_text,
        actor=_actor(),
        conversation_id=conversation_id,
    )

    assert response["tools_used"] == []
    assert response["web_evidence_status"] == "none"
    assert kernel.calls == []
    metadata = _stored_metadata(storage, response)
    assert metadata["structural"]["private_web_search_blocked"] is True


@pytest.mark.asyncio
async def test_current_unused_file_does_not_block_explicit_public_web(
    settings,
    storage,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    document = _store_generic_text(settings, storage)
    kernel = _SyntheticWebKernel(
        url=GENERIC_WEB_URL,
        fact=GENERIC_WEB_FACT,
        title="Synthetic official Nextcloud project page",
    )
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_ScriptedModel(
            {document.marker: "Синтетическая локальная сводка."},
            web_fact=GENERIC_WEB_FACT,
            web_url=GENERIC_WEB_URL,
        ),
        kernel=kernel,
    )
    query_text = "Найди в интернете официальный сайт проекта Nextcloud"
    response = await runtime.chat(
        OWNER,
        query_text,
        actor=_actor(),
        attachments=[document.attachment],
    )
    assert kernel.calls
    query = str(kernel.calls[0][1].get("query") or "")
    assert "nextcloud" in query.casefold()
    assert document.marker not in query
    assert GENERIC_WEB_FACT in response["message"]
    metadata = _stored_metadata(storage, response)
    assert metadata["structural"].get("private_web_search_blocked") is not True


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("carrier", "expected_count", "expected_restored"),
    [
        ("reply", 0, 0),
        ("replay", 0, 0),
        ("restored", 1, 1),
    ],
)
async def test_explicit_public_web_attachment_carriers_never_enter_isolated_lane(
    settings,
    storage,
    carrier: str,
    expected_count: int,
    expected_restored: int,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    conversation = storage.create_conversation(OWNER, title="synthetic private history")
    conversation_id = str(conversation["id"])
    storage.store_message(
        conversation_id,
        OWNER,
        "user",
        f"Загружен документ: {PRIVATE_HISTORY_CANARY}.docx",
        metadata={"had_attachments": True, "private_context_lineage": True},
    )
    storage.store_message(
        conversation_id,
        OWNER,
        "assistant",
        f"Приватная сводка: {PRIVATE_HISTORY_CANARY}",
        metadata={"attachment_context_used": True, "private_context_lineage": True},
    )
    document = _store_docx(settings, storage, target_chars=832, document_number=93)
    model = _ScriptedModel({document.marker: "Синтетическая локальная сводка."})
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=kernel,
    )
    query_text = "Найди в интернете официальный сайт проекта Nextcloud"
    chat_kwargs: dict[str, Any] = {}
    if carrier == "reply":
        chat_kwargs["reply_to"] = f"Приватный ответ: {PRIVATE_HISTORY_CANARY}"
    elif carrier == "replay":
        replay_source = storage.store_message(
            conversation_id,
            OWNER,
            "user",
            query_text,
            metadata={"private_context_lineage": True},
        )
        chat_kwargs["replay_source_message_id"] = str(replay_source["id"])
    else:
        await _upload_notice(runtime, document, conversation_id=conversation_id)
        query_text = "Найди в интернете официальные характеристики по этому файлу"

    response = await runtime.chat(
        OWNER,
        query_text,
        actor=_actor(),
        conversation_id=conversation_id,
        **chat_kwargs,
    )

    assert response["tools_used"] == []
    assert response["web_evidence_status"] == "none"
    assert response["attachment_context_expected_count"] == expected_count
    assert response["restored_attachment_count"] == expected_restored
    assert kernel.calls == []
    metadata = _stored_metadata(storage, response)
    assert metadata["structural"]["private_web_search_blocked"] is True


@pytest.mark.asyncio
async def test_old_attachment_lineage_beyond_prompt_tail_still_allows_only_current_news(
    settings,
    storage,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    document = _store_docx(settings, storage, target_chars=832, document_number=41)
    summary = "Синтетическая закрытая сводка с PRIVATE-HISTORICAL-ANSWER-CANARY."
    model = _ScriptedModel({document.marker: summary})
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=kernel,
    )
    uploaded = await _upload_notice(runtime, document)
    conversation_id = str(uploaded["conversation_id"])

    stale = (datetime.now(UTC) - timedelta(days=3)).isoformat()
    with storage.transaction() as connection:
        connection.execute(
            "UPDATE messages SET created_at=? WHERE conversation_id=? AND user_id=?",
            (stale, conversation_id, OWNER),
        )
    for number in range(24):
        storage.store_message(
            conversation_id,
            OWNER,
            "user",
            f"SYNTHETIC-CLEAN-USER-{number:02d}",
            metadata={"private_context_lineage": True},
        )
        storage.store_message(
            conversation_id,
            OWNER,
            "assistant",
            f"SYNTHETIC-CLEAN-ASSISTANT-{number:02d}",
            metadata={"private_context_lineage": True},
        )

    isolated_request = "Сделай сводку по новостям на зарубежных сайтах"
    response = await runtime.chat(
        OWNER,
        isolated_request,
        actor=_actor(),
        conversation_id=conversation_id,
    )

    assert response["tools_used"] == ["web_research"]
    assert response["web_evidence_status"] == "sourced"
    assert response["restored_attachment_count"] == 0
    assert len(kernel.calls) == 1
    public_payload = json.dumps(model.calls[-1], ensure_ascii=False)
    assert PRIVATE_PREFIX not in public_payload
    assert "PRIVATE-HISTORICAL-ANSWER-CANARY" not in public_payload
    # The sole public exception is current-message-only: sticky lineage remains
    # durable, but neither the old file nor intervening dialogue participates.
    assert "SYNTHETIC-CLEAN-USER" not in public_payload
    assert "SYNTHETIC-CLEAN-ASSISTANT" not in public_payload
    assert PRIVATE_PREFIX not in json.dumps(kernel.calls, ensure_ascii=False)


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("question", "lists_people"),
    [
        ("Перечисли всех людей из файла.", True),
        ("Сколько людей в файле?", False),
    ],
)
async def test_complete_signed_docx_exact_list_and_count_are_code_owned(
    settings,
    storage,
    question: str,
    lists_people: bool,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    document = _store_docx(settings, storage, target_chars=1416, document_number=52)
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=True, verify_min_answer_chars=1),
        storage,
        llm=_NeverModel(),
        kernel=kernel,
    )

    response = await runtime.chat(
        OWNER,
        question,
        actor=_actor(),
        attachments=[document.attachment],
    )

    assert response["verification_status"] == "passed"
    assert response["verified"] is True
    assert response["message_format"] == "plain"
    assert "3" in response["message"]
    if lists_people:
        assert all(name in response["message"] for name in document.names)
    else:
        assert all(name not in response["message"] for name in document.names)
    assert response["tools_used"] == []
    assert response["web_evidence_status"] == "none"
    assert response["attachment_coverage_complete"] is True
    assert kernel.calls == []
    metadata = _stored_metadata(storage, response)
    assert metadata["structural"]["verdict_kind"] == "office_exact"
    assert metadata["structural"]["model_spoke"] is False


@pytest.mark.asyncio
async def test_readable_incomplete_docx_is_partial_then_exact_followup_is_unknown(
    settings,
    storage,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    document = _store_docx(
        settings,
        storage,
        target_chars=832,
        document_number=63,
        complete=False,
        prefix="PRIVATE-SYNTHETIC-INCOMPLETE-CANARY",
    )
    partial = f"По доступной части документа видна позиция {document.names[0]}; полный состав неизвестен."
    model = _ScriptedModel({document.marker: partial})
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=kernel,
    )

    uploaded = await _upload_notice(runtime, document)

    # A structurally incomplete source may be settled before synthesis.  What
    # matters publicly is that readable bytes remain readable while whole-file
    # certainty is UNKNOWN; it must not be reported as an absent/lost upload.
    assert uploaded["message"] != OFFICE_EXACT_UNAVAILABLE_MESSAGE
    assert "не весь исходный материал" in uploaded["message"].casefold()
    assert "частич" in uploaded["message"].casefold()
    assert "прикреп" not in uploaded["message"].casefold()
    assert "файл отсутств" not in uploaded["message"].casefold()
    assert uploaded["attachment_context_available"] is True
    assert uploaded["attachment_context_readable_count"] == 1
    assert uploaded["attachment_coverage_complete"] is False
    assert uploaded["attachment_verification_complete"] is False
    assert uploaded["verification_status"] == "unknown"
    assert uploaded["tools_used"] == []
    calls_after_partial = len(model.calls)

    exact = await runtime.chat(
        OWNER,
        "Перечисли всех людей из файла.",
        actor=_actor(),
        conversation_id=str(uploaded["conversation_id"]),
    )

    assert exact["message"].endswith(OFFICE_EXACT_UNAVAILABLE_MESSAGE)
    assert "Не весь исходный материал" in exact["message"]
    assert exact["verification_status"] == "unknown"
    assert exact["verified"] is False
    assert exact["restored_attachment_count"] == 1
    assert exact["attachment_context_available"] is True
    assert exact["attachment_context_readable_count"] == 1
    assert exact["attachment_coverage_complete"] is False
    assert exact["tools_used"] == []
    # The mandatory whole-source prepass may still analyse every readable byte;
    # the exact-set guard owns the final UNKNOWN and discards any exhaustive
    # model claim when the structural source is incomplete.
    assert len(model.calls) > calls_after_partial
    assert kernel.calls == []
    metadata = _stored_metadata(storage, exact)
    assert metadata["structural"]["verdict_kind"] == "office_exact"
    assert metadata["structural"]["model_spoke"] is False


@pytest.mark.asyncio
async def test_news_inside_a_current_document_is_local_not_a_web_request(
    settings,
    storage,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    request = "Покажи новости в документе за прошедшие сутки."
    assert asks_for_the_web(request) is False
    document = _store_generic_text(settings, storage)
    answer = "В документе есть одна синтетическая локальная запись за прошедшие сутки."
    model = _ScriptedModel({document.marker: answer})
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=kernel,
    )

    response = await runtime.chat(
        OWNER,
        request,
        actor=_actor(),
        attachments=[document.attachment],
    )

    assert response["message"] == answer
    assert response["tools_used"] == []
    assert response["web_evidence_status"] == "none"
    assert response["attachment_context_expected_count"] == 1
    assert response["attachment_context_readable_count"] == 1
    assert kernel.calls == []
    metadata = _stored_metadata(storage, response)
    assert metadata["structural"].get("private_web_search_blocked") is not True


@pytest.mark.asyncio
async def test_same_sentence_document_summary_and_web_request_denies_public_research(
    settings,
    storage,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    request = "Обобщи весь документ и поищи актуальные данные в интернете."
    assert asks_for_the_web(request) is True
    document = _store_generic_text(settings, storage)
    model = _ScriptedModel({document.marker: "Синтетическая локальная сводка."})
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=kernel,
    )

    response = await runtime.chat(
        OWNER,
        request,
        actor=_actor(),
        attachments=[document.attachment],
    )

    assert kernel.calls == []
    assert response["web_evidence_status"] == "none"
    assert response["attachment_context_expected_count"] == 1
    assert "приватные вложения" not in response["message"].casefold()
    assert "Синтетическая локальная сводка." in response["message"]
    metadata = _stored_metadata(storage, response)
    assert metadata["structural"].get("private_web_search_blocked") is not True


@pytest.mark.asyncio
async def test_current_unused_file_does_not_block_public_news(settings, storage) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    document = _store_docx(settings, storage, target_chars=832, document_number=74)
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_ScriptedModel({document.marker: "Синтетическая локальная сводка."}),
        kernel=kernel,
    )
    response = await runtime.chat(
        OWNER,
        NEWS_REQUEST,
        actor=_actor(),
        attachments=[document.attachment],
    )
    assert kernel.calls
    assert kernel.calls[0][0] == "web_research"
    query = str(kernel.calls[0][1].get("query") or "")
    assert document.marker not in query
    assert "приватные вложения" not in response["message"].casefold()
    metadata = _stored_metadata(storage, response)
    assert metadata["structural"].get("private_web_search_blocked") is not True


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("carrier", "expected_count", "expected_restored"),
    [
        ("explicit_reference", 1, 1),
        ("deictic_reference", 1, 1),
        ("reply_reference", 0, 0),
        ("replayed_attachment", 1, 1),
    ],
)
async def test_attachment_derived_news_carriers_cannot_use_web(
    settings,
    storage,
    carrier: str,
    expected_count: int,
    expected_restored: int,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    document = _store_docx(settings, storage, target_chars=832, document_number=74)
    summary = f"Синтетическая сводка: {document.names[0]}."
    model = _ScriptedModel({document.marker: summary})
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=model,
        kernel=kernel,
    )

    conversation_id: str | None = None
    query = NEWS_REQUEST
    chat_kwargs: dict[str, Any] = {}
    if carrier in {"explicit_reference", "deictic_reference", "reply_reference"}:
        uploaded = await _upload_notice(runtime, document)
        conversation_id = str(uploaded["conversation_id"])
        if carrier == "explicit_reference":
            query = "Найди в интернете свежие новости из загруженного документа за прошедшие сутки."
        elif carrier == "deictic_reference":
            query = "Найди в интернете свежие новости по нему за прошедшие сутки."
        else:
            chat_kwargs["reply_to"] = f"Ответ по файлу: {document.marker}"
    else:
        first = await runtime.chat(
            OWNER,
            query,
            actor=_actor(),
            attachments=[document.attachment],
        )
        conversation_id = str(first["conversation_id"])
        source = _source_user_message(storage, first)
        chat_kwargs["replay_source_message_id"] = str(source["id"])

    kernel_calls_before = len(kernel.calls)
    assert asks_for_the_web(query) is True
    response = await runtime.chat(
        OWNER,
        query,
        actor=_actor(),
        conversation_id=conversation_id,
        **chat_kwargs,
    )

    assert response["tools_used"] == []
    assert response["web_evidence_status"] == "none"
    assert response["attachment_context_expected_count"] == expected_count
    assert response["restored_attachment_count"] == expected_restored
    assert "приватные вложения" in response["message"].casefold()
    assert kernel.calls[kernel_calls_before:] == []
    metadata = _stored_metadata(storage, response)
    assert metadata["structural"]["private_web_search_blocked"] is True
    assert metadata["structural"]["model_spoke"] is False


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "request_text",
    [
        "Найди в интернете текущие публичные правила и сравни их с этим договором.",
        "Сравни этот договор с текущими публичными правилами в интернете.",
        "Compare this contract with current public HTTP Date header rules on the web.",
    ],
)
async def test_current_file_public_web_compare_searches_only_the_public_topic(
    settings,
    storage,
    request_text: str,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    document = _store_generic_text(settings, storage)
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_ScriptedModel({document.marker: "Локальная сверка."}, web_fact=PUBLIC_FACT, web_url=PUBLIC_URL),
        kernel=kernel,
    )
    response = await runtime.chat(
        OWNER,
        request_text,
        actor=_actor(),
        attachments=[document.attachment],
    )

    assert kernel.calls, "public-web compare must prefetch web"
    assert len(kernel.calls) == 1
    query = str(kernel.calls[0][1].get("query") or "")
    folded_query = query.casefold()
    assert "публичн" in folded_query or "public" in folded_query
    assert document.marker not in query
    assert "договор" not in folded_query
    assert "файл" not in folded_query
    assert "contract" not in folded_query
    assert "file" not in folded_query
    assert "интернет" not in folded_query
    assert "web" not in folded_query
    rendered = json.dumps(runtime.llm.calls, ensure_ascii=False) if runtime.llm.calls else ""
    assert document.marker in rendered
    assert document.marker not in query
    metadata = _stored_metadata(storage, response)
    assert metadata["structural"].get("private_web_search_blocked") is not True
    assert "приватные вложения" not in response["message"].casefold()
    assert PUBLIC_FACT in response["message"]
    assert response["web_evidence_status"] == "sourced"
    assert response["web_sources"] == [{"url": PUBLIC_URL, "title": "Synthetic public news source"}]


@pytest.mark.asyncio
async def test_summarize_document_and_vague_web_keeps_local_read(
    settings,
    storage,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    request = "Обобщи весь документ и поищи актуальные данные в интернете."
    document = _store_generic_text(settings, storage)
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_ScriptedModel({document.marker: "Синтетическая локальная сводка."}),
        kernel=kernel,
    )
    response = await runtime.chat(
        OWNER,
        request,
        actor=_actor(),
        attachments=[document.attachment],
    )
    assert kernel.calls == []
    assert "приватные вложения" not in response["message"].casefold()
    assert "Синтетическая локальная сводка." in response["message"]
    metadata = _stored_metadata(storage, response)
    assert metadata["structural"].get("private_web_search_blocked") is not True


@pytest.mark.asyncio
async def test_summarize_then_compare_public_rules_searches_only_the_public_topic(
    settings,
    storage,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    request = "Обобщи документ и сравни с текущими публичными правилами в интернете."
    document = _store_generic_text(settings, storage)
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_ScriptedModel(
            {document.marker: "Синтетическая локальная сводка."},
            web_fact=PUBLIC_FACT,
            web_url=PUBLIC_URL,
        ),
        kernel=kernel,
    )
    response = await runtime.chat(
        OWNER,
        request,
        actor=_actor(),
        attachments=[document.attachment],
    )
    assert kernel.calls
    query = str(kernel.calls[0][1].get("query") or "")
    assert "публичн" in query.casefold()
    assert document.marker not in query
    assert "документ" not in query.casefold()
    assert PUBLIC_FACT in response["message"]
    metadata = _stored_metadata(storage, response)
    assert metadata["structural"].get("private_web_search_blocked") is not True


@pytest.mark.asyncio
async def test_two_current_files_open_public_web_compare(
    settings,
    storage,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    first = _store_generic_text(settings, storage, marker="PRIVATE-SYNTHETIC-COMPARE-FILE-A")
    second = _store_generic_text(settings, storage, marker="PRIVATE-SYNTHETIC-COMPARE-FILE-B")
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_ScriptedModel(
            {first.marker: "Синтетическая локальная сводка."},
            web_fact=PUBLIC_FACT,
            web_url=PUBLIC_URL,
        ),
        kernel=kernel,
    )
    response = await runtime.chat(
        OWNER,
        "Найди в интернете текущие публичные правила и сравни их с этим договором.",
        actor=_actor(),
        attachments=[first.attachment, second.attachment],
    )
    assert kernel.calls
    query = str(kernel.calls[0][1].get("query") or "")
    assert "публичн" in query.casefold()
    assert first.marker not in query
    assert second.marker not in query
    assert "приватные вложения" not in response["message"].casefold()
    assert PUBLIC_FACT in response["message"]
    metadata = _stored_metadata(storage, response)
    assert metadata["structural"].get("private_web_search_blocked") is not True


@pytest.mark.asyncio
async def test_prior_file_history_opens_public_web_compare(
    settings,
    storage,
) -> None:
    storage.ensure_user(OWNER, preset_key="owner")
    document = _store_generic_text(settings, storage)
    kernel = _SyntheticWebKernel()
    runtime = AgentRuntime(
        replace(settings, verify_answers=False),
        storage,
        llm=_ScriptedModel(
            {document.marker: "Синтетическая локальная сводка."},
            web_fact=PUBLIC_FACT,
            web_url=PUBLIC_URL,
        ),
        kernel=kernel,
    )
    first = await runtime.chat(
        OWNER,
        "Кратко изложи этот файл.",
        actor=_actor(),
        attachments=[document.attachment],
    )
    kernel.calls.clear()
    response = await runtime.chat(
        OWNER,
        "Найди в интернете текущие публичные правила и сравни их с этим договором.",
        actor=_actor(),
        conversation_id=str(first["conversation_id"]),
    )
    assert kernel.calls
    query = str(kernel.calls[0][1].get("query") or "")
    assert "публичн" in query.casefold()
    assert document.marker not in query
    assert "договор" not in query.casefold()
    assert "приватные вложения" not in response["message"].casefold()
    assert PUBLIC_FACT in response["message"]
    metadata = _stored_metadata(storage, response)
    assert metadata["structural"].get("private_web_search_blocked") is not True
