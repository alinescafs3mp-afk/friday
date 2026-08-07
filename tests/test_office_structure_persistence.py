"""The Office inventory stays content-free at every ingestion boundary.

The parser owns the literal document and returns spans over its unchanged text.
Ingestion may retain that index only beside the Raw Object: it must not let caller
metadata forge the code-owned structure, copy it into Knowledge/Inbox advice, or
turn an explicit no-save inspection into durable state.
"""

from __future__ import annotations

import hashlib
import io
import json
from types import SimpleNamespace
from typing import Any

import pytest

from friday.documents import DocumentExtractor
from friday.documents._office_structure import validate_office_structure_index
from friday.office_attestation import (
    OFFICE_STRUCTURE_ATTESTATION_METADATA_KEY,
    sign_office_structure_index,
    verify_office_structure_attestation,
)

_INDEX_KEY = "office_structure_v1"
_SOURCE_TEXT_KEY = "_office_source_text"
_ATTESTATION_KEY = OFFICE_STRUCTURE_ATTESTATION_METADATA_KEY
_PRIVATE_LITERALS = (
    "SYNTHETIC-PERSON-ZETA-4417",
    "SYNTHETIC-ROLE-OMEGA-3319",
    "SYNTHETIC-CUSTOM-STYLE-7719",
    "Synthetic roster opening",
    "Synthetic roster closing",
    "Person",
    "Role",
)


def _pipeline(settings, storage):
    from friday.ingestion import IngestionPipeline
    from friday.knowledge_graph import KnowledgeGraph

    storage.ensure_user("alice", preset_key="admin")
    return IngestionPipeline(settings, storage, KnowledgeGraph(storage))


def _add_roster(document: Any) -> None:
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Person"
    table.rows[0].cells[1].text = "Role"
    for index in range(3):
        cells = table.add_row().cells
        cells[0].text = f"{_PRIVATE_LITERALS[0]}-{index}"
        cells[1].text = f"{_PRIVATE_LITERALS[1]}-{index}"


def _docx_bytes(
    *,
    interleaved: bool,
    package_note: str = "",
    header_note: str = "",
    nested_note: str = "",
) -> bytes:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    document = Document()
    custom = document.styles.add_style(_PRIVATE_LITERALS[2], WD_STYLE_TYPE.PARAGRAPH)
    opening = document.add_paragraph("Synthetic roster opening")
    opening.style = custom
    if header_note:
        document.sections[0].header.paragraphs[0].text = header_note
    if nested_note:
        # Вложенная таблица — то, что по-прежнему НЕ читается: её содержимое не
        # попадает в текст, а полноту она отнимает. Ровно то сочетание, которое
        # нужно этой пробе: разные документы с одинаковым текстом.
        outer = document.add_table(rows=1, cols=1)
        outer.cell(0, 0).add_table(rows=1, cols=1).cell(0, 0).text = nested_note
    if interleaved:
        _add_roster(document)
    document.add_paragraph("Synthetic roster closing")
    if not interleaved:
        _add_roster(document)
    document.core_properties.comments = package_note
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _raw_metadata(storage, raw_id: str) -> tuple[dict[str, Any], dict[str, Any]]:
    raw = storage.get_raw_object(raw_id, "alice")
    assert raw is not None
    metadata = json.loads(str(raw["metadata_json"]))
    assert isinstance(metadata, dict)
    return raw, metadata


def _assert_index_has_no_document_literals(index: dict[str, Any]) -> None:
    encoded = json.dumps(index, ensure_ascii=False, sort_keys=True)
    assert all(literal not in encoded for literal in _PRIVATE_LITERALS)


def _material_counts(storage) -> dict[str, int]:
    return {
        "raw_objects": int(storage.execute("SELECT COUNT(*) FROM raw_objects").fetchone()[0]),
        "knowledge_objects": int(storage.execute("SELECT COUNT(*) FROM knowledge_objects").fetchone()[0]),
        "inbox": int(storage.execute("SELECT COUNT(*) FROM inbox").fetchone()[0]),
    }


@pytest.mark.anyio
async def test_persisted_office_index_is_raw_only_content_free_and_code_owned(settings, storage):
    pipeline = _pipeline(settings, storage)
    forged = {"schema_version": 1, "literal": "CALLER-FORGED-OFFICE-STRUCTURE"}

    outcome = await pipeline.ingest_file(
        "alice",
        None,
        _docx_bytes(interleaved=True),
        filename="synthetic-roster.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        source_ref="office:persisted",
        metadata={
            "uploaded_by": "alice",
            _INDEX_KEY: forged,
            _ATTESTATION_KEY: "f" * 64,
            _SOURCE_TEXT_KEY: "CALLER-FORGED-SOURCE-TEXT",
        },
    )

    raw, metadata = _raw_metadata(storage, str(outcome["raw_object_id"]))
    index = metadata.get(_INDEX_KEY)
    assert isinstance(index, dict)
    assert index != forged
    assert "CALLER-FORGED-OFFICE-STRUCTURE" not in json.dumps(index, ensure_ascii=False)
    assert _SOURCE_TEXT_KEY not in metadata
    assert validate_office_structure_index(index, str(raw["raw_content"])) == index
    token = metadata.get(_ATTESTATION_KEY)
    assert isinstance(token, str) and token != "f" * 64
    assert verify_office_structure_attestation(storage, index, raw["content_hash"], token)
    _assert_index_has_no_document_literals(index)

    # The internal ingestion result is the source from which API receipts are
    # projected.  Neither it nor the canonical Knowledge metadata may acquire a
    # second copy of the index.
    assert _INDEX_KEY not in json.dumps(outcome, ensure_ascii=False, default=str)
    knowledge_rows = storage.execute(
        "SELECT metadata_json FROM knowledge_objects WHERE raw_object_id=?",
        (raw["id"],),
    ).fetchall()
    assert knowledge_rows, "the normal assessed file path did not exercise Knowledge promotion"
    assert all(_INDEX_KEY not in str(row["metadata_json"] or "") for row in knowledge_rows)
    inbox_rows = storage.execute(
        "SELECT suggestions_json FROM inbox WHERE raw_object_id=?",
        (raw["id"],),
    ).fetchall()
    assert all(_INDEX_KEY not in str(row["suggestions_json"] or "") for row in inbox_rows)


def test_office_attestation_binds_the_exact_index_and_fails_closed(storage) -> None:
    index = {
        "schema_version": 1,
        "format": "docx",
        "text_sha256": "0" * 64,
        "complete": False,
    }
    source_hash = "a" * 64
    token = sign_office_structure_index(storage, index, source_hash)
    assert isinstance(token, str) and len(token) == 64
    assert verify_office_structure_attestation(storage, index, source_hash, token)

    changed = {**index, "complete": True}
    assert not verify_office_structure_attestation(storage, changed, source_hash, token)
    assert not verify_office_structure_attestation(storage, index, "b" * 64, token)
    assert not verify_office_structure_attestation(storage, index, source_hash, "f" * 64)

    class MissingKeyStorage:
        class _Cursor:
            @staticmethod
            def fetchone() -> None:
                return None

        @staticmethod
        def execute(*_args: object, **_kwargs: object) -> _Cursor:
            return MissingKeyStorage._Cursor()

    assert sign_office_structure_index(MissingKeyStorage(), index, source_hash) is None
    assert not verify_office_structure_attestation(MissingKeyStorage(), index, source_hash, token)


def test_office_attestation_cannot_move_between_distinct_files_with_the_same_text(storage) -> None:
    first_bytes = _docx_bytes(interleaved=True, package_note="SYNTHETIC-PACKAGE-A")
    second_bytes = _docx_bytes(interleaved=True, package_note="SYNTHETIC-PACKAGE-B")
    first_hash = hashlib.sha256(first_bytes).hexdigest()
    second_hash = hashlib.sha256(second_bytes).hexdigest()
    assert first_hash != second_hash

    extractor = DocumentExtractor()
    first = extractor.extract(first_bytes, "first.docx")
    second = extractor.extract(second_bytes, "second.docx")
    assert first.text == second.text
    assert first.office_structure_index == second.office_structure_index
    assert isinstance(first.office_structure_index, dict)

    token = sign_office_structure_index(storage, first.office_structure_index, first_hash)
    assert isinstance(token, str)
    assert verify_office_structure_attestation(storage, first.office_structure_index, first_hash, token)
    assert not verify_office_structure_attestation(storage, second.office_structure_index, second_hash, token)


@pytest.mark.anyio
async def test_no_save_office_structure_and_full_source_are_ephemeral_only(settings, storage):
    pipeline = _pipeline(settings, storage)
    before = _material_counts(storage)

    transient = await pipeline.inspect_file_transient(
        _docx_bytes(interleaved=True),
        filename="private-roster.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        preview_chars=1_000,
    )

    index = transient.get(_INDEX_KEY)
    source_text = transient.get(_SOURCE_TEXT_KEY)
    assert isinstance(index, dict)
    assert isinstance(source_text, str) and source_text
    assert validate_office_structure_index(index, source_text) == index
    _assert_index_has_no_document_literals(index)
    after = _material_counts(storage)
    assert after == before
    assert not any(path.is_file() for path in settings.files_dir.rglob("*"))


@pytest.mark.anyio
async def test_no_save_propagates_the_extractors_own_text_budget(settings, storage):
    pipeline = _pipeline(settings, storage)
    pipeline._doc_extractor = SimpleNamespace(
        extract=lambda *_args, **_kwargs: SimpleNamespace(
            text="bounded extractor result",
            metadata={"text_truncated": True},
            success=True,
            error="",
            office_structure_index=None,
        )
    )

    transient = await pipeline.inspect_file_transient(
        b"synthetic",
        filename="bounded.docx",
        preview_chars=48_000,
    )

    assert transient["text_preview"] == "bounded extractor result"
    assert transient["text_truncated"] is True


@pytest.mark.anyio
async def test_same_flat_office_text_with_different_structure_is_not_deduplicated(settings, storage):
    pipeline = _pipeline(settings, storage)
    first = await pipeline.ingest_file(
        "alice",
        None,
        _docx_bytes(interleaved=False),
        filename="roster-a.docx",
        source_ref="office:layout:a",
        force_review=True,
        metadata={"uploaded_by": "alice"},
    )
    second = await pipeline.ingest_file(
        "alice",
        None,
        _docx_bytes(interleaved=True),
        filename="roster-b.docx",
        source_ref="office:layout:b",
        force_review=True,
        metadata={"uploaded_by": "alice"},
    )

    first_raw, first_metadata = _raw_metadata(storage, str(first["raw_object_id"]))
    second_raw, second_metadata = _raw_metadata(storage, str(second["raw_object_id"]))
    assert first_raw["raw_content"] == second_raw["raw_content"], "fixture lost the legacy flat-text premise"
    assert first_raw["content_hash"] != second_raw["content_hash"]
    assert first_raw["id"] != second_raw["id"]
    assert first_metadata[_INDEX_KEY] != second_metadata[_INDEX_KEY]
    assert storage.execute("SELECT COUNT(*) FROM raw_objects WHERE user_id='alice'").fetchone()[0] == 2


@pytest.mark.parametrize("office_first", [False, True])
@pytest.mark.anyio
async def test_cross_format_text_match_does_not_depend_on_upload_order(
    settings,
    storage,
    office_first,
):
    from friday.documents import DocumentExtractor

    pipeline = _pipeline(settings, storage)
    office_bytes = _docx_bytes(interleaved=True)
    office_text = DocumentExtractor().extract(office_bytes, "roster.docx").text
    office_upload = {
        "file_content": office_bytes,
        "filename": "roster.docx",
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "source_ref": "office:cross-format",
    }
    text_upload = {
        "file_content": office_text.encode("utf-8"),
        "filename": "roster.txt",
        "mime_type": "text/plain",
        "source_ref": "text:cross-format",
    }
    first_upload, second_upload = (
        (office_upload, text_upload) if office_first else (text_upload, office_upload)
    )

    first = await pipeline.ingest_file(
        "alice",
        None,
        force_review=True,
        metadata={"uploaded_by": "alice"},
        **first_upload,
    )
    second = await pipeline.ingest_file(
        "alice",
        None,
        force_review=True,
        metadata={"uploaded_by": "alice"},
        **second_upload,
    )

    assert second.get("idempotent_replay") is not True
    assert second["raw_object_id"] != first["raw_object_id"]
    assert storage.execute("SELECT COUNT(*) FROM raw_objects WHERE user_id='alice'").fetchone()[0] == 2


@pytest.mark.anyio
async def test_equivalent_valid_office_indexes_keep_the_existing_text_dedup(settings, storage):
    pipeline = _pipeline(settings, storage)
    first = await pipeline.ingest_file(
        "alice",
        None,
        _docx_bytes(interleaved=True, package_note="package-a"),
        filename="roster-a.docx",
        source_ref="office:equivalent:a",
        force_review=True,
        metadata={"uploaded_by": "alice"},
    )
    second = await pipeline.ingest_file(
        "alice",
        None,
        _docx_bytes(interleaved=True, package_note="package-b"),
        filename="roster-b.docx",
        source_ref="office:equivalent:b",
        force_review=True,
        metadata={"uploaded_by": "alice"},
    )

    assert second.get("idempotent_replay") is True
    assert second["raw_object_id"] == first["raw_object_id"]
    assert storage.execute("SELECT COUNT(*) FROM raw_objects WHERE user_id='alice'").fetchone()[0] == 1


@pytest.mark.anyio
async def test_equal_incomplete_office_indexes_never_authorize_text_dedup(settings, storage):
    pipeline = _pipeline(settings, storage)
    first = await pipeline.ingest_file(
        "alice",
        None,
        _docx_bytes(interleaved=True, nested_note="OMITTED-NESTED-ALPHA"),
        filename="roster-incomplete-a.docx",
        source_ref="office:incomplete:a",
        force_review=True,
        metadata={"uploaded_by": "alice"},
    )
    second = await pipeline.ingest_file(
        "alice",
        None,
        _docx_bytes(interleaved=True, nested_note="OMITTED-NESTED-BETA"),
        filename="roster-incomplete-b.docx",
        source_ref="office:incomplete:b",
        force_review=True,
        metadata={"uploaded_by": "alice"},
    )

    first_raw, first_metadata = _raw_metadata(storage, str(first["raw_object_id"]))
    second_raw, second_metadata = _raw_metadata(storage, str(second["raw_object_id"]))
    first_index = first_metadata[_INDEX_KEY]
    second_index = second_metadata[_INDEX_KEY]
    assert first_raw["raw_content"] == second_raw["raw_content"]
    assert first_index == second_index
    assert first_index["complete"] is False
    # Причина сменилась вместе с тем, что читается: колонтитул с 0.190.0 читается
    # и полноты не отнимает, вложенная таблица — по-прежнему нет.
    assert "nested_table" in first_index["coverage"]["reasons"]
    assert second.get("idempotent_replay") is not True
    assert first_raw["id"] != second_raw["id"]
