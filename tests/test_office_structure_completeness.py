"""Synthetic completeness regressions for native Office structure indexes."""

from __future__ import annotations

import copy
import io
import zipfile
from collections.abc import Callable
from types import SimpleNamespace
from typing import Any
from xml.etree import ElementTree

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.comments import Comment

import friday.documents._office_structure as office_structure
from friday.documents import DocumentExtractor, validate_office_structure_index
from friday.documents._office_structure import build_docx_text_and_structure


def _docx_bytes(document: Document) -> bytes:
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _xlsx_bytes(workbook: Workbook) -> bytes:
    output = io.BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


def _rewrite_zip_member(
    content: bytes,
    member_name: str,
    rewrite: Callable[[bytes], bytes],
) -> bytes:
    output = io.BytesIO()
    with (
        zipfile.ZipFile(io.BytesIO(content)) as source,
        zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as target,
    ):
        for item in source.infolist():
            value = source.read(item.filename)
            target.writestr(item, rewrite(value) if item.filename == member_name else value)
    return output.getvalue()


def _index(result: Any) -> dict[str, Any]:
    index = result.office_structure_index
    assert isinstance(index, dict)
    assert validate_office_structure_index(index, result.text) == index
    return index


def _small_roster() -> tuple[str, dict[str, Any]]:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Person", "Role"])
    sheet.append(["SYNTHETIC-ALICE", "Engineer"])
    sheet.append(["SYNTHETIC-BOB", "Manager"])
    result = DocumentExtractor().extract(_xlsx_bytes(workbook), "mutations.xlsx")
    return result.text, _index(result)


def _remove_last_record(index: dict[str, Any]) -> tuple[dict[str, Any], dict[str, Any]]:
    mutated = copy.deepcopy(index)
    sheet = mutated["blocks"][0]
    removed_row = sheet["rows"].pop()
    sheet["text_span"][1] = sheet["rows"][-1]["text_span"][1]
    removed_id = removed_row["id"]
    record_set = mutated["record_sets"][0]
    record_set["record_ids"].remove(removed_id)
    record_set["records_total"] -= 1
    mutated["candidate_refs"] = [
        candidate for candidate in mutated["candidate_refs"] if candidate["record_id"] != removed_id
    ]
    mutated["coverage"]["rows_seen"] -= 1
    mutated["coverage"]["rows_indexed"] -= 1
    mutated["coverage"]["cells_seen"] -= len(removed_row["cells"])
    mutated["coverage"]["cells_indexed"] -= len(removed_row["cells"])
    return mutated, removed_row


def test_validator_rejects_removed_expanded_and_duplicated_rows() -> None:
    text, original = _small_roster()

    removed, _ = _remove_last_record(original)
    assert validate_office_structure_index(removed, text) is None

    expanded, removed_row = _remove_last_record(original)
    alice_row = expanded["blocks"][0]["rows"][-1]
    alice_row["text_span"][1] = removed_row["text_span"][1]
    alice_row["cells"][0]["text_span"][1] = removed_row["cells"][0]["text_span"][1]
    expanded["candidate_refs"][0]["text_span"] = list(alice_row["cells"][0]["text_span"])
    expanded["blocks"][0]["text_span"][1] = removed_row["text_span"][1]
    assert validate_office_structure_index(expanded, text) is None

    duplicated = copy.deepcopy(original)
    sheet = duplicated["blocks"][0]
    duplicate = copy.deepcopy(sheet["rows"][-1])
    duplicate["id"] = f"{sheet['id']}:r000004"
    duplicate["source_row"] = 4
    for column, cell in enumerate(duplicate["cells"], start=1):
        cell["id"] = f"{duplicate['id']}:c{column:06d}"
        cell["coordinate"] = f"{chr(64 + column)}4"
        cell["merge_anchor"] = cell["id"]
    sheet["rows"].append(duplicate)
    duplicated["coverage"]["rows_seen"] += 1
    duplicated["coverage"]["rows_indexed"] += 1
    duplicated["coverage"]["cells_seen"] += len(duplicate["cells"])
    duplicated["coverage"]["cells_indexed"] += len(duplicate["cells"])
    duplicated["record_sets"][0]["record_ids"].append(duplicate["id"])
    duplicated["record_sets"][0]["records_total"] += 1
    duplicate_candidate = copy.deepcopy(duplicated["candidate_refs"][-1])
    duplicate_candidate["id"] = "cand000003"
    duplicate_candidate["record_id"] = duplicate["id"]
    duplicate_candidate["cell_id"] = duplicate["cells"][0]["id"]
    duplicated["candidate_refs"].append(duplicate_candidate)
    assert validate_office_structure_index(duplicated, text) is None


@pytest.mark.parametrize(
    "rows",
    [
        [("SYNTHETIC-ALICE", "Employee"), ("SYNTHETIC-BOB", "Manager")],
        [("Name", "Engineer"), ("SYNTHETIC-BOB", "Manager")],
    ],
)
def test_data_rows_with_people_words_are_not_misclassified_as_headers(
    rows: list[tuple[str, str]],
) -> None:
    workbook = Workbook()
    sheet = workbook.active
    for row in rows:
        sheet.append(row)
    index = _index(DocumentExtractor().extract(_xlsx_bytes(workbook), "not-header.xlsx"))
    assert index["complete"] is True
    assert index["record_sets"] == []
    assert index["candidate_refs"] == []


@pytest.mark.parametrize(
    "container_name",
    [
        "header",
        "even_page_header",
        "first_page_header",
        "footer",
        "even_page_footer",
        "first_page_footer",
    ],
)
def test_every_docx_header_footer_variant_revokes_completeness(container_name: str) -> None:
    document = Document()
    document.add_paragraph("VISIBLE-BODY")
    getattr(document.sections[0], container_name).paragraphs[0].text = "AUXILIARY-TEXT"
    index = _index(DocumentExtractor().extract(_docx_bytes(document), "header-footer.docx"))
    assert index["complete"] is False
    assert "header_footer" in index["coverage"]["reasons"]


@pytest.mark.parametrize(
    "field_xml",
    [
        ('<w:fldSimple {ns} w:instr="DATE"><w:r><w:t>DISPLAYED-FIELD</w:t></w:r></w:fldSimple>'),
        (
            '<w:r {ns}><w:fldChar w:fldCharType="begin"/></w:r>'
            "<w:r {ns}><w:instrText>DATE</w:instrText></w:r>"
            '<w:r {ns}><w:fldChar w:fldCharType="separate"/></w:r>'
            "<w:r {ns}><w:t>DISPLAYED-FIELD</w:t></w:r>"
            '<w:r {ns}><w:fldChar w:fldCharType="end"/></w:r>'
        ),
    ],
)
def test_docx_fields_revoke_completeness(field_xml: str) -> None:
    document = Document()
    paragraph = document.add_paragraph("VISIBLE")
    namespace = nsdecls("w")
    if field_xml.startswith("<w:fldSimple"):
        paragraph._p.append(parse_xml(field_xml.format(ns=namespace)))
    else:
        for fragment in field_xml.format(ns=namespace).split("</w:r>"):
            if fragment:
                paragraph._p.append(parse_xml(f"{fragment}</w:r>"))
    index = _index(DocumentExtractor().extract(_docx_bytes(document), "field.docx"))
    assert index["complete"] is False
    assert "unsupported_body_content" in index["coverage"]["reasons"]


@pytest.mark.parametrize(
    ("fragment", "omitted"),
    [
        (
            "<m:oMath {ns}><m:r><m:t>MATH-OMITTED</m:t></m:r></m:oMath>",
            "MATH-OMITTED",
        ),
        (
            '<w:dir {ns} w:val="rtl"><w:r><w:t>DIR-OMITTED</w:t></w:r></w:dir>',
            "DIR-OMITTED",
        ),
        (
            '<w:bdo {ns} w:val="rtl"><w:r><w:t>BDO-OMITTED</w:t></w:r></w:bdo>',
            "BDO-OMITTED",
        ),
        (
            "<w:ruby {ns}><w:rubyPr/><w:rt><w:r><w:t>RUBY-ANNOTATION</w:t></w:r></w:rt>"
            "<w:rubyBase><w:r><w:t>RUBY-BASE</w:t></w:r></w:rubyBase></w:ruby>",
            "RUBY-BASE",
        ),
    ],
)
def test_docx_visible_wrappers_omitted_by_python_docx_revoke_completeness(
    fragment: str,
    omitted: str,
) -> None:
    document = Document()
    paragraph = document.add_paragraph("VISIBLE")
    paragraph._p.append(parse_xml(fragment.format(ns=nsdecls("w", "m"))))
    result = DocumentExtractor().extract(_docx_bytes(document), "wrapped.docx")
    index = _index(result)
    assert omitted not in result.text
    assert index["complete"] is False
    assert "unsupported_body_content" in index["coverage"]["reasons"]


def test_xlsx_comments_charts_and_headers_revoke_completeness() -> None:
    for feature in ("comment", "chart", "header"):
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["Person", "Value"])
        sheet.append(["SYNTHETIC-ALICE", 3])
        if feature == "comment":
            sheet["A2"].comment = Comment("AUXILIARY-COMMENT", "Synthetic")
        elif feature == "chart":
            chart = BarChart()
            chart.add_data(Reference(sheet, min_col=2, min_row=1, max_row=2), titles_from_data=True)
            sheet.add_chart(chart, "D4")
        else:
            sheet.oddHeader.center.text = "AUXILIARY-HEADER"
        index = _index(DocumentExtractor().extract(_xlsx_bytes(workbook), f"{feature}.xlsx"))
        assert index["complete"] is False
        expected = "header_footer" if feature == "header" else "unsupported_body_content"
        assert expected in index["coverage"]["reasons"]


def test_xlsx_phonetic_text_omitted_by_openpyxl_revokes_completeness() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Person"])
    sheet.append(["SYNTHETIC-ALICE"])
    content = _xlsx_bytes(workbook)

    def add_phonetic_text(xml: bytes) -> bytes:
        root = ElementTree.fromstring(xml)
        namespace = root.tag[1:].split("}", 1)[0]
        for cell in root.iter(f"{{{namespace}}}c"):
            if cell.attrib.get("r") != "A2":
                continue
            inline = next(child for child in cell if child.tag.rsplit("}", 1)[-1] == "is")
            phonetic = ElementTree.SubElement(inline, f"{{{namespace}}}rPh", {"sb": "0", "eb": "5"})
            ElementTree.SubElement(phonetic, f"{{{namespace}}}t").text = "PHONETIC-BOB"
            break
        return ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)

    content = _rewrite_zip_member(content, "xl/worksheets/sheet1.xml", add_phonetic_text)
    result = DocumentExtractor().extract(content, "phonetic.xlsx")
    index = _index(result)
    assert "PHONETIC-BOB" not in result.text
    assert index["complete"] is False
    assert "unsupported_body_content" in index["coverage"]["reasons"]


@pytest.mark.parametrize("layout", ["hidden_row", "hidden_column", "outline", "filter"])
def test_xlsx_hidden_or_filtered_layout_revokes_completeness(layout: str) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Person", "Role"])
    sheet.append(["SYNTHETIC-ALICE", "Engineer"])
    if layout == "hidden_row":
        sheet.row_dimensions[2].hidden = True
    elif layout == "hidden_column":
        sheet.column_dimensions["A"].hidden = True
    elif layout == "outline":
        sheet.row_dimensions[2].outlineLevel = 1
    else:
        sheet.auto_filter.ref = "A1:B2"
        sheet.auto_filter.add_filter_column(0, ["SYNTHETIC-ALICE"])
    index = _index(DocumentExtractor().extract(_xlsx_bytes(workbook), f"{layout}.xlsx"))
    assert index["complete"] is False
    assert "hidden_layout" in index["coverage"]["reasons"]


@pytest.mark.parametrize(
    "footer_value",
    ["Grand Total", "Итог", "Общий итог", "TOTALS", "Average", "Среднее", "16"],
)
def test_footer_like_terminal_rows_never_become_people(footer_value: str) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Person", "Role"])
    sheet.append(["SYNTHETIC-ALICE", "Total Quality Manager"])
    sheet.append([footer_value, ""])
    index = _index(DocumentExtractor().extract(_xlsx_bytes(workbook), "footer.xlsx"))
    assert index["complete"] is True
    assert index["record_sets"] == []
    assert index["candidate_refs"] == []


def test_xlsx_cached_formula_is_complete_when_visible_value_is_proven() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Person", "Value"])
    sheet.append(["SYNTHETIC-ALICE", "=1+1"])
    content = _xlsx_bytes(workbook)

    def add_cached_value(xml: bytes) -> bytes:
        root = ElementTree.fromstring(xml)
        for cell in root.iter():
            if cell.tag.rsplit("}", 1)[-1] != "c" or cell.attrib.get("r") != "B2":
                continue
            cached = next(child for child in cell if child.tag.rsplit("}", 1)[-1] == "v")
            cached.text = "2"
            break
        return ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)

    content = _rewrite_zip_member(content, "xl/worksheets/sheet1.xml", add_cached_value)
    result = DocumentExtractor().extract(content, "cached-formula.xlsx")
    index = _index(result)
    assert result.text.endswith("SYNTHETIC-ALICE | 2")
    assert index["complete"] is True
    assert not any("formula" in reason for reason in index["coverage"]["reasons"])


def test_xlsx_stale_dimension_and_overlapping_merges_fail_closed() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Person", "Value"])
    sheet["A100"] = "OMITTED-BY-STALE-DIMENSION"
    sheet["B100"] = "=1+1"
    content = _xlsx_bytes(workbook)

    def stale_dimension(xml: bytes) -> bytes:
        assert b'ref="A1:B100"' in xml
        return xml.replace(b'ref="A1:B100"', b'ref="A1:B2"', 1)

    stale = _rewrite_zip_member(content, "xl/worksheets/sheet1.xml", stale_dimension)
    stale_result = DocumentExtractor().extract(stale, "stale-dimension.xlsx")
    stale_index = _index(stale_result)
    assert "OMITTED-BY-STALE-DIMENSION" not in stale_result.text
    assert stale_index["complete"] is False
    assert "formula_alignment" in stale_index["coverage"]["reasons"]

    merged = Workbook()
    sheet = merged.active
    sheet.append(["Person", "B", "C", "D"])
    sheet.append(["SYNTHETIC-ALICE", "ONE", "TWO", "THREE"])
    sheet.merge_cells("B2:C2")
    sheet.merge_cells("C2:D2")
    merged_index = _index(DocumentExtractor().extract(_xlsx_bytes(merged), "overlap.xlsx"))
    assert merged_index["complete"] is False
    assert "merge_topology" in merged_index["coverage"]["reasons"]


def test_xlsx_cell_visit_budget_stops_before_the_next_wide_row(monkeypatch) -> None:
    monkeypatch.setattr(office_structure, "_MAX_XLSX_VISITED_CELLS", 4)
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Person", "Value"])
    sheet.append(["SYNTHETIC-ALICE", "VISIBLE"])
    sheet.append(["TAIL-FORMULA", "=1+1"])
    result = DocumentExtractor().extract(_xlsx_bytes(workbook), "cell-budget.xlsx")
    index = _index(result)
    assert "SYNTHETIC-ALICE" in result.text
    assert "TAIL-FORMULA" not in result.text
    assert result.metadata["rows_read"] == 2
    assert result.metadata["extraction_truncated"] is True
    assert "row_budget" in index["coverage"]["reasons"]
    assert "formula_without_cached_value" not in index["coverage"]["reasons"]


def test_xlsx_real_cell_visit_cap_bounds_a_max_width_declared_grid() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet["A1"] = "VISIBLE-PREFIX"
    sheet["XFD1000"] = "TAIL-OUTSIDE-CELL-BUDGET"
    result = DocumentExtractor().extract(_xlsx_bytes(workbook), "wide-grid.xlsx")
    index = _index(result)
    assert "VISIBLE-PREFIX" in result.text
    assert "TAIL-OUTSIDE-CELL-BUDGET" not in result.text
    assert result.metadata["rows_read"] <= office_structure._MAX_XLSX_VISITED_CELLS // 16_384
    assert result.metadata["extraction_truncated"] is True
    assert "row_budget" in index["coverage"]["reasons"]


def test_dimensionless_write_only_xlsx_uses_bounded_xml_extents_without_losing_rows() -> None:
    workbook = Workbook(write_only=True)
    sheet = workbook.create_sheet("Data")
    for number in range(1, 7):
        sheet.append((f"ROW-{number}", f"VALUE-{number}"))
    result = DocumentExtractor().extract(_xlsx_bytes(workbook), "dimensionless.xlsx")
    index = _index(result)
    assert result.text == "\n".join(
        [
            "--- Sheet: Data ---",
            *(f"ROW-{number} | VALUE-{number}" for number in range(1, 7)),
        ]
    )
    assert result.metadata["rows_read"] == 6
    assert index["complete"] is True
    assert "formula_alignment" not in index["coverage"]["reasons"]


def test_multiline_office_values_keep_an_incomplete_but_bound_index() -> None:
    document = Document()
    paragraph = document.add_paragraph("LINE-ONE")
    paragraph.add_run().add_break()
    paragraph.add_run("LINE-TWO")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "CELL-ONE"
    table.cell(0, 0).add_paragraph("CELL-TWO")
    docx_result = DocumentExtractor().extract(_docx_bytes(document), "multiline.docx")
    docx_index = _index(docx_result)
    assert "LINE-ONE\nLINE-TWO" in docx_result.text
    assert "CELL-ONE\nCELL-TWO" in docx_result.text
    assert docx_index["complete"] is False
    assert "unsupported_body_content" in docx_index["coverage"]["reasons"]

    workbook = Workbook()
    workbook.active["A1"] = "CELL-ONE\nCELL-TWO"
    xlsx_result = DocumentExtractor().extract(_xlsx_bytes(workbook), "multiline.xlsx")
    xlsx_index = _index(xlsx_result)
    assert "CELL-ONE\nCELL-TWO" in xlsx_result.text
    assert xlsx_index["complete"] is False
    assert "unsupported_body_content" in xlsx_index["coverage"]["reasons"]


def test_docx_text_budget_does_not_visit_the_ignored_tail_or_package() -> None:
    class Element:
        tag = "p"

    first_element = Element()

    class Body:
        @staticmethod
        def iterchildren():
            yield first_element

        @staticmethod
        def iter():
            raise AssertionError("completeness scan traversed the ignored tail")

    class FakeDocument:
        element = SimpleNamespace(body=Body())

        @property
        def paragraphs(self):
            yield SimpleNamespace(text="X" * 20_000, _element=first_element, style=None)
            raise AssertionError("paragraph after the text boundary was visited")

        @property
        def tables(self):
            raise AssertionError("table tail was materialized after the text boundary")

    text, index, truncated = build_docx_text_and_structure(
        FakeDocument(),
        max_text_chars=10_000,
        content=b"not-even-a-zip-because-it-must-not-be-scanned",
    )
    assert text == "X" * 10_000
    assert truncated is True
    assert index["coverage"]["reasons"] == ["text_budget"]
    assert validate_office_structure_index(index, text) == index
