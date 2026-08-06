"""OfficeStructureIndex v1: exact text, content-free structure, fail-closed coverage."""

from __future__ import annotations

import copy
import hashlib
import io
import json
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from openpyxl import Workbook

from friday.documents import DocumentExtractor, DocumentResult, validate_office_structure_index
from friday.documents._office_structure import OFFICE_STRUCTURE_MAX_SERIALIZED_BYTES


def _docx_bytes(document: Document) -> bytes:
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _xlsx_bytes(workbook: Workbook) -> bytes:
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def _span_text(text: str, span: list[int]) -> str:
    return text[span[0] : span[1]]


def _table_block(index: dict) -> dict:
    return next(block for block in index["blocks"] if block["kind"] in {"table", "sheet"})


def _sixteen_row_docx(*, merged_header: bool = False) -> bytes:
    document = Document()
    first = document.add_paragraph()
    first.add_run("IN")
    first.add_run("TRO")
    table = document.add_table(rows=17, cols=2)
    if merged_header:
        table.cell(0, 0).merge(table.cell(0, 1)).text = "ФИО"
    else:
        table.cell(0, 0).text = "ФИО"
        table.cell(0, 1).text = "Роль"
    names = [f"PERSON-{number:02d}" for number in range(1, 16)] + ["PERSON-01"]
    for row_number, name in enumerate(names, start=1):
        person_cell = table.cell(row_number, 0)
        person_cell.text = ""
        person_cell.paragraphs[0].add_run(name[:6])
        person_cell.paragraphs[0].add_run(name[6:])
        table.cell(row_number, 1).text = f"ROLE-{row_number:02d}"
    document.add_paragraph("OUTRO")
    return _docx_bytes(document)


def test_docx_keeps_legacy_text_but_indexes_true_source_order_and_exact_spans():
    result = DocumentExtractor().extract(_sixteen_row_docx(), "positions.docx")
    rows = [
        "ФИО | Роль",
        *(f"PERSON-{number:02d} | ROLE-{number:02d}" for number in range(1, 16)),
        "PERSON-01 | ROLE-16",
    ]
    expected = "\n".join(["INTRO", "OUTRO", *rows])

    assert result.text == expected
    assert result.text.encode("utf-8") == expected.encode("utf-8")
    index = result.office_structure_index
    assert index is not None
    assert index["text_sha256"] == hashlib.sha256(expected.encode("utf-8")).hexdigest()
    assert validate_office_structure_index(index, result.text) == index

    # The old corpus text remains paragraphs-first/table-last. The independent
    # index restores the actual body order without moving one byte in that text.
    assert [block["kind"] for block in index["blocks"]] == ["paragraph", "table", "paragraph"]
    assert [block["source_order"] for block in index["blocks"]] == [0, 1, 2]
    assert [_span_text(result.text, block["text_span"]) for block in index["blocks"]] == [
        "INTRO",
        "\n".join(rows),
        "OUTRO",
    ]
    first_runs = index["blocks"][0]["runs"]
    assert [_span_text(result.text, run["text_span"]) for run in first_runs] == ["IN", "TRO"]

    record_set = index["record_sets"][0]
    assert record_set["authoritative"] is True
    assert record_set["kind"] == "person_rows"
    assert record_set["records_total"] == 16
    assert len(record_set["record_ids"]) == 16
    assert len(index["candidate_refs"]) == 16
    assert all(candidate["type"] == "person" for candidate in index["candidate_refs"])
    assert all(candidate["basis"] == "declared_person_column" for candidate in index["candidate_refs"])
    values = [_span_text(result.text, candidate["text_span"]) for candidate in index["candidate_refs"]]
    assert len(values) == 16
    assert len({" ".join(value.casefold().split()) for value in values}) == 15


def test_xlsx_keeps_legacy_text_and_sheet_title_only_as_a_span():
    workbook = Workbook()
    sheet = workbook.active
    title = "PRIVATE-SHEET-TITLE-SYNTHETIC"
    sheet.title = title
    sheet.append(["ФИО", "Роль"])
    for number in range(1, 17):
        sheet.append([f"XLSX-PERSON-{number:02d}", f"R-{number:02d}"])

    result = DocumentExtractor().extract(_xlsx_bytes(workbook), "positions.xlsx")
    expected = "\n".join(
        [
            f"--- Sheet: {title} ---",
            "ФИО | Роль",
            *(f"XLSX-PERSON-{number:02d} | R-{number:02d}" for number in range(1, 17)),
        ]
    )
    assert result.text == expected
    index = result.office_structure_index
    assert index is not None and index["complete"] is True
    assert validate_office_structure_index(index, result.text) == index
    sheet_block = _table_block(index)
    assert sheet_block["kind"] == "sheet"
    assert _span_text(result.text, sheet_block["title_span"]) == title
    assert index["record_sets"][0]["records_total"] == 16
    assert len(index["candidate_refs"]) == 16

    # A sheet title and cell values are private literals. They occur once in
    # raw_content and nowhere in the durable index JSON.
    encoded_index = json.dumps(index, ensure_ascii=False, sort_keys=True)
    assert title not in encoded_index
    assert "XLSX-PERSON-16" not in encoded_index


def test_custom_style_and_cell_literals_never_enter_the_index():
    document = Document()
    style_secret = "PRIVATE-CUSTOM-STYLE-SYNTHETIC"
    custom_style = document.styles.add_style(style_secret, WD_STYLE_TYPE.PARAGRAPH)
    paragraph = document.add_paragraph("PRIVATE-PARAGRAPH-VALUE-SYNTHETIC")
    paragraph.style = custom_style
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "ФИО"
    table.cell(1, 0).text = "PRIVATE-PERSON-VALUE-SYNTHETIC"

    result = DocumentExtractor().extract(_docx_bytes(document), "private.docx")
    index = result.office_structure_index
    assert index is not None
    encoded_index = json.dumps(index, ensure_ascii=False, sort_keys=True)
    assert style_secret not in encoded_index
    assert "PRIVATE-PARAGRAPH-VALUE-SYNTHETIC" not in encoded_index
    assert "PRIVATE-PERSON-VALUE-SYNTHETIC" not in encoded_index
    assert index["blocks"][0]["style_role"] == "other"


def test_horizontally_merged_person_header_is_header_but_not_an_authoritative_set():
    result = DocumentExtractor().extract(
        _sixteen_row_docx(merged_header=True),
        "merged-header.docx",
    )
    index = result.office_structure_index
    assert index is not None
    table = _table_block(index)
    header = table["rows"][0]
    assert header["role"] == "header"
    assert header["cells"][0]["merge_anchor"] == header["cells"][1]["merge_anchor"]
    assert index["record_sets"] == []
    assert index["candidate_refs"] == []
    assert validate_office_structure_index(index, result.text) == index

    workbook = Workbook()
    sheet = workbook.active
    sheet.merge_cells("A1:B1")
    sheet["A1"] = "ФИО"
    sheet.append(["XLSX-MERGED-PERSON", "ROLE"])
    xlsx_result = DocumentExtractor().extract(_xlsx_bytes(workbook), "merged-header.xlsx")
    xlsx_index = xlsx_result.office_structure_index
    assert xlsx_index is not None
    assert _table_block(xlsx_index)["rows"][0]["role"] == "header"
    assert xlsx_index["record_sets"] == []
    assert xlsx_index["candidate_refs"] == []


def test_nested_table_late_header_footer_and_text_box_make_docx_coverage_incomplete():
    document = Document()
    document.add_paragraph("BODY")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).add_table(rows=1, cols=1).cell(0, 0).text = "NESTED-OMITTED"

    second = document.add_section(WD_SECTION.NEW_PAGE)
    second.header.is_linked_to_previous = False
    second.footer.is_linked_to_previous = False
    second.header.paragraphs[0].text = "LATE-HEADER-OMITTED"
    second.footer.paragraphs[0].text = "LATE-FOOTER-OMITTED"

    paragraph = document.add_paragraph("VISIBLE")
    paragraph._p.append(
        parse_xml(
            f"<w:r {nsdecls('w')}><w:pict><w:txbxContent><w:p><w:r>"
            "<w:t>TEXT-BOX-OMITTED</w:t></w:r></w:p></w:txbxContent></w:pict></w:r>"
        )
    )

    result = DocumentExtractor().extract(_docx_bytes(document), "omissions.docx")
    index = result.office_structure_index
    assert index is not None and index["complete"] is False
    assert {"nested_table", "header_footer", "text_box"} <= set(index["coverage"]["reasons"])
    assert "NESTED-OMITTED" not in result.text
    assert "LATE-HEADER-OMITTED" not in result.text
    assert "TEXT-BOX-OMITTED" not in result.text


def test_tracked_change_and_footnote_part_fail_closed_without_copying_their_text():
    document = Document()
    paragraph = document.add_paragraph("VISIBLE-BODY")
    paragraph._p.append(
        parse_xml(f'<w:ins {nsdecls("w")} w:id="7"><w:r><w:t>TRACKED-OMITTED</w:t></w:r></w:ins>')
    )
    original = _docx_bytes(document)
    output = io.BytesIO()
    with (
        zipfile.ZipFile(io.BytesIO(original)) as source,
        zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as target,
    ):
        for item in source.infolist():
            target.writestr(item, source.read(item.filename))
        target.writestr(
            "word/footnotes.xml",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                f'<w:footnotes {nsdecls("w")}><w:footnote w:id="1"><w:p><w:r>'
                "<w:t>FOOTNOTE-OMITTED</w:t></w:r></w:p></w:footnote></w:footnotes>"
            ),
        )

    result = DocumentExtractor().extract(output.getvalue(), "tracked.docx")
    index = result.office_structure_index
    assert result.text == "VISIBLE-BODY"
    assert index is not None and index["complete"] is False
    assert "unsupported_body_content" in index["coverage"]["reasons"]
    encoded_index = json.dumps(index, ensure_ascii=False)
    assert "TRACKED-OMITTED" not in encoded_index
    assert "FOOTNOTE-OMITTED" not in encoded_index


def test_xlsx_formula_without_cached_result_is_an_explicit_coverage_reason():
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["ФИО", "Сумма"])
    sheet.append(["FORMULA-PERSON", "=1+1"])

    result = DocumentExtractor().extract(_xlsx_bytes(workbook), "formula.xlsx")
    assert result.text.endswith("FORMULA-PERSON | ")
    index = result.office_structure_index
    assert index is not None and index["complete"] is False
    assert "formula_without_cached_value" in index["coverage"]["reasons"]
    assert validate_office_structure_index(index, result.text) == index


def test_text_budget_keeps_exact_legacy_prefix_and_marks_both_metadata_and_index():
    document = Document()
    document.add_paragraph("X" * 10_100)
    document.add_paragraph("AFTER-BUDGET")
    result = DocumentExtractor(max_text_chars=10_000).extract(
        _docx_bytes(document),
        "large.docx",
    )

    assert result.text == "X" * 10_000
    assert result.metadata["extraction_truncated"] is True
    assert result.metadata["text_truncated"] is True
    index = result.office_structure_index
    assert index is not None and index["complete"] is False
    assert "text_budget" in index["coverage"]["reasons"]
    assert validate_office_structure_index(index, result.text) == index


def test_large_index_is_bounded_by_dropping_only_whole_rows():
    workbook = Workbook(write_only=True)
    sheet = workbook.create_sheet("Data")
    sheet.append(("ФИО", "Роль"))
    for number in range(400):
        sheet.append((f"BOUNDED-PERSON-{number:04d}", f"ROLE-{number:04d}"))

    result = DocumentExtractor().extract(_xlsx_bytes(workbook), "bounded.xlsx")
    index = result.office_structure_index
    assert index is not None and index["complete"] is False
    assert "index_budget" in index["coverage"]["reasons"]
    serialized = json.dumps(index, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode()
    assert len(serialized) <= OFFICE_STRUCTURE_MAX_SERIALIZED_BYTES
    assert validate_office_structure_index(index, result.text) == index
    for block in index["blocks"]:
        for row in block.get("rows", []):
            if row["text_span"] is None:
                continue
            assert all(cell["text_span"] is not None for cell in row["cells"])


def test_validator_rejects_structural_and_coordinated_inventory_mutations():
    result = DocumentExtractor().extract(_sixteen_row_docx(), "mutations.docx")
    base = result.office_structure_index
    assert base is not None and base["complete"] is True

    mutations: list[dict] = []

    wrong_hash = copy.deepcopy(base)
    wrong_hash["text_sha256"] = "0" * 64
    mutations.append(wrong_hash)

    lost_source_order = copy.deepcopy(base)
    del lost_source_order["blocks"][0]["source_order"]
    mutations.append(lost_source_order)

    escaped_span = copy.deepcopy(base)
    escaped_span["candidate_refs"][0]["text_span"] = [0, len(result.text) + 1]
    mutations.append(escaped_span)

    missing_candidate = copy.deepcopy(base)
    missing_candidate["candidate_refs"].pop()
    mutations.append(missing_candidate)

    wrong_person_column = copy.deepcopy(base)
    first_candidate = wrong_person_column["candidate_refs"][0]
    table = _table_block(wrong_person_column)
    first_record = next(row for row in table["rows"] if row["id"] == first_candidate["record_id"])
    other_cell = next(cell for cell in first_record["cells"] if cell["column"] == 2)
    first_candidate["cell_id"] = other_cell["id"]
    first_candidate["text_span"] = other_cell["text_span"]
    mutations.append(wrong_person_column)

    # A coordinated, internally consistent-looking 15/16 subset must still
    # fail: exact composition is re-derived from the bound source text.
    coordinated_subset = copy.deepcopy(base)
    record_set = coordinated_subset["record_sets"][0]
    removed_record_id = record_set["record_ids"].pop()
    record_set["records_total"] = 15
    coordinated_subset["candidate_refs"] = [
        candidate
        for candidate in coordinated_subset["candidate_refs"]
        if candidate["record_id"] != removed_record_id
    ]
    coordinated_table = _table_block(coordinated_subset)
    next(row for row in coordinated_table["rows"] if row["id"] == removed_record_id)["role"] = "unknown"
    mutations.append(coordinated_subset)

    unknown_literal_field = copy.deepcopy(base)
    unknown_literal_field["private_copy"] = "MUST-NOT-BE-ACCEPTED"
    mutations.append(unknown_literal_field)

    wrong_schema = copy.deepcopy(base)
    wrong_schema["schema_version"] = 2
    mutations.append(wrong_schema)

    for mutation in mutations:
        assert validate_office_structure_index(mutation, result.text) is None
    assert validate_office_structure_index(base, result.text + "\ud800") is None


def test_document_result_positional_contract_stays_compatible():
    result = DocumentResult("text", {"format": "synthetic"}, False, "synthetic-error")
    assert result.text == "text"
    assert result.metadata == {"format": "synthetic"}
    assert result.success is False
    assert result.error == "synthetic-error"
    assert result.office_structure_index is None
    assert "office_structure_index" not in result.to_dict()
